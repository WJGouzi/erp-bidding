"""版面感知的文档解析器。

支持 DOCX（标题层级 + 表格结构）和 PDF（fitz 文本页 + PaddleOCR 扫描页）的
结构化解析，输出统一的结构化文档模型 JSON。
"""

import json
import logging
import tempfile
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import fitz

logger = logging.getLogger(__name__)



# ── 表格列宽和合并单元格提取辅助函数（用于序列化/反序列化） ──


def _extract_table_column_widths(table):
    """从 python-docx Table 提取列宽（twips）。"""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    column_widths = []
    try:
        tbl = getattr(table, '_tbl', None)
        if tbl is not None:
            tbl_grid = tbl.find(f'.//{ns}tblGrid')
            if tbl_grid is not None:
                for gc in tbl_grid.findall(f'{ns}gridCol'):
                    w = gc.get(f'{ns}w')
                    if w:
                        column_widths.append(int(w))
    except Exception:
        pass
    return column_widths


def _extract_table_row_heights(table):
    """从 python-docx Table 提取每行高度。

    Returns:
        list[dict]: [{"val": int, "rule": str}, ...]
        val=0 表示行高未指定，rule 通常为 'atLeast' 或 'exactly'
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    heights = []
    for row in table.rows:
        try:
            tr_pr = row._tr.find(f'{ns}trPr')
            if tr_pr is not None:
                tr_height = tr_pr.find(f'{ns}trHeight')
                if tr_height is not None:
                    heights.append({
                        "val": int(tr_height.get(f'{ns}val', 0)),
                        "rule": tr_height.get(f'{ns}rule', 'atLeast'),
                    })
                    continue
        except Exception:
            pass
        heights.append({"val": 0, "rule": "atLeast"})
    return heights


def _extract_table_merge_cells(table):
    """从 python-docx Table 提取合并单元格信息（基于 XML <tc> 直接遍历，无重复记录）。

    直接遍历 <w:tr> 下的 <w:tc> XML 元素，而非 python-docx 虚拟单元格。
    每个物理 <w:tc> 只出现一次，避免 gridSpan 被虚拟单元格重复读取。
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    merges = []
    _covered_vmerge_ranges = {}
    try:
        for row_idx, row in enumerate(table.rows):
            tr = row._tr
            tc_elements = tr.findall(f'{ns}tc')
            col_idx = 0
            for tc in tc_elements:
                try:
                    tcPr = tc.find(f'{ns}tcPr')
                    if tcPr is None:
                        col_idx += 1
                        continue

                    gs = tcPr.find(f'{ns}gridSpan')
                    span_val = int(gs.get(f'{ns}val', '1')) if gs is not None else 1

                    if gs is not None:
                        if span_val > 1:
                            merges.append({"type": "horizontal", "row": row_idx, "col": col_idx, "span": span_val})

                    vm = tcPr.find(f'{ns}vMerge')
                    if vm is not None:
                        v_val = vm.get(f'{ns}val', 'continue')
                        if v_val == 'restart' or v_val is None:
                            merge_count = 1
                            cell_text = table.rows[row_idx].cells[col_idx].text.strip()
                            for nr in range(row_idx + 1, len(table.rows)):
                                _vstop = False
                                try:
                                    next_cell = table.rows[nr].cells[col_idx]
                                    nv = next_cell._tc.find(f'.//{ns}vMerge')
                                    if nv is not None:
                                        nv_val = nv.get(f'{ns}val', 'continue')
                                        if nv_val == 'continue':
                                            merge_count += 1
                                        elif nv_val == 'restart':
                                            if next_cell.text.strip() == cell_text:
                                                merge_count += 1
                                            else:
                                                _vstop = True
                                        else:
                                            _vstop = True
                                    else:
                                        _vstop = True
                                except Exception:
                                    _vstop = True
                                if _vstop:
                                    break
                            if merge_count > 1:
                                already_covered = False
                                for vr, vspan in _covered_vmerge_ranges.get(col_idx, []):
                                    if vr <= row_idx < vr + vspan:
                                        already_covered = True
                                        break
                                if not already_covered:
                                    merges.append({"type": "vertical", "row": row_idx, "col": col_idx, "span": merge_count})
                                    _covered_vmerge_ranges.setdefault(col_idx, []).append((row_idx, merge_count))

                    col_idx += span_val
                except Exception:
                    col_idx += 1
    except Exception:
        pass
    return merges
class StructuredDocument:
    """结构化文档模型，统一表示 DOCX/PDF 的解析结果。"""

    def __init__(self, file_name="", file_sha256="", parse_version="1.0"):
        self.file_name = file_name
        self.file_sha256 = file_sha256
        self.parse_version = parse_version
        self.sections = []  # list[Section]
        self.tables = []  # list of python-docx Table objects

    def to_dict(self) -> dict:
        return {
            "file_name": self.file_name,
            "file_sha256": self.file_sha256,
            "parse_version": self.parse_version,
            "sections": [s.to_dict() for s in self.sections],
            "tables": [
                {
                    "headers": (
                        # python-docx Table: 表头在第一行
                        [cell.text.strip() for cell in t.rows[0].cells]
                        if hasattr(t, "rows") and not hasattr(t, "headers")
                        # TableStub: headers 属性直接可用
                        else list(t.headers)
                    ),
                    "rows": (
                        # python-docx Table: 从第二行开始取数据
                        [[cell.text.strip() for cell in row.cells] for row in list(t.rows)[1:]]
                        if hasattr(t, "rows") and not hasattr(t, "headers")
                        # TableStub: rows 属性直接可用
                        else [list(row) for row in t.rows]
                    ),
                    "column_widths": (
                        _extract_table_column_widths(t)
                        if hasattr(t, "_tbl")
                        else (list(t.column_widths) if hasattr(t, "column_widths") else [])
                    ),
                    "merge_cells": (
                        _extract_table_merge_cells(t)
                        if hasattr(t, "_tbl")
                        else (list(t.merge_cells) if hasattr(t, "merge_cells") else [])
                    ),
                    "row_heights": (
                        _extract_table_row_heights(t)
                        if hasattr(t, "_tbl")
                        else (list(t.row_heights) if hasattr(t, "row_heights") else [])
                    ),
                }
                for t in self.tables
            ] if self.tables else [],
        }

    def to_json(self) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=False)

    def to_text(self) -> str:
        """提取纯文本内容，递归遍历章节和子章节。"""
        texts = []
        for section in self.sections:
            texts.extend(self._section_to_texts(section))
        return "\n".join(texts)

    @staticmethod
    def _section_to_texts(section) -> "list[str]":
        result = []
        if section.title:
            result.append(section.title)
        for block in section.content:
            if block.type in (ContentBlock.TYPE_PARAGRAPH, ContentBlock.TYPE_HEADING, ContentBlock.TYPE_LIST):
                if block.text:
                    result.append(block.text)
            elif block.type == ContentBlock.TYPE_TABLE:
                parts = []
                if block.headers:
                    parts.append(" | ".join(block.headers))
                for row in block.rows:
                    parts.append(" | ".join(row))
                if parts:
                    result.append("\n".join(parts))
        for child in section.children:
            result.extend(StructuredDocument._section_to_texts(child))
        return result


    def build_section_index(self) -> list:
        """构建扁平章节索引。

        递归遍历所有 section，输出扁平列表，每项包含：
        - id: 章节唯一标识（如 "sec_1", "sec_2_1"）
        - title: 章节标题
        - level: 章节层级
        - page_range: 页码范围
        - parent_id: 父章节 id（根节点为 None）
        - children_ids: 子章节 id 列表

        Returns:
            list[dict]: 扁平章节索引
        """
        index = []
        counter = [0]  # 使用 list 实现闭包可变

        def _walk(sections, parent_id=None):
            for section in sections:
                counter[0] += 1
                sec_id = f"sec_{counter[0]}"
                entry = {
                    "id": sec_id,
                    "title": section.title if section.title else "",
                    "level": section.level,
                    "page_range": section.page_range or [],
                    "parent_id": parent_id,
                    "children_ids": [],
                }
                # 先添加 entry 到 index，再处理 children
                index.append(entry)
                if section.children:
                    child_ids = []
                    _walk(section.children, parent_id=sec_id)
                    for child in section.children:
                        # 在 index 中找到这个 child
                        for e in index:
                            if e["title"] == child.title and e["parent_id"] == sec_id:
                                child_ids.append(e["id"])
                                break
                    entry["children_ids"] = child_ids

        _walk(self.sections)
        # 后处理去重：同名+同级+同父节点 → 保留有内容的，删掉空的
        index = _dedup_section_index(index)
        return index

    @classmethod
    def from_dict(cls, data: dict) -> "StructuredDocument":
        doc = cls(data.get("file_name", ""), data.get("file_sha256", ""), data.get("parse_version", "1.0"))
        for s_data in data.get("sections", []):
            section = Section.from_dict(s_data)
            doc.sections.append(section)
        # 表格数据以纯文本形式缓存（不可序列化 python-docx 原生对象）
        # 缓存的表格数据在 to_dict 中已转为 headers/rows 格式
        table_data = data.get("tables", [])
        if table_data:
            from collections import namedtuple
            class TableStub:
                """Stub for table data when loaded from cache."""
                def __init__(self, headers, rows, column_widths=None, merge_cells=None, row_heights=None):
                    self.headers = headers
                    self.rows = rows
                    self.column_widths = column_widths or []
                    self.merge_cells = merge_cells or []
                    self.row_heights = row_heights or []
            doc.tables = [
                TableStub(
                    headers=t.get("headers", []),
                    rows=t.get("rows", []),
                    column_widths=t.get("column_widths", []),
                    merge_cells=t.get("merge_cells", []),
                    row_heights=t.get("row_heights", []),
                )
                for t in table_data
            ]
        return doc


class Section:
    """文档中的一个章节或区块。"""

    def __init__(self, title="", level=1, page_range=None):
        self.title = title
        self.level = level  # 1=一级标题, 2=二级标题, ...
        self.content = []   # list[ContentBlock]
        self.children = []  # list[Section]（子章节）
        self.page_range = page_range or []  # [start_page, end_page]

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "level": self.level,
            "content": [c.to_dict() for c in self.content],
            "children": [c.to_dict() for c in self.children],
            "page_range": self.page_range,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "Section":
        s = cls(data.get("title", ""), data.get("level", 1), data.get("page_range", []))
        for c_data in data.get("content", []):
            s.content.append(ContentBlock.from_dict(c_data))
        for c_data in data.get("children", []):
            s.children.append(cls.from_dict(c_data))
        return s


class ContentBlock:
    """文档内容块（段落、表格等）。"""

    TYPE_PARAGRAPH = "paragraph"
    TYPE_TABLE = "table"
    TYPE_HEADING = "heading"
    TYPE_LIST = "list"

    def __init__(self, type_="paragraph", text="", level=0):
        self.type = type_
        self.text = text
        self.level = level  # 列表缩进层级
        # 表格专用字段
        self.row_heights = []
        self.per_cell_data = None  # 唯一存储：per_cell 模型数据
        # 字体元数据（封面渲染用）
        self.font_name = ""         # 字体名称，如"宋体""黑体"
        self.font_size = None       # 字号（Pt），如 16.0
        self.bold = False           # 是否加粗
        self.alignment = None       # 对齐方式: left/center/right

    # -----------------------------------------------------------
    #  Stage 2: aggregated column widths / merge_map / cell bounds
    # -----------------------------------------------------------

    @property
    def actual_column_widths(self):
        if not self.column_widths:
            return []
        if not self.merge_cells:
            return list(self.column_widths)

        h_merges = [m for m in self.merge_cells if m['type'] == 'horizontal']
        if not h_merges:
            return list(self.column_widths)

        row0_h = sorted([m for m in h_merges if m['row'] == 0], key=lambda x: x['col'])
        if not row0_h:
            return list(self.column_widths)

        widths = []
        ci = 0
        merged_idx = 0
        while ci < len(self.column_widths):
            if merged_idx < len(row0_h) and ci == row0_h[merged_idx]['col']:
                span = row0_h[merged_idx]['span']
                widths.append(sum(self.column_widths[ci:ci + span]))
                ci += span
                merged_idx += 1
            else:
                widths.append(self.column_widths[ci])
                ci += 1
        return widths

    @property
    def merge_map(self):
        mm = {}
        for m in self.merge_cells:
            key = (m['row'], m['col'])
            if m['type'] == 'horizontal':
                mm.setdefault(key, {})['h_span'] = m['span']
            elif m['type'] == 'vertical':
                mm.setdefault(key, {})['v_span'] = m['span']
        return mm

    def get_cell_bounds(self, row, col):
        mm = self.merge_map
        key = (row, col)
        info = mm.get(key, {})

        col_span = info.get('h_span', 1)
        row_span = info.get('v_span', 1)

        cw = self.column_widths or []
        actual_w = sum(cw[col:col + col_span]) if col < len(cw) else 0

        rh = []
        for h in (self.row_heights or []):
            if isinstance(h, dict):
                rh.append(h.get('val', 0))
            else:
                rh.append(h if isinstance(h, int) else 0)
        actual_h = sum(rh[row:row + row_span]) if row < len(rh) else 0

        return {
            'col_span': col_span,
            'row_span': row_span,
            'width': actual_w,
            'height': actual_h,
            'is_hidden': col_span == 0 and row_span == 0,
        }

    def get_width_matrix(self):
        if not self.rows:
            return []
        nrows = len(self.rows)
        ncols = max(len(r) for r in self.rows) if self.rows else 0
        mm = self.merge_map
        cw = self.actual_column_widths

        matrix = []
        for r in range(nrows):
            row_w = []
            col_idx = 0
            while col_idx < min(ncols, len(cw)):
                info = mm.get((r, col_idx), {})
                span = info.get('h_span', 1)
                w = sum(cw[col_idx:col_idx + span]) if col_idx < len(cw) else 0
                row_w.append(w)
                col_idx += span
            matrix.append(row_w)
        return matrix

    # ──────────────────────────────────────────────────────
    #  向后兼容属性：优先从 per_cell_data 读取，降级到内部 _storage
    # ──────────────────────────────────────────────────────

    @property
    def headers(self):
        if self.per_cell_data:
            rows = self.per_cell_data.get('rows', [])
            if rows:
                return [c.get('text', '') for c in rows[0].get('cells', [])]
        return getattr(self, '_headers', [])

    @headers.setter
    def headers(self, value):
        self._headers = value

    @property
    def rows(self):
        if self.per_cell_data:
            rows = self.per_cell_data.get('rows', [])
            return [[c.get('text', '') for c in r.get('cells', [])] for r in rows[1:]]
        return getattr(self, '_rows', [])

    @rows.setter
    def rows(self, value):
        self._rows = value

    @property
    def merge_cells(self):
        if self.per_cell_data:
            return self.per_cell_data.get('merge_cells', [])
        return getattr(self, '_merge_cells', [])

    @merge_cells.setter
    def merge_cells(self, value):
        self._merge_cells = value

    @property
    def column_widths(self):
        if self.per_cell_data:
            return self.per_cell_data.get('gridCols', [])
        return getattr(self, '_column_widths', [])

    @column_widths.setter
    def column_widths(self, value):
        self._column_widths = value

    @property
    def grid_col_count(self):
        return len(self.actual_column_widths)

    def to_per_cell(self):
        if self.type != self.TYPE_TABLE or not self.per_cell_data:
            return None
        return self.per_cell_data

    def to_dict(self) -> dict:
        d = {"type": self.type}
        if self.type in (self.TYPE_PARAGRAPH, self.TYPE_HEADING, self.TYPE_LIST):
            d["text"] = self.text
            if self.level:
                d["level"] = self.level
            # 字体元数据
            if self.font_name:
                d["font_name"] = self.font_name
            if self.font_size is not None:
                d["font_size"] = self.font_size
            if self.bold:
                d["bold"] = True
            if self.alignment:
                d["alignment"] = self.alignment
        elif self.type == self.TYPE_TABLE:
            d["row_heights"] = self.row_heights
            d["headers"] = self.headers
            d["rows"] = self.rows
            d["merge_cells"] = self.merge_cells
            d["column_widths"] = self.column_widths
            if self.per_cell_data:
                d["per_cell_data"] = self.per_cell_data
        return d

    @classmethod
    def from_dict(cls, data: dict) -> "ContentBlock":
        cb = cls(data.get("type", "paragraph"), data.get("text", ""), data.get("level", 0))
        cb.font_name = data.get("font_name", "") or ""
        cb.font_size = data.get("font_size")
        cb.bold = data.get("bold", False) or False
        cb.alignment = data.get("alignment") or None
        cb.row_heights = data.get("row_heights", [])
        # 优先从 per_cell_data 还原，降级到旧字段
        pcd = data.get("per_cell_data")
        if pcd:
            cb.per_cell_data = pcd
        else:
            # 重建 per_cell_data（旧格式兼容）
            from app.infrastructure.table_codec import to_per_cell, to_dict
            hd = data.get("headers", [])
            rw = data.get("rows", [])
            mc = data.get("merge_cells", [])
            cw = data.get("column_widths", [])
            try:
                td = to_per_cell(hd, rw, mc, cw)
                cb.per_cell_data = to_dict(td)
            except Exception:
                cb.per_cell_data = None
        return cb


# ── 通用标题前置符剥离 ──
_HEADING_PREFIX_RE = __import__('re').compile(r'^[★◆●■▲➢※▪▶•·❤\s]+')


def strip_heading_prefix(text: str) -> str:
    """剥离标题前导装饰字符，保留标题实质内容。
    
    招标文件中标题经常带特殊符号标记，如 ★◆●■▲➢※ 等，
    这些符号会阻塞标题正则检测，导致整个章节被当作普通段落。
    
    剥离规则：去掉开头的连续非中文、非英文、非数字字符。
    
    示例:
      "★二、商务要求"  → "二、商务要求"
      "●1.技术要求"     → "1.技术要求"
      "【重要】三、须知"  → "三、须知"
      "◆ 四、资格要求"   -> " 四、资格要求"
      "比选邀请"        → "比选邀请"  (无前缀，保持不变)
    """
    if not text:
        return text
    return _HEADING_PREFIX_RE.sub('', text)


def _has_real_content(entry: dict) -> bool:
    """判断章节节点是否有真实内容。"""
    has_content = bool(entry.get("content"))
    has_children = bool(entry.get("children_ids"))
    page_range = entry.get("page_range", []) or []
    has_pages = len(page_range) >= 2 and (page_range[1] or 0) > (page_range[0] or 0)
    return has_content or has_children or has_pages


class DocumentParser:
    """版面感知的文档解析器。"""

    PARSE_VERSION = "3.0"

    def __init__(self, ocr_client=None):
        self.ocr_client = ocr_client
    # ========== WPS/Word 自动编号解析 ==========

    _CHINESE_NUMS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']

    @staticmethod
    def _int_to_chinese(n: int) -> str:
        """阿拉伯数字转中文数字（支持 1~999）。"""
        if n < 1:
            return str(n)
        if n <= 10:
            return DocumentParser._CHINESE_NUMS[n]
        if n < 20:
            return '十' + (DocumentParser._CHINESE_NUMS[n - 10] if n > 10 else '')
        if n < 100:
            tens = n // 10
            ones = n % 10
            return DocumentParser._CHINESE_NUMS[tens] + '十' + (DocumentParser._CHINESE_NUMS[ones] if ones else '')
        if n < 1000:
            hundreds = n // 100
            rest = n % 100
            result = DocumentParser._CHINESE_NUMS[hundreds] + '百'
            if rest:
                if rest < 10:
                    result += '零' + DocumentParser._CHINESE_NUMS[rest]
                else:
                    result += DocumentParser._int_to_chinese(rest)
            return result
        return str(n)

    _NUMFMT_HANDLERS = {
        'chineseCounting': _int_to_chinese,
        'decimal': str,
        'ordinal': str,
        'cardinalText': str,
    }

    def _load_numbering_defs(self, payload: bytes) -> dict:
        """从 DOCX payload 中加载自动编号定义（numbering.xml）。

        Returns:
            dict: {num_id: {ilvl: {numFmt, lvlText, start}}}
        """
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        numbering_defs = {}
        try:
            with ZipFile(BytesIO(payload)) as zf:
                if 'word/numbering.xml' not in zf.namelist():
                    return numbering_defs
                tree = ET.parse(zf.open('word/numbering.xml'))
                root = tree.getroot()

                # 1. Build abstractNum lookup
                abstract_nums = {}
                for anum in root.findall(f'.//{ns}abstractNum'):
                    anum_id = anum.get(f'{ns}abstractNumId')
                    levels = {}
                    for lvl in anum.findall(f'{ns}lvl'):
                        ilvl = lvl.get(f'{ns}ilvl')
                        numFmt_elem = lvl.find(f'{ns}numFmt')
                        lvlText_elem = lvl.find(f'{ns}lvlText')
                        start_elem = lvl.find(f'{ns}start')
                        numFmt = numFmt_elem.get(f'{ns}val') if numFmt_elem is not None else 'decimal'
                        lvlText = lvlText_elem.get(f'{ns}val') if lvlText_elem is not None else '%1.'
                        start = int(start_elem.get(f'{ns}val')) if start_elem is not None else 1
                        levels[int(ilvl)] = {
                            'numFmt': numFmt,
                            'lvlText': lvlText,
                            'start': start,
                        }
                    abstract_nums[anum_id] = levels

                # 2. Build num lookup (numId -> abstractNumId)
                for num_elem in root.findall(f'.//{ns}num'):
                    num_id = num_elem.get(f'{ns}numId')
                    anum_ref = num_elem.find(f'{ns}abstractNumId')
                    if anum_ref is not None:
                        anum_id = anum_ref.get(f'{ns}val')
                        numbering_defs[num_id] = abstract_nums.get(anum_id, {})
        except Exception as e:
            logger.warning("[parser] 加载编号定义失败: %s", e)
        return numbering_defs

    def _resolve_numbering_prefix(self, para_element, numbering_defs: dict,
                                   num_counters: dict) -> str:
        """从段落 XML 中解析自动编号前缀。

        Args:
            para_element: python-docx 段落的 _element (lxml Element)
            numbering_defs: _load_numbering_defs 返回的编号定义
            num_counters: 跟踪每个 numId+ilvl 已出现的次数（跨段落维护）

        Returns:
            str: 解析出的编号前缀（如 "三、"），无编号时返回空字符串
        """
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        try:
            pPr = para_element.find(f'{ns}pPr')
            if pPr is None:
                return ''

            numPr = pPr.find(f'{ns}numPr')
            if numPr is None:
                return ''

            numId_elem = numPr.find(f'{ns}numId')
            ilvl_elem = numPr.find(f'{ns}ilvl')

            if numId_elem is None:
                return ''

            num_id = numId_elem.get(f'{ns}val')
            ilvl = int(ilvl_elem.get(f'{ns}val', '0')) if ilvl_elem is not None else 0

            levels = numbering_defs.get(num_id, {})
            level = levels.get(ilvl, {})

            if not level or level.get('numFmt') in ('none', 'bullet'):
                return ''

            numFmt = level['numFmt']
            lvlText = level['lvlText']
            start = level['start']

            # Track sequential count for this numId+ilvl
            key = (num_id, ilvl)
            count = num_counters.get(key, 0)
            num_counters[key] = count + 1

            current_num = start + count

            # Format the number according to numFmt
            handler = self._NUMFMT_HANDLERS.get(numFmt, str)
            number_text = handler(current_num)

            prefix = lvlText.replace('%1', number_text)
            # 清除多级编号的残留引用 %2, %3...
            prefix = re.sub(r'%\d+', '', prefix).strip()

            return prefix
        except Exception as e:
            logger.debug("[parser] 解析编号前缀失败: %s", e)
            return ''



    # ========== 统一入口 ==========

    def parse_structured(self, filename: str, payload: bytes, file_sha256: str = "") -> StructuredDocument:
        """统一入口：按扩展名选择解析器，返回结构化文档。

        Args:
            filename: 文件名（用于判断扩展名）
            payload: 文件二进制内容
            file_sha256: 文件 SHA256（可选）

        Returns:
            StructuredDocument 结构化文档
        """
        doc = StructuredDocument(file_name=filename, file_sha256=file_sha256, parse_version=self.PARSE_VERSION)
        ext = Path(filename).suffix.lower().lstrip(".") if filename else ""

        if ext == "docx":
            self._parse_docx_structured(payload, doc)
        elif ext in ("pdf",):
            self._parse_pdf_structured(payload, doc)
        elif ext in ("doc",):
            self._parse_doc_structured(payload, doc)
        elif ext in ("xlsx", "xls"):
            self._parse_spreadsheet_structured(payload, doc, ext)
        else:
            # 纯文本兜底
            text = payload.decode("utf-8", errors="replace")
            section = Section(title="全文", level=1)
            section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
            doc.sections.append(section)

        return doc

    # ========== DOCX 结构化解析 ==========

    def _parse_docx_structured(self, payload: bytes, doc: StructuredDocument):
        """解析 DOCX，保留标题层级和表格结构。"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("[parser] python-docx 未安装，使用降级解析")
            self._parse_docx_fallback(payload, doc)
            return

        try:
            document = DocxDocument(BytesIO(payload))
        except Exception as exc:
            logger.warning("[parser] python-docx 解析失败，使用降级解析: %s", exc)
            self._parse_docx_fallback(payload, doc)
            return

        # Heading 样式映射（含 toc 样式）
        heading_map = {}
        for i in range(1, 10):
            style_name = f"Heading {i}"
            try:
                heading_map[style_name] = i
            except Exception:
                pass


        # 文本内容级别的标题检测模式
        text_heading_patterns = [
            (1, r'^第[一二三四五六七八九十零〇百千万亿]+[章节篇部]'),      # 第一章
            (2, r'^[一二三四五六七八九十零〇]+[、，,．.]'),              # 一、
            (2, r'^\d+[、，,．.]'),                                   # 1.
            (2, r'^\d+\.\d+\s'),                                   # 1.1
            (3, r'^（[一二三四五六七八九十零〇]+）'),                    # （一）
            (3, r'^\d+\.\d+\.\d+\s'),                            # 1.1.1
            # 中文括号封面/封皮标记，如（资格性响应文件封面、封皮）
            (2, r'^（[^）]*封面[^）]*）'),
            (2, r'^（[^）]*封皮[^）]*）'),

        ]

        # 加载 WPS/Word 自动编号定义
        numbering_defs = self._load_numbering_defs(payload)
        num_counters = {}

        stack = [Section(title="__root__", level=0)]
        current_section = stack[-1]
        # ── 跟踪 body 子元素到章节的映射 ──
        body_child_section_map = []  # list of (body_child_index, section)
        _body_child_counter = 0

        for para in document.paragraphs:
            _body_child_counter += 1
            text = para.text.strip()
            # 解析 WPS/Word 自动编号前缀
            if text and numbering_defs and hasattr(para, '_element') and para._element is not None:
                _num_prefix = self._resolve_numbering_prefix(para._element, numbering_defs, num_counters)
                if _num_prefix:
                    text = _num_prefix + text
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            heading_level = heading_map.get(style_name, 0)

            if heading_level > 0:
                # 覆盖检查：如果样式层级与文本内容不匹配（如 Heading2 的"第七章"），
                # 使用文本检测的层级（更准确）
                _text_level_override = 0
                for _tl, _tp in text_heading_patterns:
                    if re.match(_tp, text):
                        _text_level_override = _tl
                        break
                if _text_level_override > 0 and _text_level_override < heading_level:
                    heading_level = _text_level_override
                # 创建新章节
                new_section = Section(title=text, level=heading_level)
                # 弹出比当前层级深或相等的章节
                while stack and stack[-1].level >= heading_level:
                    stack.pop()
                # 仅剩 __root__ 时直接加到 doc.sections（修复章节消失 bug）
                if len(stack) == 1 and stack[0].level == 0:
                    doc.sections.append(new_section)
                elif stack:
                    stack[-1].children.append(new_section)
                else:
                    doc.sections.append(new_section)
                stack.append(new_section)
                current_section = new_section
                body_child_section_map.append((_body_child_counter, current_section))
            else:
                # 文本内容级标题检测（当样式为 Normal 但内容像标题时）
                text_heading = 0
                # 剥离前导装饰字符后再匹配（如 ★二、商务要求 → 二、商务要求）
                text_for_heading = strip_heading_prefix(text) if text else text
                for level, pattern in text_heading_patterns:
                    if re.match(pattern, text_for_heading):
                        text_heading = level
                        break
                
                if text_heading > 0 and heading_level == 0:
                    # ── 安检门1: 结尾标点拦截 ──
                    # 以句号/分号/冒号结尾 → 判定为内容段落
                    # 例外：纯中文编号标题（"一、项目概况。"）保留标题资格
                    _is_pure_chinese_num = bool(
                        re.match(r'^[一二三四五六七八九十]+[、，]', text)
                        and len(text) <= 15
                    )
                    if text.endswith(('。', '；', '：', ';', ':')) and not _is_pure_chinese_num:
                        block = ContentBlock(ContentBlock.TYPE_PARAGRAPH, text)
                        current_section.content.append(block)
                        continue

                    # ── 安检门3: 标题过长拦截（≥100字 → 内容段落）──
                    # 有标题关键词（承诺、声明、保证、格式、封面、封皮、应答、一览表、情况表）时仍保留标题资格
                    _heading_kw = re.search(r'承诺|声明|保证|格式|封面|封皮|应答|一览表|情况表|响应函|投标函|报价函|授权|证明|合同|协议|方案|说明|须知', text)
                    if not _heading_kw and len(text) >= 100:
                        block = ContentBlock(ContentBlock.TYPE_PARAGRAPH, text)
                        current_section.content.append(block)
                        continue

                    # 跳过目录项（含 tab 或纯数字页码的短标题）
                    if "\t" in text:
                        # 来自 TOC 目录的条目，不作为章节
                        block = ContentBlock(ContentBlock.TYPE_PARAGRAPH, text)
                        current_section.content.append(block)
                        continue
                    new_section = Section(title=text, level=text_heading)
                    while stack and stack[-1].level >= text_heading:
                        stack.pop()
                    # 仅剩 __root__ 时直接加到 doc.sections（修复章节消失 bug）
                    if len(stack) == 1 and stack[0].level == 0:
                        doc.sections.append(new_section)
                    elif stack:
                        stack[-1].children.append(new_section)
                    else:
                        doc.sections.append(new_section)
                    stack.append(new_section)
                    current_section = new_section
                    body_child_section_map.append((_body_child_counter, current_section))
                else:
                    # ── 安检门4: 格式检测 — 加粗+大号字体的潜在标题（无编号前缀）──
                    try:
                        _has_bold = any(getattr(r, "bold", None) for r in para.runs)
                        _large_font = any(
                            getattr(getattr(r, "font", None), "size", None) and
                            getattr(r.font, "size", None) >= 152400
                            for r in para.runs
                        )
                    except Exception:
                        _has_bold = False
                        _large_font = False
                    if _has_bold and _large_font and 3 <= len(text) < 80 and not text.endswith(('。', '；', '：', ';', ':', '、', '）')) and '：' not in text:
                        _new_sec = Section(title=text, level=2)
                        while stack and stack[-1].level >= 2:
                            stack.pop()
                        if len(stack) == 1 and stack[0].level == 0:
                            doc.sections.append(_new_sec)
                        elif stack:
                            stack[-1].children.append(_new_sec)
                        else:
                            doc.sections.append(_new_sec)
                        stack.append(_new_sec)
                        current_section = _new_sec
                    else:
                        # 提取字体元数据（封面渲染用）
                        _font_name = ""
                        _font_size = None
                        _bold = False
                        _alignment = None
                        try:
                            for _r in para.runs:
                                if not _font_name:
                                    try:
                                        _fn = getattr(_r.font, "name", None) or ""
                                        if _fn:
                                            _font_name = _fn
                                    except Exception:
                                        pass
                                if _font_size is None:
                                    try:
                                        _fs = getattr(_r.font, "size", None)
                                        if _fs:
                                            _font_size = _fs / 12700.0  # EMU → Pt
                                    except Exception:
                                        pass
                                if not _bold:
                                    try:
                                        if getattr(_r, "bold", None):
                                            _bold = True
                                    except Exception:
                                        pass
                            # 段落对齐方式
                            _pPr = para._p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr') if hasattr(para, '_p') else None
                            if _pPr is not None:
                                _jc = _pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
                                if _jc is not None:
                                    _jc_val = _jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                                    if _jc_val:
                                        _alignment = _jc_val
                        except Exception:
                            pass
                        block = ContentBlock(ContentBlock.TYPE_PARAGRAPH, text)
                        block.font_name = _font_name
                        block.font_size = _font_size
                        block.bold = _bold
                        block.alignment = _alignment
                        # 尝试判断列表
                        num_prefix = re.match(r'^[\d一二三四五六七八九十]+[、.．\s]', text)
                        bullet_prefix = re.match(r'^[-•●○■\s]', text)
                        if num_prefix or bullet_prefix:
                            block.type = ContentBlock.TYPE_LIST
                            block.level = 0
                        current_section.content.append(block)

        # 在解析表格前，先保存栈中的段落内容
        if stack:
            root_section = stack[0]
            # 将 root section 中的非标题内容保存为一个前言章节
            if root_section.content:
                from_title = root_section.content[0].text[:30] if root_section.content[0].text else ""
                preamble = Section(title=from_title or "前言", level=1)
                preamble.content = list(root_section.content)
                root_section.content = []
                doc.sections.insert(0, preamble)
            # 将 root 下的子章节（文本检测到的标题）转移到 doc.sections
            if root_section.children:
                for child in root_section.children:
                    child.level = 1  # 提升到顶级
                    doc.sections.append(child)
                root_section.children = []

        # 直接遍历 body 子元素，按出现顺序将每个 table 分配给对应的活跃章节
        _para_counter = 0  # 对应 document.paragraphs 中的索引（0-based）
        _table_counter = 0
        _current_section_for_table = doc.sections[-1] if doc.sections else None
        # 追踪每个章节已分配元素的最后内容索引，用于表格的正确插入位置
        _element_last_pos = {}  # id(deep_section) -> last content index

        try:
            body = document.element.body
            for child in body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("p", "pPr"):
                    if _para_counter < len(document.paragraphs):
                        # 从 body_child_section_map 中查找该段落的章节
                        for _bidx, _sec in reversed(body_child_section_map):
                            if _bidx <= _para_counter + 1:
                                _current_section_for_table = _sec
                                break
                    # 更新章节的最后元素索引（段落已在第1遍加入 content，需追踪其在 content 中的位置）
                    # 注意：仅计数实际产生 ContentBlock 的段落（非 heading、非空段）
                    # heading 段落和空段在 Pass 1 中不会添加到 section.content，因此不应影响 _insert_pos
                    if _current_section_for_table:
                        # 判断此段落是否会产生 ContentBlock
                        _is_content_paragraph = True
                        _para_texts = [t.text for t in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if t.text]
                        _para_text = ''.join(_para_texts).strip()
                        # 空段不产生 ContentBlock
                        if not _para_text:
                            _is_content_paragraph = False
                        else:
                            # style 标题不产生 ContentBlock
                            _pPr = child.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                            if _pPr is not None:
                                _pStyle = _pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                                if _pStyle is not None:
                                    _pStyle_val = _pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                                    if _pStyle_val in heading_map:
                                        _is_content_paragraph = False
                            # 文本模式标题不产生 ContentBlock
                            if _is_content_paragraph:
                                _stripped = strip_heading_prefix(_para_text) if _para_text else _para_text
                                for _tl, _tp in text_heading_patterns:
                                    if re.match(_tp, _stripped):
                                        _is_content_paragraph = False
                                        break
                        if _is_content_paragraph:
                            _deep = _current_section_for_table
                            while _deep and _deep.children:
                                _deep = _deep.children[-1]
                            if _deep:
                                _sec_key = id(_deep)
                                _cur = _element_last_pos.get(_sec_key, -1)
                                _element_last_pos[_sec_key] = _cur + 1
                    _para_counter += 1
                elif tag == "tbl":
                    if _table_counter < len(document.tables):
                        # 计算此表格的插入位置：在此章节上一个已分配元素的后面
                        _insert_pos = -1
                        if _current_section_for_table:
                            _deep = _current_section_for_table
                            while _deep and _deep.children:
                                _deep = _deep.children[-1]
                            if _deep:
                                _sec_key = id(_deep)
                                _insert_pos = _element_last_pos.get(_sec_key, -1) + 1
                        self._parse_table(document.tables[_table_counter], doc,
                                          table_index=_table_counter, docx_document=document,
                                          _position_hint=_current_section_for_table,
                                          _insert_position=_insert_pos)
                        # 更新章节的最后元素索引
                        if _insert_pos >= 0 and _current_section_for_table:
                            _deep = _current_section_for_table
                            while _deep and _deep.children:
                                _deep = _deep.children[-1]
                            if _deep:
                                _sec_key = id(_deep)
                                _element_last_pos[_sec_key] = _insert_pos
                        _table_counter += 1
        except Exception:
            logger.warning("[parser] 直接 body 遍历失败，使用降级定位: %s", exc_info=True)
            # 降级：使用 body_child_counter 近似定位
            for table_idx, table in enumerate(document.tables):
                _table_body_idx = _body_child_counter + table_idx + 1
                _nearest_section = None
                for _bidx, _sec in reversed(body_child_section_map):
                    if _bidx < _table_body_idx:
                        _nearest_section = _sec
                        break
                self._parse_table(table, doc, table_index=table_idx, docx_document=document, _position_hint=_nearest_section)

        # 保存原始表格对象供 table_parser 使用
        doc.tables = list(document.tables)

        # 清理空的根章节
        doc.sections = [s for s in doc.sections if s.title != "__root__"]
        
        # 安检门2: 事后扫描 — 移除连续 3+ 同级别空节点（列举列表误判）
        self._cleanup_fake_headings(doc.sections)

        # 如果没有检测到任何标题层级（纯文本），把内容放到一个根章节下
        if not doc.sections:
            root_content = []
            if stack and stack[0].children:
                doc.sections = stack[0].children
            elif stack:
                root = Section(title="全文", level=1)
                for s in stack:
                    root.content.extend(s.content)
                if root.content or root.children:
                    doc.sections = [root]

    def _parse_table(self, table, doc: StructuredDocument, table_index: int = 0, docx_document=None, _position_hint=None, _insert_position=-1):
        """从 python-docx Table 对象提取结构化表格，并尝试分配到正确的章节。

        Args:
            table: python-docx Table 对象
            doc: StructuredDocument 目标文档
            table_index: 表格在 document.tables 中的索引（用于定位章节）
            docx_document: 可选的 python-docx Document，用于定位表格在正文中的位置
        """
        block = ContentBlock(ContentBlock.TYPE_TABLE)
        rows_data = []
        merge_cells = []
        column_widths = []
        # 提取列宽
        try:
            for col in table.columns:
                column_widths.append(col.width)
        except Exception:
            column_widths = []
        block.column_widths = column_widths

        # 提取行高
        block.row_heights = _extract_table_row_heights(table)

        # 提取合并单元格（gridSpan / vMerge）— 基于 XML <tc> 直接遍历，无重复
        _ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        _covered_vmerge_ranges = {}
        for row_idx, row in enumerate(table.rows):
            cells = []
            # 第1遍：通过 python-docx 虚拟单元格提取文本（正确展开合并格的文本）
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                cells.append(text)

            # 第2遍：通过 XML <tc> 元素遍历检测合并（避免虚拟单元格导致重复）
            tr = row._tr
            tc_elements = tr.findall(f'{_ns}tc')
            col_idx = 0
            for tc in tc_elements:
                try:
                    tcPr = tc.find(f'{_ns}tcPr')
                    if tcPr is None:
                        col_idx += 1
                        continue

                    # 水平合并
                    gs = tcPr.find(f'{_ns}gridSpan')
                    span_val = int(gs.get(f'{_ns}val', '1')) if gs is not None else 1
                    if gs is not None and span_val > 1:
                        merge_cells.append({"type": "horizontal", "row": row_idx, "col": col_idx, "span": span_val})

                    # 垂直合并
                    vm = tcPr.find(f'{_ns}vMerge')
                    if vm is not None:
                        vm_val = vm.get(f'{_ns}val')
                        if vm_val is None or vm_val == 'restart':
                            merge_span = 1
                            cell_text = table.rows[row_idx].cells[col_idx].text.strip()
                            for nr in range(row_idx + 1, len(table.rows)):
                                _vstop = False
                                try:
                                    next_cell = table.rows[nr].cells[col_idx]
                                    next_tc = next_cell._tc
                                    next_vm = next_tc.find(f'.//{_ns}vMerge')
                                    if next_vm is not None:
                                        next_val = next_vm.get(f'{_ns}val')
                                        if next_val is None or next_val == 'continue':
                                            merge_span += 1
                                        elif next_val == 'restart':
                                            if next_cell.text.strip() == cell_text:
                                                merge_span += 1
                                            else:
                                                _vstop = True
                                        else:
                                            _vstop = True
                                    else:
                                        _vstop = True
                                except Exception:
                                    _vstop = True
                                if _vstop:
                                    break
                            if merge_span > 1:
                                _already_covered = False
                                for (_vr, _vspan) in _covered_vmerge_ranges.get(col_idx, []):
                                    if _vr <= row_idx < _vr + _vspan:
                                        _already_covered = True
                                        break
                                if not _already_covered:
                                    merge_cells.append({"type": "vertical", "row": row_idx, "col": col_idx, "span": merge_span})
                                    _covered_vmerge_ranges.setdefault(col_idx, []).append((row_idx, merge_span))

                    col_idx += span_val
                except Exception:
                    col_idx += 1

            if row_idx == 0:
                _header_cells = cells
            else:
                rows_data.append(cells)

        # Stage 5: 构建 per_cell_data 作为唯一存储
        from app.infrastructure.table_codec import to_per_cell, to_dict
        try:
            all_rows = [_header_cells] + rows_data
            td = to_per_cell(_header_cells, all_rows, merge_cells, column_widths, block.row_heights)
            pcd = to_dict(td)
            pcd["merge_cells"] = merge_cells
            block.per_cell_data = pcd
        except Exception:
            # 降级：直接存储旧格式（per_cell 未构建成功）
            if not block.per_cell_data:
                block.per_cell_data = None
            logger.warning("[parser] to_per_cell 构建失败，降级到旧格式")
        # ===== WPS 伪合并检测（无 OOXML gridSpan/vMerge 的合并） =====
        if not merge_cells and rows_data:
            # 水平伪合并：逐行检测连续相同非空文本
            all_rows = [_header_cells] + rows_data
            for r_idx in range(len(all_rows)):
                c = 0
                while c < len(all_rows[r_idx]):
                    cell_text = all_rows[r_idx][c].strip()
                    if not cell_text:
                        c += 1
                        continue
                    span = 1
                    while c + span < len(all_rows[r_idx]) and all_rows[r_idx][c + span].strip() == cell_text:
                        span += 1
                    if span > 1:
                        merge_cells.append({"type": "horizontal", "row": r_idx, "col": c, "span": span})
                        c += span
                    else:
                        c += 1
        if not doc.sections:
            s = Section(title="表格", level=1)
            s.content.append(block)
            doc.sections.append(s)
            return

        # 尝试定位表格在文档体中的位置
        target_section = _position_hint or doc.sections[-1]  # 默认：最后章节或位置提示
        if docx_document and hasattr(docx_document, "element") and hasattr(docx_document.element, "body"):
            try:
                body = docx_document.element.body
                # 遍历 body 子元素，找到所有表格的索引
                table_elements = []
                for child in body:
                    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                    if tag == "tbl":
                        table_elements.append(child)
                
                # 找到当前表格前面的段落文本
                if table_index < len(table_elements):
                    tbl_elem = table_elements[table_index]
                    prev_text = ""
                    # 找表格前最近的段落文本
                    for child in body:
                        if child is tbl_elem:
                            break
                        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                        if tag in ("p", "pPr"):
                            # 提取段落文本
                            texts = child.itertext() if hasattr(child, "itertext") else []
                            for t in texts:
                                if t.strip():
                                    prev_text = t.strip()
                        elif tag == "tbl":
                            prev_text = ""  # 表格后的内容
                    
                    if prev_text and not _position_hint:
                        # 按标题定位（仅当没有位置提示时使用文本查找作为后备）
                        target_section = self._find_section_by_text(doc, prev_text)
            except Exception as exc:
                pass

        # 添加到目标章节
        target = target_section
        while target.children:
            target = target.children[-1]
        if _insert_position >= 0 and _insert_position < len(target.content):
            target.content.insert(_insert_position, block)
        else:
            target.content.append(block)

    def _find_section_by_text(self, doc, text: str):
        """根据文本片段找到包含它的章节。"""
        if not text:
            return doc.sections[-1] if doc.sections else None

        best_section = None
        best_match_len = 0

        def _search(node, depth=0):
            nonlocal best_section, best_match_len
            node_text = getattr(node, "title", "") or ""
            if node_text and text in node_text and len(text) > best_match_len:
                best_section = node
                best_match_len = len(text)
            for block in getattr(node, "content", []):
                block_text = getattr(block, "text", "") or ""
                if block_text and text in block_text and len(text) > best_match_len:
                    best_section = node
                    best_match_len = len(text)
            for child in getattr(node, "children", []):
                _search(child, depth + 1)

        for section in doc.sections:
            _search(section)

        return best_section or (doc.sections[-1] if doc.sections else None)

    def _parse_docx_fallback(self, payload: bytes, doc: StructuredDocument):
        """DOCX 降级解析：使用 docx2python 提取纯文本。"""
        try:
            from docx2python import docx2python
            with docx2python(BytesIO(payload)) as result:
                text = (result.text or "").strip()
                if text:
                    section = Section(title="全文", level=1)
                    section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
                    doc.sections.append(section)
        except Exception as exc:
            logger.warning("[parser] docx2python 也失败: %s", exc)
            text = payload.decode("utf-8", errors="replace")
            section = Section(title="全文", level=1)
            section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
            doc.sections.append(section)

    # ══════════════════════════════════════════════════════════════
    #  安检门2: 事后扫描 — 移除连续 3+ 同级别空节点（列举列列表误判）
    # ══════════════════════════════════════════════════════════════

    def _cleanup_fake_headings(self, sections_list):
        """事后扫描：移除连续 3+ 同级别的空节点（列举列表误判）
    
        同时处理顶层 sections 列表和子章节。
        """
        def _is_legitimate_empty(section):
            """判断空节点是否为合法结构节点（封面、编号标题等，不应被移除）。"""
            title = getattr(section, 'title', '') or ''
            if '封面' in title or '封皮' in title:
                return True
            if re.match(r'^[一二三四五六七八九十零〇]+[、，,．.]', title):
                return True
            if re.match(r'^（[^）]*）', title):
                return True
            return False

        def _scan_siblings(children, parent=None):
            if not children:
                return
            to_delete = []
            i = 0
            while i < len(children):
                j = i
                level = children[i].level if hasattr(children[i], 'level') else 0
                while j < len(children):
                    c = children[j]
                    same_level = c.level == level if hasattr(c, 'level') else True
                    no_content = (len(getattr(c, 'content', []) or []) == 0
                                  and len(getattr(c, 'children', []) or []) == 0)
                    if same_level and no_content:
                        j += 1
                    else:
                        break
                count = j - i
                # 检查这批空节点中是否有合法结构节点（封面、编号等）
                _has_legit = any(_is_legitimate_empty(children[k]) for k in range(i, j))
                if count >= 3 and not _has_legit:
                    logger.debug("[parser] 安检门2: 将 %d 个连续空章节降级为内容", count)
                    for k in range(i, j):
                        block = ContentBlock(ContentBlock.TYPE_PARAGRAPH,
                                             getattr(children[k], 'title', ''))
                        if parent is not None:
                            parent.content.append(block)
                    to_delete.extend(range(i, j))
                    i = j
                elif count == 0:
                    _scan_siblings(getattr(children[i], 'children', []) or [], children[i])
                    i += 1
                else:
                    for k in range(i, j):
                        _scan_siblings(getattr(children[k], 'children', []) or [], children[k])
                    i = j
            for idx in reversed(to_delete):
                del children[idx]
    
        if sections_list:
            _scan_siblings(sections_list)
            for section in list(sections_list):
                _scan_siblings(getattr(section, 'children', []) or [], section)
    
    # ========== PDF 结构化解析 ==========

    def _parse_pdf_structured(self, payload: bytes, doc: StructuredDocument):
        """解析 PDF：fitz 逐页判断类型，混合策略提取。"""
        try:
            pdf_doc = fitz.open(stream=payload, filetype="pdf")
        except Exception as exc:
            logger.error("[parser] fitz 打开 PDF 失败: %s", exc)
            text = payload.decode("utf-8", errors="replace")
            section = Section(title="全文", level=1)
            section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
            doc.sections.append(section)
            return

        total_pages = len(pdf_doc)
        text_pages = []       # list of (page_text, page_no)
        ocr_pages = []        # list of (image_bytes, page_no)

        # 第一步：逐页判断类型
        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            page_text = page.get_text().strip()
            images = page.get_images()

            is_scan = (len(page_text) < 50 and len(images) > 0) or (len(page_text) < 20)
            is_mixed = len(page_text) < 200 and len(images) > 0

            # 表格检测（对所有类型的页面都尝试）
            page_tables = self._detect_tables_in_pdf_page(page)
            
            if is_scan and self.ocr_client:
                # 有OCR：渲染为图片后 OCR
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                ocr_pages.append((img_bytes, page_num + 1))
                if page_text:
                    text_pages.append((page_text, page_num + 1))
                if page_tables:
                    text_pages.append(("", page_num + 1))
            elif is_scan:
                # 无OCR：使用已有文本，不足则标记
                text = page_text or f"【第{page_num + 1}页为扫描页，无可用文本】"
                text_pages.append((text, page_num + 1))
                if page_tables:
                    for t in page_tables:
                        text_pages.append((t, page_num + 1))
            elif is_mixed and self.ocr_client:
                # 混合页：文本 + OCR 补充
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                ocr_pages.append((img_bytes, page_num + 1))
                if page_text:
                    text_pages.append((page_text, page_num + 1))
            else:
                # 纯文本页
                if page_tables:
                    # 有表格时：用表格文本替换纯文本
                    table_text_parts = []
                    for t in page_tables:
                        h = " | ".join(t.headers) if t.headers else ""
                        rows = [" | ".join(r) for r in t.rows]
                        table_text_parts.append(h + "\n" + "\n".join(rows))
                    combined_table_text = "\n".join(table_text_parts)
                    if page_text:
                        combined_table_text = page_text + "\n" + combined_table_text
                    text_pages.append((combined_table_text, page_num + 1))
                else:
                    text_pages.append((page_text, page_num + 1))

        pdf_doc.close()

        # 第二步：OCR 识别
        ocr_results = {}  # page_no -> list of (text, box)
        if ocr_pages and self.ocr_client:
            import asyncio
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

            images = [img for img, _ in ocr_pages]
            results = loop.run_until_complete(self.ocr_client.recognize_images_batch(images))
            for (_, page_no), page_items in zip(ocr_pages, results):
                ocr_results[page_no] = [(item["text"], item.get("box")) for item in page_items]

        # 第三步：按页码合并，重建版面
        all_pages_content = {}  # page_no -> list[ContentBlock]
        for text, page_no in text_pages:
            blocks = self._parse_text_page(text)
            all_pages_content[page_no] = blocks
        for page_no, items in ocr_results.items():
            texts = [t for t, _ in items]
            combined = "\n".join(texts)
            if page_no in all_pages_content:
                existing = all_pages_content[page_no]
                existing.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, combined))
            else:
                all_pages_content[page_no] = [ContentBlock(ContentBlock.TYPE_PARAGRAPH, combined)]

        # 第四步：按页码排序，合并为章节
        sorted_pages = sorted(all_pages_content.items())
        full_text = []
        for page_no, blocks in sorted_pages:
            for b in blocks:
                full_text.append(b.text)

        combined = "\n".join(full_text)
        self._build_sections_from_text(combined, doc)


    def _detect_table_in_text(self, text_block: str) -> "Optional[ContentBlock]":
        """启发式检测文本中是否包含表格结构，尝试重建为 ContentBlock(type=table)。
        
        检测条件：
        1. 连续 >=3 行，每行有相同的列数（按 3+空格 / | / \t 分割）
        2. 首行可能为表头
        
        返回 ContentBlock(type=table) 或 None
        """
        if not text_block or not isinstance(text_block, str):
            return None
        
        lines = [l.strip() for l in text_block.split("\n") if l.strip()]
        if len(lines) < 3:
            return None
        
        # 尝试多种分隔符
        separators = [
            lambda x: [c.strip() for c in re.split(r"\s{3,}", x) if c.strip()],  # 3+空格
            lambda x: [c.strip() for c in x.split("\t") if c.strip()],           # tab
            lambda x: [c.strip() for c in x.split("|") if c.strip()],             # pipe
        ]
        
        for sep in separators:
            split_lines = [sep(line) for line in lines]
            col_counts = [len(sl) for sl in split_lines]
            
            # 检查是否有 >=3 行有相同的列数 >=2
            from collections import Counter
            count_counter = Counter(col_counts)
            most_common_count, occurrences = count_counter.most_common(1)[0]
            
            if occurrences >= 3 and most_common_count >= 2:
                # 判定为表格
                table_lines = [sl for sl in split_lines if len(sl) == most_common_count]
                if len(table_lines) < 3:
                    continue
                
                block = ContentBlock(ContentBlock.TYPE_TABLE)
                block.headers = table_lines[0]
                block.rows = table_lines[1:]
                return block
        
        return None

    def _detect_tables_in_pdf_page(self, page) -> "list[ContentBlock]":
        """用 fitz 内置表格检测提取 PDF 页面的表格。
        
        优先用 find_tables()（检测网格线），
        失败则对页面文本用启发式检测。
        """
        tables = []
        
        # 方法1：fitz find_tables() - 检测有网格线的表格
        try:
            found = page.find_tables()
            if found and found.tables:
                for ft in found.tables:
                    data = ft.extract()
                    if not data or len(data) < 2:
                        continue
                    block = ContentBlock(ContentBlock.TYPE_TABLE)
                    block.headers = [str(c).strip() for c in data[0]]
                    block.rows = [[str(c).strip() for c in row] for row in data[1:]]
                    tables.append(block)
                if tables:
                    return tables
        except Exception as exc:
            logger.debug("[parser] fitz 表格检测异常: %s", exc)
        
        # 方法2：启发式 - 从页面文本中检测
        try:
            page_text = page.get_text().strip()
            if page_text:
                block = self._detect_table_in_text(page_text)
                if block:
                    tables.append(block)
        except Exception as exc:
            logger.debug("[parser] 启发式表格检测异常: %s", exc)
        
        return tables


    def _parse_text_page(self, text: str) -> list:
        """从纯文本中提取内容块。"""
        blocks = []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            block_type = ContentBlock.TYPE_PARAGRAPH
            # 尝试识别标题
            heading_match = re.match(r'^(第[一二三四五六七八九十]+[章节篇]|[\d]+[.、]\s*|（[\d一二三四五六七八九十]+）)', line)
            if heading_match:
                block_type = ContentBlock.TYPE_HEADING
            blocks.append(ContentBlock(block_type, line))
        return blocks

    def _build_sections_from_text(self, text: str, doc: StructuredDocument):
        """从合并文本中重建章节结构。"""
        lines = text.split("\n")
        # 加载 WPS/Word 自动编号定义
        numbering_defs = self._load_numbering_defs(payload)
        num_counters = {}

        stack = [Section(title="__root__", level=0)]

        for line in lines:
            line = line.strip()
            if not line:
                continue

            heading_level = self._detect_heading_level(line)
            if heading_level > 0:
                section = Section(title=line, level=heading_level)
                while stack and stack[-1].level >= heading_level:
                    stack.pop()
                if stack:
                    stack[-1].children.append(section)
                else:
                    doc.sections.append(section)
                stack.append(section)
            else:
                if stack:
                    stack[-1].content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, line))

        # 清理空根节点
        doc.sections = [s for s in doc.sections if s.title != "__root__"]
        # 如果没有章节结构，创建一个
        if not doc.sections and stack and stack[0].children:
            doc.sections = stack[0].children

    def _detect_heading_level(self, text: str) -> int:
        """检测文本是否是标题，返回标题层级（0=不是标题）。"""
        # 一级标题：第X章/X篇
        if re.match(r'^第[一二三四五六七八九十零〇百千万亿]+[章节篇部]', text):
            return 1
        # 二级标题：一、 二、 或 1. 2. 或 1.1
        if re.match(r'^[一二三四五六七八九十零〇]+[、，,．.]', text):
            return 2
        if re.match(r'^\d+[、，,．.]', text):
            return 2
        if re.match(r'^\d+\.\d+\s', text):
            return 2
        # 三级标题：（一）（二）或 1.1.1
        if re.match(r'^（[一二三四五六七八九十零〇]+）', text):
            return 3
        if re.match(r'^\d+\.\d+\.\d+\s', text):
            return 3
        return 0


    def _detect_text_tables_in_text(self, text):
        """从纯文本中检测表格，返回 ContentBlock(type=table) 列表。
        
        用于 PDF 文本页和扫描页 OCR 结果中的非原生表格文本。
        """
        tables = []
        lines = text.split("\n")
        current_table = []
        in_table = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if in_table and len(current_table) >= 3:
                    table = self._parse_text_table(current_table)
                    if table:
                        tables.append(table)
                current_table = []
                in_table = False
                continue

            is_table_line = False
            if "|" in stripped or "\u2502" in stripped:
                is_table_line = True
            elif "\t" in stripped:
                is_table_line = True
            elif stripped and stripped[0] in "\u250c\u2510\u2514\u2518\u251c\u2524\u252c\u2534\u253c\u2550\u2500\u2502\u2503\u2554\u2557\u255a\u255d\u2560\u2563\u2566\u2569\u256c":
                is_table_line = True

            if is_table_line:
                if not all(c in "\u250c\u2510\u2514\u2518\u251c\u2524\u252c\u2534\u253c\u2550\u2500\u2502\u2503\u2554\u2557\u255a\u255d\u2560\u2563\u2566\u2569\u256c " for c in stripped):
                    current_table.append(stripped)
                    in_table = True
            else:
                if in_table and len(current_table) >= 3:
                    table = self._parse_text_table(current_table)
                    if table:
                        tables.append(table)
                current_table = []
                in_table = False

        if in_table and len(current_table) >= 3:
            table = self._parse_text_table(current_table)
            if table:
                tables.append(table)

        return tables

    def _parse_text_table(self, lines):
        """从文本表格行解析为 ContentBlock(type=table)。"""
        if not lines:
            return None

        pipe_lines = [l for l in lines if "|" in l or "\u2502" in l]
        if pipe_lines:
            import re
            data_lines = [l for l in pipe_lines if not re.match(r"^[\s\|\u2502\-\u2501\u2550\u2500\+]+$", l)]
            if len(data_lines) < 2:
                return None
            parsed_rows = []
            for line in data_lines:
                cells = [c.strip() for c in re.split(r"[\||\u2502]", line) if c.strip()]
                if cells:
                    parsed_rows.append(cells)
            if len(parsed_rows) >= 2:
                block = ContentBlock(ContentBlock.TYPE_TABLE)
                block.headers = parsed_rows[0]
                block.rows = parsed_rows[1:]
                return block

        tab_lines = [l for l in lines if "\t" in l]
        if tab_lines and len(tab_lines) >= 2:
            parsed_rows = []
            for line in tab_lines:
                cells = [c.strip() for c in line.split("\t") if c.strip()]
                if cells:
                    parsed_rows.append(cells)
            if len(parsed_rows) >= 2:
                block = ContentBlock(ContentBlock.TYPE_TABLE)
                block.headers = parsed_rows[0]
                block.rows = parsed_rows[1:]
                return block

        return None


    # ========== 语义切片 ==========

    CHUNK_MIN_CHARS = 200
    CHUNK_MAX_CHARS = 1500

    def semantic_chunk(self, doc: StructuredDocument) -> list[dict]:
        """按标题/表格自然边界切片。

        Args:
            doc: 结构化文档

        Returns:
            list[dict]: 每个元素包含 text, section_path, content_type, page_range, metadata
        """
        chunks = []
        self._chunk_sections(doc.sections, [], chunks)
        return chunks

    def _chunk_sections(self, sections: list, parent_path: list, chunks: list):
        """递归遍历章节，生成切片。"""
        for section in sections:
            path = parent_path + [section.title] if section.title else parent_path
            section_path = " > ".join(path) if path else ""

            # 如果该章节有独立内容，作为一个 chunk
            if section.content:
                texts = []
                content_types = set()
                for block in section.content:
                    if block.type == ContentBlock.TYPE_TABLE:
                        texts.append(self._table_to_text(block))
                    elif block.text:
                        texts.append(block.text)
                    content_types.add(block.type)

                combined = "\n".join(texts)
                if combined and len(combined) >= self.CHUNK_MIN_CHARS:
                    chunks.append({
                        "text": combined,
                        "section_path": section_path,
                        "content_type": "mixed" if len(content_types) > 1 else (content_types.pop() if content_types else "paragraph"),
                        "page_range": section.page_range,
                        "metadata": {"section_level": section.level},
                    })
                elif combined and len(combined) < self.CHUNK_MIN_CHARS and path:
                    # 短内容合并到前一个 chunk
                    if chunks and chunks[-1].get("section_path", "").startswith(section_path.rsplit(" > ", 1)[0] if " > " in section_path else ""):
                        chunks[-1]["text"] += "\n" + combined
                    else:
                        chunks.append({
                            "text": combined,
                            "section_path": section_path,
                            "content_type": "paragraph",
                            "page_range": section.page_range,
                            "metadata": {"section_level": section.level},
                        })

            # 递归子章节
            if section.children:
                self._chunk_sections(section.children, path, chunks)

    def _table_to_text(self, block: ContentBlock) -> str:
        """将表格块转为结构化文本。"""
        lines = []
        if block.headers:
            lines.append(" | ".join(block.headers))
            lines.append(" | ".join(["---"] * len(block.headers)))
        for row in block.rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    # ========== 旧版兼容 ==========

    def parse_bytes(self, filename: str, payload: bytes) -> str:
        """兼容旧接口：返回纯文本。"""
        doc = self.parse_structured(filename, payload)
        texts = []
        for chunk in self.semantic_chunk(doc):
            texts.append(chunk["text"])
        return "\n".join(texts)

    def split_text_chunks(self, text: str, max_length=1200, overlap=120):
        """兼容旧接口。"""
        return self.semantic_chunk(self._text_to_doc(text))

    def _text_to_doc(self, text: str) -> StructuredDocument:
        doc = StructuredDocument()
        self._build_sections_from_text(text, doc)
        return doc

    # ========== 辅助方法 ==========

    def _parse_doc_structured(self, payload: bytes, doc: StructuredDocument):
        """解析旧版 DOC 格式。"""
        import subprocess
        import tempfile
        from pathlib import Path

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(payload)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", tmp_path],
                capture_output=True, text=True, check=False, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                # 尝试检测表格
                table_block = self._detect_table_in_text(text)
                if table_block:
                    # 有表格：把表格信息嵌入文本后再建章节
                    table_text = "表格内容：\n"
                    if table_block.headers:
                        table_text += " | ".join(table_block.headers) + "\n"
                    for row in table_block.rows:
                        table_text += " | ".join(row) + "\n"
                    text = text + "\n" + table_text
                self._build_sections_from_text(text, doc)
            else:
                text = payload.decode("utf-8", errors="replace")
                section = Section(title="全文", level=1)
                section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
                doc.sections.append(section)
        except Exception as exc:
            text = payload.decode("utf-8", errors="replace")
            section = Section(title="全文", level=1)
            section.content.append(ContentBlock(ContentBlock.TYPE_PARAGRAPH, text))
            doc.sections.append(section)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def _parse_spreadsheet_structured(self, payload: bytes, doc: StructuredDocument, ext: str):
        """解析电子表格。"""
        namespace = {
            "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        }
        try:
            with ZipFile(BytesIO(payload)) as z:
                # 读取共享字符串
                shared_strings = []
                if "xl/sharedStrings.xml" in z.namelist():
                    root = ET.fromstring(z.read("xl/sharedStrings.xml"))
                    for si in root.findall(".//main:si", namespace):
                        parts = [t.text or "" for t in si.findall(".//main:t", namespace)]
                        shared_strings.append("".join(parts))

                # 读取第一个 sheet
                if "xl/worksheets/sheet1.xml" in z.namelist():
                    sheet_root = ET.fromstring(z.read("xl/worksheets/sheet1.xml"))
                    rows = sheet_root.findall(".//main:row", namespace)
                    for row in rows:
                        cells = []
                        for cell in row.findall("main:c", namespace):
                            cell_ref = cell.get("r", "")
                            cell_type = cell.get("t", "")
                            value_elem = cell.find("main:v", namespace)
                            raw_value = value_elem.text if value_elem is not None else ""
                            if cell_type == "s" and raw_value.isdigit() and int(raw_value) < len(shared_strings):
                                cells.append(shared_strings[int(raw_value)])
                            else:
                                cells.append(raw_value)
                        if cells:
                            section = Section(title=f"行 {row.get('r', '')}", level=1)
                            block = ContentBlock(ContentBlock.TYPE_PARAGRAPH, " | ".join(cells))
                            section.content.append(block)
                            doc.sections.append(section)
        except Exception as exc:
            logger.warning("[parser] 电子表格解析失败: %s", exc)

    def _normalize_text(self, text: str) -> str:
        """清理文本。"""
        if not text:
            return ""
        lines = []
        for raw_line in str(text).replace("\x00", "").splitlines():
            line = " ".join(raw_line.split())
            if line:
                lines.append(line)
        return "\n".join(lines).strip()

def _dedup_section_index(index: list) -> list:
    """去重：同名+同级+同父节点 → 保留有内容的版本，删除空洞节点。

    key = (title, level, parent_id) 三元组。
    防止同名但不同位置的真章节被误删。
    """
    content_keys = set()
    content_ids = set()
    for entry in index:
        title = entry.get("title", "") or ""
        if title and _has_real_content(entry):
            key = (title, entry.get("level"), entry.get("parent_id"))
            content_keys.add(key)
            content_ids.add(entry.get("id"))

    clean = []
    seen_keys = set()
    for entry in index:
        title = entry.get("title", "") or ""
        if not title:
            clean.append(entry)
            continue
        key = (title, entry.get("level"), entry.get("parent_id"))
        if key in content_keys:
            if entry.get("id") in content_ids:
                clean.append(entry)
        else:
            clean.append(entry)
    return clean
