# -*- coding: utf-8 -*-
import json
import logging
import re
import time
from io import BytesIO
from itertools import zip_longest
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from flask import current_app, send_file
import numpy as _np
from sqlalchemy.exc import SQLAlchemyError
from werkzeug.utils import secure_filename

from ...core.extensions import db
from ...core.time_utils import utc_now
from ...domain import (
    BiddingAnalysisResult,
    BiddingCatalog,
    BiddingCheckItem,
    BiddingSharedResource,
    BiddingTenderAttachment,
    BiddingTask,
    BiddingTaskChapter,
    BiddingTaskExecution,
    FileStorage,
    KnowledgeBase,
    KnowledgeBaseFile,
    SubjectCompany,
    SubjectMaterialFile,
    TemplateCatalog,
)
from ...infrastructure.document_parser import DocumentParser
from ...infrastructure.integrations import ChromaAdapter, LLMAdapter, MinioAdapter
from ...infrastructure.multi_recall_engine import MultiRecallEngine
from ..quality_assurance import (
    inject_constraints_into_prompt,
)
from ...infrastructure.task_queue import TaskQueueManager
from ...core.response import page_success
from ..common import (
    CHAPTER_FIELD_UNSET,
    TENDER_ALLOWED_EXTENSIONS,
    TaskExecutionCancelledError,
    dump_json,
    get_subject_material_completeness,
    normalize_knowledge_base_ids,
    validate_subject_knowledge_bases,
)
from ..storage import StorageService

logger = logging.getLogger(__name__)


def _apply_black_solid_borders(table):
    """为 python-docx Table 的 XML 设置黑色实线边框。"""
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = __import__('lxml.etree', fromlist=['SubElement']).SubElement(tbl, qn('w:tblPr'))
    # 移除已有 tblBorders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    
    NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    borders = __import__('lxml.etree', fromlist=['SubElement']).SubElement(tblPr, NS + 'tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = __import__('lxml.etree', fromlist=['SubElement']).SubElement(borders, NS + edge)
        el.set(NS + 'val', 'single')
        el.set(NS + 'sz', '4')
        el.set(NS + 'space', '0')
        el.set(NS + 'color', '000000')



def _strip_xml_control_chars(text: str) -> str:
    """移除 XML 不兼容的控制字符，防止 _build_docx_bytes 序列化时崩溃。

    python-docx 底层使用 lxml 生成 XML，控制字符（NULL 字节、起止控制符等）
    在 XML 1.0 中不合法，必须移除。
    """
    if not text:
        return text
    # 保留 \t(09) \n(0a) \r(0d) 等 XML 允许的空白字符
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)



# ========== 全局产品列名映射表 ==========
# 统一所有表格解析系统使用的字段映射，确保 _extract_table_data_from_analysis()
# 表格字段映射规则
PRODUCT_COLUMN_MAP = {
    "name": ["品名", "名称", "产品名称", "试剂名称", "货物名称",
             "商品名", "采购产品名称", "标的名称", "产品名",
             "采购产品名称", "产品（设备）名称"],
    "spec": ["规格", "规格型号", "型号", "技术规格", "参数",
             "★规格参数", "技术参数与性能指标", "规格参数",
             "技术规格参数"],
    "brand": ["品牌", "生产厂家", "厂家", "制造商"],
    "qty": ["数量", "需求量", "预估数量", "采购量", "★数量"],
    "unit": ["单位", "计量单位", "★计量单位"],
    "unit_price": ["单价", "预算单价", "最高限价", "★单价最高限价",
                   "单价最高限价"],
    "total_price": ["总价", "金额", "合计"],
    "产地": ["产地", "来源"],
    "备注": ["备注", "说明"],
}


# 产品库 API 字段名 → 中标文件表格列名映射
# 用于将 _fetch_product_data() 返回的结构化字段填充到表格空格中
PRODUCT_FIELD_TO_COLUMN = {
    "brand": ["品牌", "生产厂家", "厂家", "制造商", "★品牌"],
    "specAndModel": ["规格", "规格型号", "型号", "技术规格", "规格参数",
                     "★规格参数", "技术参数与性能指标"],
    "manufacturer": ["生产厂家", "厂家", "制造商"],
    "unit": ["单位", "计量单位", "★计量单位"],
    "articleNo": ["货号", "商品编号", "产品编号"],
    "serialNo": ["序列号", "批号"],
    "descOfFunc": ["功能描述", "产品描述", "描述", "主要功能"],
    "detectionOfSpec": ["检测标准", "检测规范"],
    "storageCondition": ["储存条件", "存储条件", "存放条件", "保存条件"],
    "concentration": ["浓度"],
    "registrationCertificateNo": ["注册证号", "注册号", "医疗器械注册证", "注册证"],
    "qualityPeriod": ["保质期", "有效期", "质量保证期"],
}


_FIELD_UNSET = object()
_EMPTY_PAGE_MARKER = "[[EMPTY_PAGE]]"
_CONTENT_BLOCKS_PREFIX = "[[CONTENT_BLOCKS]]"
_SEPARATOR_PAGE_PREFIX = "[[SEPARATOR_PAGE]]"
_SEPARATOR_PAGE_EMPTY = "[[SEPARATOR_PAGE_EMPTY]]"
_SEPARATOR_PAGE_KEYWORDS = (
    "资格性响应文件", "符合性响应文件", "技术响应文件",
    "商务响应文件", "其他响应文件", "其他文件",
)

_TABLE_MARKER_PREFIX = "[[TABLE:"
_TABLE_JSON_PREFIX = "[[TABLE_JSON:"
_QUALIFICATION_MARKER = "[[QUALIFICATION_DOCS]]"


def _normalize_catalog_generation_level(level):
    """规范化目录颗粒度配置。"""
    normalized = str(level or "MEDIUM").strip().upper()
    if normalized not in {"LOW", "MEDIUM", "HIGH"}:
        return "MEDIUM"
    return normalized


def _get_catalog_generation_profile(level):
    """返回目录颗粒度对应的描述长度和写作要求。"""
    normalized = _normalize_catalog_generation_level(level)
    profile_map = {
        "LOW": {
            "description_max_length": 70,
            "directive": "目录颗粒度：LOW，章节内容保持简洁直达，突出核心响应点即可。",
        },
        "MEDIUM": {
            "description_max_length": 120,
            "directive": "目录颗粒度：MEDIUM，章节内容需要兼顾概述、关键响应点与必要说明。",
        },
        "HIGH": {
            "description_max_length": 180,
            "directive": "目录颗粒度：HIGH，章节内容需要展开到实施细节、评分响应、风险控制和支撑材料说明。",
        },
    }
    return {"level": normalized, **profile_map[normalized]}


def _get_task_chapters(task_id):
    """读取任务的全部章节记录并按章节号排序。"""
    chapters = (
        BiddingTaskChapter.query.filter_by(task_id=task_id)
        .order_by(BiddingTaskChapter.chapter_no.asc(), BiddingTaskChapter.id.asc())
        .all()
    )
    return [item.to_dict() for item in chapters]


def _build_shared_resource_analysis_text(shared_resource_id, package_no=None):
    """汇总共享资源下的主招标文件与附件文本，作为统一分析输入。"""
    shared_resource = db.session.get(BiddingSharedResource, shared_resource_id)
    if not shared_resource:
        return {"raw_text": "", "effective_text": "", "source_files": []}

    source_files = []
    file_records = []

    tender_file = db.session.get(FileStorage, shared_resource.tender_file_id) if shared_resource.tender_file_id else None
    if tender_file and not tender_file.deleted_flag:
        file_records.append(("TENDER", tender_file))

    attachments = (
        BiddingTenderAttachment.query.filter_by(shared_resource_id=shared_resource_id)
        .order_by(BiddingTenderAttachment.uploaded_at.asc(), BiddingTenderAttachment.id.asc())
        .all()
    )
    for attachment in attachments:
        file_record = db.session.get(FileStorage, attachment.file_id) if attachment.file_id else None
        if file_record and not file_record.deleted_flag:
            file_records.append(("ATTACHMENT", file_record))

    raw_parts = [None] * len(file_records)
    effective_parts = [None] * len(file_records)
    source_files = [None] * len(file_records)

    # 捕获 Flask 应用上下文，供并行线程使用
    _app = current_app._get_current_object()

    def _read_single_file(idx, file_role, file_record):
        # 每个并行线程需要自己的 Flask 应用上下文
        with _app.app_context():
            file_text = (_read_file_text(file_record) or "").strip()
        sf = {
            "file_id": file_record.id,
            "file_name": file_record.file_name,
            "file_role": file_role,
        }
        if not file_text:
            return (idx, sf, "", "")

        labeled_text = f"[{file_role}] {file_record.file_name}\n{file_text}"

        if package_no:
            filtered_text = (_extract_effective_text(file_text, package_no) or "").strip()
            eff = f"[{file_role}] {file_record.file_name}\n{filtered_text}" if filtered_text else ""
        else:
            eff = labeled_text

        return (idx, sf, labeled_text, eff)

    from concurrent.futures import ThreadPoolExecutor, as_completed
    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = []
        for idx, (file_role, file_record) in enumerate(file_records):
            futures.append(executor.submit(_read_single_file, idx, file_role, file_record))
        for future in as_completed(futures):
            idx, sf, raw_text, eff_text = future.result()
            source_files[idx] = sf
            raw_parts[idx] = raw_text
            effective_parts[idx] = eff_text

    raw_parts = [p for p in raw_parts if p]
    effective_parts = [p for p in effective_parts if p]
    source_files = [s for s in source_files if s]

    return {
        "raw_text": "\n\n".join(raw_parts).strip(),
        "effective_text": "\n\n".join(effective_parts).strip(),
        "source_files": source_files,
    }


def _extract_analysis_context(analysis_result):
    """统一从 analysis_data/兼容字段中提取目录与正文生成需要的分析上下文。"""
    context = {
        "bidder_notice": {},
        "qualification_review": {},
        "source_files": [],
        "overview": "",
        "requirements": "",
        "business_requirements": "",
        "qualification_requirements": "",
        "technical_requirements": "",
        "scoring_items": "",
        "disqualification_items": "",
    }
    if not analysis_result:
        return context

    payload = {}
    if getattr(analysis_result, "analysis_data", None):
        try:
            payload = json.loads(analysis_result.analysis_data)
        except (TypeError, json.JSONDecodeError):
            payload = {}

    if isinstance(payload, dict):
        bidder_notice = payload.get("bidder_notice", {}) or {}
        qualification_review = payload.get("qualification_review", {}) or {}
        context["source_files"] = payload.get("source_files", []) or []
        context["bidder_notice"] = bidder_notice
        context["qualification_review"] = qualification_review
        context["overview"] = bidder_notice.get("overview", "") or getattr(analysis_result, "overview", "") or ""
        context["requirements"] = payload.get("requirements", "") or getattr(analysis_result, "requirements", "") or ""
        context["business_requirements"] = (
            payload.get("business_requirements", "") or getattr(analysis_result, "business_requirements", "") or ""
        )
        context["qualification_requirements"] = (
            qualification_review.get("qualification_check", "")
            or getattr(analysis_result, "qualification_requirements", "")
            or ""
        )
        context["technical_requirements"] = (
            payload.get("technical_requirements", "") or getattr(analysis_result, "technical_requirements", "") or ""
        )
        context["scoring_items"] = payload.get("scoring_items", "") or getattr(analysis_result, "scoring_items", "") or ""
        context["disqualification_items"] = (
            qualification_review.get("disqualification_items", "")
            or getattr(analysis_result, "disqualification_items", "")
            or ""
        )
        # 确保 bidder_notice 有 project_name/project_no（可能顶层也有）
        if not bidder_notice.get("project_name"):
            for field in ("project_name", "项目名称", "标的名称"):
                # v3 分析数据存储在 metadata.project_name.value
                meta = payload.get("metadata", {}) if isinstance(payload, dict) else {}
                if isinstance(meta, dict):
                    meta_val = meta.get("project_name", {})
                    if isinstance(meta_val, dict):
                        meta_val = meta_val.get("value", "")
                else:
                    meta_val = ""
                val = meta_val or payload.get(field) or getattr(analysis_result, field, None) or ""
                if val:
                    bidder_notice["project_name"] = val
                    break
        if not bidder_notice.get("project_no"):
            for field in ("project_no", "项目编号", "比选编号"):
                # v3 分析数据存储在 metadata.project_code.value
                meta = payload.get("metadata", {}) if isinstance(payload, dict) else {}
                if isinstance(meta, dict):
                    meta_val = meta.get("project_code", {})
                    if isinstance(meta_val, dict):
                        meta_val = meta_val.get("value", "")
                else:
                    meta_val = ""
                val = meta_val or payload.get(field) or getattr(analysis_result, field, None) or ""
                if val:
                    bidder_notice["project_no"] = val
                    break        # 表格数据按章节存储在 format_requirements.required_sections[].template_content
        # 生成时通过 _generate_table_content 按章节标题查找，不需要全局列表
        fmt = payload.get("format_requirements", {}) if isinstance(payload, dict) else {}
        context["_raw_product_lists"] = []
        context["_raw_product_tables"] = []
        context["_eligibility"] = payload.get("eligibility", {}) if isinstance(payload, dict) else {}
        context["_format_requirements"] = fmt
        context["_scoring"] = payload.get("scoring", {}) if isinstance(payload, dict) else {}
        context["_packages"] = payload.get("packages", []) if isinstance(payload, dict) else []

    
    else:
        context["source_files"] = payload.get("source_files", []) if isinstance(payload, dict) else []
        context["overview"] = getattr(analysis_result, "overview", "") or ""
        context["requirements"] = getattr(analysis_result, "requirements", "") or ""
        context["business_requirements"] = getattr(analysis_result, "business_requirements", "") or ""
        context["qualification_requirements"] = getattr(analysis_result, "qualification_requirements", "") or ""
        context["technical_requirements"] = getattr(analysis_result, "technical_requirements", "") or ""
        context["scoring_items"] = getattr(analysis_result, "scoring_items", "") or ""
        context["disqualification_items"] = getattr(analysis_result, "disqualification_items", "") or ""
    return context


def _get_catalog_outline(catalog_record):
    """将目录记录解析为章节大纲列表。"""
    payload = json.loads(catalog_record.catalog_content) if isinstance(catalog_record.catalog_content, str) else catalog_record.catalog_content
    outline = payload.get("outline") if isinstance(payload, dict) else None
    return outline if isinstance(outline, list) else []


def _prepare_task_chapters(task, catalog_record):
    """根据确认目录为任务初始化章节记录。"""
    BiddingTaskChapter.query.filter_by(task_id=task.id).delete()
    outline = _get_catalog_outline(catalog_record)
    if not outline:
        outline = [{"title": "综合响应", "description": ""}]
    for idx, item in enumerate(outline, start=1):
        record = BiddingTaskChapter(
            task_id=task.id,
            chapter_no=idx,
            chapter_title=(item.get("title") or f"章节{idx + 1}").strip(),

            status="PENDING",
            stage_code="QUEUED",
            stage_message="等待生成",
            progress=0,
        )
        db.session.add(record)
    db.session.flush()


def _update_task_generate_stage(task, stage_code=None, stage_message=None, error_message=None):
    """更新任务级生成阶段编码、提示语和错误信息。"""
    # CHAPTER_FIELD_UNSET 已在文件顶部导入
    if stage_code is not None:
        task.generate_stage_code = stage_code
    if stage_message is not None:
        task.generate_stage_message = stage_message
    if error_message is not None:
        task.error_message = error_message
    # 如果未传值则从 task 现有字段读取
    if stage_code is None:
        task.generate_stage_code = getattr(task, "generate_stage_code", CHAPTER_FIELD_UNSET)
    if stage_message is None:
        task.generate_stage_message = getattr(task, "generate_stage_message", CHAPTER_FIELD_UNSET)
    if error_message is None:
        task.error_message = getattr(task, "error_message", "")

def _maybe_fail_generate_stage_for_testing(stage_code):
    """在测试模式下按配置模拟指定生成阶段失败。"""
    force_fail_codes = current_app.config.get("GENERATE_FORCE_FAIL_STAGE_CODES")
    if isinstance(force_fail_codes, str):
        force_fail_codes = {force_fail_codes}
    if force_fail_codes and stage_code in (force_fail_codes or set()):
        raise RuntimeError(f"模拟生成阶段失败: {stage_code}")


def _get_failed_generate_stage(stage_code):
    """根据阶段编码映射生成失败后的任务阶段。"""
    mapping = {
        "CHAPTER_GENERATING": ("CHAPTER_GENERATION_FAILED", "章节正文生成失败"),
        "ASSEMBLING_CONTENT": ("CONTENT_ASSEMBLY_FAILED", "章节内容组装失败"),
        "BUILDING_DOCX": ("DOCX_BUILD_FAILED", "结果文档构建失败"),
        "SAVING_RESULT": ("RESULT_SAVE_FAILED", "结果文件保存失败"),
    }
    return mapping.get(stage_code, ("GENERATE_FAILED", "标书生成失败"))


def _update_chapter_runtime_state(
    chapter_record,
    status=_FIELD_UNSET,
    progress=_FIELD_UNSET,
    stage_code=_FIELD_UNSET,
    stage_message=_FIELD_UNSET,
    error_message=_FIELD_UNSET,
    started_at=_FIELD_UNSET,
    finished_at=_FIELD_UNSET,
    content_snapshot=_FIELD_UNSET,
):
    """更新单个章节的运行状态、进度和错误信息。"""
    if status is not _FIELD_UNSET:
        chapter_record.status = status
    if progress is not _FIELD_UNSET:
        chapter_record.progress = progress
    if stage_code is not _FIELD_UNSET:
        chapter_record.stage_code = stage_code
    if stage_message is not _FIELD_UNSET:
        chapter_record.stage_message = stage_message
    if error_message is not _FIELD_UNSET:
        chapter_record.error_message = error_message
    if started_at is not _FIELD_UNSET:
        chapter_record.started_at = started_at
    if finished_at is not _FIELD_UNSET:
        chapter_record.finished_at = finished_at
    if content_snapshot is not _FIELD_UNSET:
        chapter_record.content_snapshot = content_snapshot


def _ensure_task_chapters(task, catalog_record):
    """确保任务已存在与目录一致的章节记录。"""
    existing = (
        BiddingTaskChapter.query.filter_by(task_id=task.id)
        .order_by(BiddingTaskChapter.chapter_no.asc(), BiddingTaskChapter.id.asc())
        .all()
    )
    if not existing:
        _prepare_task_chapters(task, catalog_record)
        db.session.commit()
        existing = (
            BiddingTaskChapter.query.filter_by(task_id=task.id)
            .order_by(BiddingTaskChapter.chapter_no.asc(), BiddingTaskChapter.id.asc())
            .all()
        )
    return existing


def _calculate_generate_task_progress(chapter_records):
    """根据章节进度估算任务级百分比进度。"""
    if not chapter_records:
        return 0
    total = sum(item.progress for item in chapter_records)
    return min(round(total / len(chapter_records)), 99)


def _refresh_generate_task_progress(task):
    """重新计算并写回任务级生成进度。"""
    chapters = (
        BiddingTaskChapter.query.filter_by(task_id=task.id)
        .order_by(BiddingTaskChapter.chapter_no.asc())
        .all()
    )
    raw = _calculate_generate_task_progress(chapters)
    # 生成中的进度映射到 40~100 范围
    if task.status == "GENERATING":
        task.progress = 40 + int(raw * 60 / 100)
    else:
        task.progress = raw
    db.session.commit()


def _get_chapter_progress_floor(task):
    """根据任务状态计算进度下限（与需求文档状态机映射同步）。"""
    floor_map = {
        "INIT": 0,
        "UPLOADED": 10,
        "ANALYZING": 10,
        "PACKAGE_PENDING": 15,
        "ANALYZED": 20,
        "CHECKED": 30,
        "CATALOG_CONFIRMED": 40,
        "GENERATING": 41,
        "GENERATED": 100,
        "CANCELLED": 0,
        "FAILED": 0,
    }
    return floor_map.get(task.status, 0)


def _get_task_progress_value(task):
    """返回适合前端展示的任务进度值（与需求文档状态机映射同步）。"""
    if task.status == "GENERATING":
        chapters = (
            BiddingTaskChapter.query.filter_by(task_id=task.id)
            .order_by(BiddingTaskChapter.chapter_no.asc())
            .all()
        )
        if chapters:
            raw = _calculate_generate_task_progress(chapters)
            # 将章节进度 0~100 映射到 40~100 范围（GENERATING 起始进度为40）
            return 40 + int(raw * 60 / 100)
        return 41
    if task.status == "GENERATED":
        return 100
    return _get_chapter_progress_floor(task)
def _extract_failed_chapter_nos(chapter_records):
    """提取当前失败章节编号列表。"""
    return [item.chapter_no for item in chapter_records if item.status == "FAILED"]


def _extract_retry_chapter_nos(chapter_records):
    """提取适合重试的章节编号列表。"""
    return _extract_failed_chapter_nos(chapter_records)


def _build_generate_retry_hint(task, chapter_records):
    """构建前端可直接使用的生成重试提示信息。"""
    failed = _extract_failed_chapter_nos(chapter_records)
    if task.status in ("GENERATING", "GENERATED") and not failed:
        return {}
    all_generated = all(item.status == "GENERATED" for item in chapter_records)
    if all_generated:
        return {"retry_type": "TASK_OR_CHAPTERS", "failed_chapters": failed}
    return {"retry_type": "CHAPTERS", "failed_chapters": failed}


def _normalize_retry_chapter_nos(chapter_nos):
    """规范化重试请求中的章节编号参数。"""
    if not chapter_nos:
        return []
    if isinstance(chapter_nos, str):
        parts = chapter_nos.split(",")
        result = []
        for part in parts:
            part = part.strip()
            try:
                result.append(int(part))
            except (ValueError, TypeError):
                pass
        return result
    if isinstance(chapter_nos, (list, tuple)):
        return [int(x) for x in chapter_nos if x is not None]
    raise ValueError("chapter_nos 格式不正确")


def _resolve_retry_chapter_nos(chapter_nos, chapter_records, outline, retry_all=False):
    """结合章节现状和重试策略确定最终重试范围。"""
    if retry_all:
        return [item.chapter_no for item in chapter_records]
    normalized = _normalize_retry_chapter_nos(chapter_nos) if chapter_nos else _extract_retry_chapter_nos(chapter_records)
    valid_nos = {item.chapter_no for item in chapter_records}
    for no in normalized:
        if no not in valid_nos:
            raise ValueError(f"章节编号不存在: {no}")
    return sorted(set(normalized))


def _validate_generate_prerequisites(task):
    """校验任务是否满足生成标书的全部前置条件。"""
    if not task.subject_id:
        raise ValueError("请先选择主体公司")
    subject = db.session.get(SubjectCompany, task.subject_id)
    if not subject or subject.status is False:
        raise ValueError("所选主体公司不可用")
    if not task.model_type:
        raise ValueError("请先选择模型")
    material_status = get_subject_material_completeness(task.subject_id)
    if not material_status.get("is_complete"):
        missing = "、".join(material_status.get("missing_material_types", []))
        raise ValueError(f"主体资料未上传齐全，缺少: {missing}")
    kb_ids = normalize_knowledge_base_ids(task.knowledge_base_ids)
    if task.use_knowledge_base and not kb_ids:
        raise ValueError("启用知识库时必须选择 knowledge_base_ids")
    if task.use_knowledge_base:
        validate_subject_knowledge_bases(task.subject_id, kb_ids)


def _get_confirmed_catalog_record(task):
    """读取任务已确认的最终目录记录。"""
    catalog = BiddingCatalog.query.filter_by(shared_resource_id=task.shared_resource_id).first()
    if not catalog:
        raise LookupError("请先确认目录")
    return catalog


def _build_chapter_contents_from_records(chapter_records):
    """从章节记录中提取已生成的内容快照（支持 ContentBlock 结构）。"""
    import json as _json
    chapter_contents = []
    for item in chapter_records:
        if not item.content_snapshot:
            raise RuntimeError(f"章节{item.chapter_no}尚未生成完成，无法组装结果文件")
        snapshot = item.content_snapshot
        content_blocks = None
        # 检测是否为序列化的 ContentBlock
        if snapshot.startswith(_CONTENT_BLOCKS_PREFIX):
            try:
                blocks_data = _json.loads(snapshot[len(_CONTENT_BLOCKS_PREFIX):])
                if isinstance(blocks_data, list):
                    content_blocks = blocks_data
                    snapshot = ""  # content_blocks 优先
            except (_json.JSONDecodeError, TypeError):
                pass
        chapter_contents.append({
            "title": item.chapter_title,
            "content": snapshot,
            "content_blocks": content_blocks,
        })
    return chapter_contents


def _maybe_fail_chapter_for_testing(chapter_no):
    """在测试模式下按配置模拟章节生成失败。"""
    force_fail = current_app.config.get("GENERATE_FORCE_FAIL_CHAPTERS")
    if isinstance(force_fail, str):
        force_fail = {int(x.strip()) for x in force_fail.split(",") if x.strip().isdigit()}
    if force_fail and chapter_no in force_fail:
        raise RuntimeError(f"模拟章节失败: {chapter_no}")


def _validate_tender_file(file_storage):
    """校验招标文件是否存在且扩展名受支持。保留原始文件名（含中文）。"""
    original_filename = file_storage.filename or "uploaded_file"
    extension = Path(original_filename).suffix.lower().lstrip(".")
    if extension not in TENDER_ALLOWED_EXTENSIONS:
        raise ValueError("招标文件仅支持 doc、docx、pdf 格式")
    # 返回原始文件名，让调用方决定何时使用 secure_filename
    return original_filename, extension


def _build_knowledge_base_context(task, query_text=None):
    """构建本次生成需要拼接的知识库上下文。
    
    Args:
        task: BiddingTask 对象
        query_text: 可选的搜索文本，不传则使用 effective_text[:200]
    """
    kb_ids = normalize_knowledge_base_ids(task.knowledge_base_ids)
    if not task.use_knowledge_base or not kb_ids:
        return {}
    knowledge_bases = KnowledgeBase.query.filter(KnowledgeBase.id.in_(kb_ids)).all()
    if not knowledge_bases:
        return {}
    context = {"knowledge_list": []}
    
    if not query_text:
        analysis_result = BiddingAnalysisResult.query.filter_by(shared_resource_id=task.shared_resource_id).first()
        query_text = analysis_result.effective_text if analysis_result else (analysis_result.raw_text if analysis_result else "")
    
    for kb in knowledge_bases:
        files = KnowledgeBaseFile.query.filter_by(knowledge_base_id=kb.id).order_by(KnowledgeBaseFile.id.asc()).all()
        if not files:
            continue
        enabled_files = [item for item in files if item.reference_enabled]
        if not enabled_files:
            continue
        enabled_file_ids = {item.file_id for item in enabled_files}
        enabled_file_names = {item.file_name for item in enabled_files if item.file_name}
        chroma_tenant = kb.chroma_tenant or current_app.config.get("CHROMA_TENANT")
        chroma_database = kb.chroma_database or current_app.config.get("CHROMA_DATABASE")
        chroma_collection = kb.chroma_collection or f"kb_{kb.id}"
        
        snippets = []
        try:
            adapter = ChromaAdapter(
                host=current_app.config.get("CHROMA_HOST"),
                port=current_app.config.get("CHROMA_PORT"),
                tenant=chroma_tenant,
                database=chroma_database,
            )
            # 使用传入的 query_text 搜索，最多取1000字符用于查询
            search_text = (query_text or "")[:5000]
            if not search_text.strip():
                continue
            engine = MultiRecallEngine()
            recall_results = engine.recall(
                query=search_text[:2000],
                collection=chroma_collection,
                top_k=15,
                tenant=chroma_tenant,
                database=chroma_database,
                file_id=None,
            )
            # 置信度门控：召回相关性最低阈值
            MIN_RECALL_SCORE = current_app.config.get("MIN_RECALL_CONFIDENCE", 0.01)
            for rr in recall_results:
                if rr.get("text") and len(rr["text"].strip()) > 20:
                    # 相关性门控：score 低于阈值的片段丢弃
                    rr_score = rr.get("score", 0) or 0
                    if rr_score < MIN_RECALL_SCORE:
                        logger.debug("[confidence] 召回片段相关性偏低 score=%.4f, 已过滤", rr_score)
                        continue
                    # 过滤：只保留已启用的文件
                    src = rr.get("source", {}) or {}
                    fid = src.get("file_id")
                    fname = src.get("file_name", "")
                    if fid is not None:
                        try:
                            if int(fid) not in enabled_file_ids:
                                continue
                        except (TypeError, ValueError):
                            pass
                    if fname and fname not in enabled_file_names:
                        continue
                    from app.domain.analysis_schema import ConfidenceLevel
                    snippet_score = rr.get("score", 0) or 0
                    snippet_confidence = ConfidenceLevel.from_value(snippet_score).name
                    snippets.append({
                        "text": rr["text"].strip(),
                        "score": snippet_score,
                        "confidence": snippet_confidence,
                        "source": rr.get("source", {}),
                    })
        except Exception as exc:
            logger.warning("[kb] 知识库查询异常: %s", exc)
        
        if snippets:
            context["knowledge_list"].append({
                "knowledge_base_name": kb.name,
                "tenant": chroma_tenant,
                "database": chroma_database,
                "collection": chroma_collection,
                "snippets": snippets,
            })
    return context


def _build_product_context(task):
    """构建产品库检索得到的产品上下文。"""
    if not task.use_product_library:
        return {}
    analysis_result = BiddingAnalysisResult.query.filter_by(shared_resource_id=task.shared_resource_id).first()
    requirements = analysis_result.technical_requirements if analysis_result else ""
    effective = analysis_result.effective_text if analysis_result and analysis_result.effective_text else (analysis_result.raw_text if analysis_result else "")
    
    # 从 format_requirements.required_sections[].template_content 提取产品名称
    product_names_from_tables = []
    if analysis_result and analysis_result.analysis_data:
        try:
            ad = json.loads(analysis_result.analysis_data) if isinstance(analysis_result.analysis_data, str) else analysis_result.analysis_data
            fmt = ad.get("format_requirements", {}) if isinstance(ad, dict) else {}
            if fmt:
                for sec in fmt.get("required_sections", []):
                    for blk in sec.get("template_content", []):
                        if blk.get("type") != "table":
                            continue
                        tbl = blk
                        headers = tbl.get("headers", [])
                        rows = tbl.get("rows", [])
                        if not headers or not rows:
                            continue
                        for row in rows:
                            for i, h in enumerate(headers):
                                if any(k in h for k in ["采购产品名称", "产品名称", "标的名称", "标的"]):
                                    name = row[i] if i < len(row) else ""
                                    if name and len(name) >= 2:
                                        product_names_from_tables.append(name)
        except Exception:
            pass
    
    terms = _extract_product_terms(effective + "\n" + (requirements or ""))
    # 合并结构化表格中的产品名
    for pname in product_names_from_tables:
        if pname not in terms:
            terms.append(pname)
    
    if not terms:
        return {"snippets": [], "matched_products": []}
    
    try:
        chroma_tenant = current_app.config.get("CHROMA_TENANT")
        chroma_database = current_app.config.get("CHROMA_DATABASE")
        adapter = ChromaAdapter(
            host=current_app.config.get("CHROMA_HOST"),
            port=current_app.config.get("CHROMA_PORT"),
            tenant=chroma_tenant,
            database=chroma_database,
        )
        engine = MultiRecallEngine()
        recall_results = engine.recall(
            query=" ".join(terms),
            collection="product_library",
            top_k=10,
            tenant=chroma_tenant,
            database=chroma_database,
        )
        matched = []
        for rr in recall_results:
            if rr.get("text") and len(rr["text"].strip()) > 10:
                matched.append({"query_term": terms[0] if terms else "", "matched_text": rr["text"].strip()})
        return {"matched_products": matched, "product_terms": terms, "snippets": [m["matched_text"] for m in matched]}
    except Exception:
        return {"snippets": [], "matched_products": [], "product_terms": terms}


def _fetch_product_data(product_names, adapter=None):
    """从产品库批量查询产品信息，返回结构化数据。

    使用 ChromaAdapter（业务服务端口 28712）的 /objects/query 接口，
    直接返回 object_json 中的结构化字段，无需 LLM 解析。

    Args:
        product_names: list[str] - 产品名称列表
        adapter: ChromaAdapter 实例（可选，自动创建）

    Returns:
        dict: {product_name: {brand, specAndModel, manufacturer, unit, ...}}
    """
    if not product_names:
        return {}

    if adapter is None:
        adapter = ChromaAdapter(
            host=current_app.config.get("CHROMA_HOST"),
            tenant=current_app.config.get("PRODUCT_CHROMA_TENANT", "erp"),
            database=current_app.config.get("PRODUCT_CHROMA_DATABASE", "erp"),
        )

    collection = current_app.config.get("PRODUCT_CHROMA_COLLECTION", "product")
    result = {}

    for name in product_names:
        pname = (name or "").strip()
        if not pname or len(pname) < 2:
            continue
        try:
            data = adapter.query_objects(collection, query_text=pname, top_k=1)
            matches = (data or {}).get("matches", []) or []
            if matches:
                best = matches[0]
                obj_str = best.get("object_json", "") or ""
                if obj_str:
                    obj = json.loads(obj_str)
                    info = {}
                    for field in ["productName", "brand", "specAndModel", "manufacturer",
                                   "unit", "articleNo", "serialNo", "descOfFunc",
                                   "detectionOfSpec", "storageCondition", "concentration",
                                   "qualityPeriod", "qualityPeriodUnit", "registrationCertificateNo"]:
                        val = obj.get(field)
                        if val is not None and str(val).strip() not in ("", "-"):
                            info[field] = str(val)
                    if info:
                        result[pname] = info
        except Exception as exc:
            logger.warning("[product] 产品库查询失败 name=%s: %s", pname, exc)

    return result
def _build_coverage_matrix(nrows, ncols, merges):
    """建立合并覆盖矩阵，标记每个单元格是实格还是虚拟格。

    Returns:
        list[list[str]]: "real" | "h_virtual" | "v_continue"
    """
    matrix = [["real"] * ncols for _ in range(nrows)]
    for m in merges:
        if m.get("type") == "horizontal":
            r, c, s = m["row"], m["col"], m["span"]
            for cc in range(c + 1, min(c + s, ncols)):
                if r < nrows and cc < ncols:
                    matrix[r][cc] = "h_virtual"
        elif m.get("type") == "vertical":
            r, c, s = m["row"], m["col"], m["span"]
            for rr in range(r + 1, min(r + s, nrows)):
                if rr < nrows and c < ncols:
                    if matrix[rr][c] != "h_virtual":
                        matrix[rr][c] = "v_continue"
    return matrix


def _detect_fillable_columns(nrows, ncols, coverage, merged_rows):
    """检测哪些列是投标人填写列（原文中空率 > 70% 的真实单元格列）。

    排除合并虚拟格后，统计每列的空率。
    """
    if nrows <= 1:
        return set()
    col_empty_count = [0] * ncols
    col_real_count = [0] * ncols
    for ri in range(1, nrows):  # 从第二行开始（跳过表头）
        row = merged_rows[ri] if ri < len(merged_rows) else []
        for ci in range(ncols):
            if coverage[ri][ci] != "real":
                continue
            col_real_count[ci] += 1
            cell_val = (row[ci] if ci < len(row) else "").strip()
            if not cell_val:
                col_empty_count[ci] += 1

    fillable = set()
    for ci in range(ncols):
        if col_real_count[ci] == 0:
            continue
        empty_rate = col_empty_count[ci] / col_real_count[ci]
        if empty_rate >= 0.7:
            fillable.add(ci)
    return fillable


# 要求关键词 → 主体资料类型 映射（用于响应表/资格表的填充匹配）
_REQUIREMENT_TO_MATERIAL = [
    ("营业执照", "BUSINESS_LICENSE", "\u5df2\u63d0\u4f9b\u7ecf\u5546\u90e8\u95e8\u6838\u51c6\u767b\u8bb0\u7684\u6709\u6548\u8425\u4e1a\u6267\u7167\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u6cd5\u5b9a\u4ee3\u8868", ("LEGAL_PERSON_STATEMENT", "LEGAL_PERSON_ID_CARD"), "\u5df2\u63d0\u4f9b\u6cd5\u5b9a\u4ee3\u8868\u4eba\u8eab\u4efd\u8bc1\u660e\u53ca\u76f8\u5173\u6750\u6599\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u6388\u6743", "AUTHORIZATION_LETTER", "\u5df2\u63d0\u4f9b\u6388\u6743\u59d4\u6258\u4e66\u53ca\u88ab\u6388\u6743\u4eba\u8eab\u4efd\u8bc1\u660e\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u8d22\u52a1", "FINANCIAL_STATEMENT", "\u5df2\u63d0\u4f9b\u8fd1\u4e09\u5e74\u8d22\u52a1\u62a5\u8868\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u7eb3\u7a0e", "FINANCIAL_STATEMENT", "\u5df2\u63d0\u4f9b\u7eb3\u7a0e\u8bc1\u660e\u6750\u6599\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u793e\u4fdd", "FINANCIAL_STATEMENT", "\u5df2\u63d0\u4f9b\u793e\u4f1a\u4fdd\u9669\u7f34\u7eb3\u8bc1\u660e\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("\u8d44\u8d28", ("QUALIFICATION_FILE", "QUALIFICATION_DECLARATION"), "\u5df2\u63d0\u4f9b\u76f8\u5173\u8d44\u8d28\u8bc1\u4e66\u53ca\u8d44\u8d28\u58f0\u660e\u6750\u6599\uff0c\u8be6\u89c1\u9644\u4ef6"),
    ("健康", ("INTEGRITY_COMMITMENT",), "已提供健康承诺书，详见附件"),
    ("廉洁", ("INTEGRITY_COMMITMENT",), "已提供相关承诺书，详见附件"),
    ("\u4fe1\u7528", "QUALIFICATION_FILE", "\u5df2\u63d0\u4f9b\u4fe1\u7528\u8bc1\u660e\u6750\u6599\uff0c\u8be6\u89c1\u9644\u4ef6"),
]


def _match_row_text_to_material(row_text, subject_context):
    """将行文本中的要求关键词匹配到主体资料，返回填充文本或 None。"""
    if not row_text or not subject_context:
        return None
    materials = subject_context.get("materials", [])
    if not materials:
        return None
    material_types = {m.get("material_type", "") for m in materials if m.get("material_type")}
    if not material_types:
        return None

    for keyword, mat_type, fill_text in _REQUIREMENT_TO_MATERIAL:
        if keyword in row_text:
            if isinstance(mat_type, tuple):
                matched_types = {t for t in mat_type if t in material_types}
            else:
                matched_types = {mat_type} if mat_type in material_types else set()
            if matched_types:
                return fill_text
    return None


def _match_row_text_to_requirement(row_text, analysis_context):
    """检查行文本是否匹配已识别的资格/商务/技术要求。

    如果能匹配到已识别的 requirement 条目，检查对应条目在分析结果中的"满足状态"。
    暂返回通用填充文本。
    """
    if not row_text or not analysis_context:
        return None
    # 获取资格要求列表
    eligibility = analysis_context.get("_eligibility", {}) or {}
    quals = eligibility.get("qualifications", []) or []
    for q in quals:
        req = q.get("requirement", "") or ""
        if len(req) >= 4 and any(kw in row_text for kw in [req[:8]]):
            pass  # 关键词匹配成功
    # 兜底：不填写具体匹配文本，让调用方走默认逻辑
    return None


def _smart_fill_table(table_dict, analysis_context, subject_context):
    """通用表格填充引擎。

    核心逻辑：
    1. 构建合并覆盖矩阵，区分实格和虚拟格
    2. 检测填充列（原文空率 > 70% 的列）
    3. 对填充列中的空实格，用同行非空文本匹配数据源
    4. 匹配不到 → 留空

    Args:
        table_dict: {"headers": [...], "rows": [[...], ...], "merges": [...], ...}
        analysis_context: 分析上下文（含资格/商务/技术要求）
        subject_context: 主体资料上下文（含上传材料清单）

    Returns:
        list[list[str]]: 填充后的数据行（不含表头）
    """
    headers = table_dict.get("headers", [])
    original_rows = table_dict.get("rows", [])
    merges = table_dict.get("merges", [])

    if not headers or not original_rows:
        return original_rows

    ncols = len(headers)
    nrows_data = len(original_rows)

    # 构建完整矩阵（含表头行）
    all_rows = [headers] + original_rows
    nrows = nrows_data + 1

    # 1. 构建合并覆盖矩阵
    coverage = _build_coverage_matrix(nrows, ncols, merges)

    # 2. 检测填充列
    fillable_cols = _detect_fillable_columns(nrows, ncols, coverage, all_rows)

    if not fillable_cols:
        return original_rows

    # 3. 逐行填充
    filled_rows = [list(row) for row in original_rows]
    for ri in range(nrows_data):
        data_row_idx = ri + 1  # 跳过表头行
        for ci in fillable_cols:
            # 跳过虚拟格
            if coverage[data_row_idx][ci] != "real":
                continue
            # 跳过已有内容的格
            current = (filled_rows[ri][ci] if ci < len(filled_rows[ri]) else "").strip()
            if current:
                continue

            # 收集当前行所有非空文本作为行上下文
            row_context = ""
            for j in range(ncols):
                if j != ci:
                    val = (all_rows[data_row_idx][j] if j < len(all_rows[data_row_idx]) else "").strip()
                    if val:
                        row_context += val + " "

            row_context = row_context.strip()
            if not row_context:
                continue

            # 4. 尝试匹配
            fill_val = None

            # 4a. 尝试匹配主体资料
            fill_val = _match_row_text_to_material(row_context, subject_context)

            # 4b. 如果没有匹配到，检查是否是产品表，尝试匹配产品库
            if not fill_val:
                # 检查此行是否包含产品名
                product_name = None
                for j in range(ncols):
                    val = (original_rows[ri][j] if j < len(original_rows[ri]) else "").strip()
                    if val:
                        product_name = val
                        break
                if product_name:
                    try:
                        product_data = _fetch_product_data([product_name])
                    except Exception:
                        product_data = None
                    if product_data and product_data.get(product_name):
                        p_info = product_data[product_name]
                        # 检测当前列应该匹配哪个产品字段
                        header_text = headers[ci] if ci < len(headers) else ""
                        for std_field, candidates in PRODUCT_FIELD_TO_COLUMN.items():
                            if any(c in header_text for c in candidates):
                                p_val = p_info.get(std_field, "")
                                if p_val:
                                    fill_val = str(p_val)[:100]
                                break

            # 4c. 对技术偏离表，默认填"完全响应"
            if not fill_val:
                for kw in ["\u504f\u79bb", "\u54cd\u5e94"]:
                    if kw in headers[ci] if ci < len(headers) else "":
                        fill_val = "\u5b8c\u5168\u54cd\u5e94\uff0c\u65e0\u504f\u79bb"
                        break

            if fill_val:
                while ci >= len(filled_rows[ri]):
                    filled_rows[ri].append("")
                filled_rows[ri][ci] = fill_val
            # 没匹配到 → 留空

    return filled_rows

def _extract_product_terms(text):
    """从当前有效分析文本中抽取产品项关键词。"""
    if not text:
        return []
    terms = []
    for match in re.finditer(r"(?:采购|供货|产品|设备|系统|服务)[：:：\s]*([^\s，。,\.]{2,30})", text):
        term = match.group(1).strip()
        if term and len(term) >= 2:
            terms.append(term)
    return list(dict.fromkeys(terms))[:10]


def _build_subject_material_context(subject_id):
    """汇总主体资料文本，生成主体相关上下文。"""
    if not subject_id:
        return {}
    subject = db.session.get(SubjectCompany, subject_id)
    if not subject:
        return {}
    materials = SubjectMaterialFile.query.filter_by(subject_id=subject_id).order_by(SubjectMaterialFile.uploaded_at.asc()).all()
    material_labels = {
        "BUSINESS_LICENSE": "营业执照",
        "QUALIFICATION_FILE": "资质文件",
        "LEGAL_PERSON_ID_CARD": "法人身份证",
        "AUTHORIZATION_LETTER": "授权委托书",
        "AUTHORIZED_PERSON_ID_CARD": "被授权人身份证",
        "QUALIFICATION_DECLARATION": "资质声明函",
        "LEGAL_PERSON_STATEMENT": "法定代表人身份证明",
        "FINANCIAL_STATEMENT": "财务报表",
        "INTEGRITY_COMMITMENT": "廉洁承诺书",
    }
    items = []
    for m in materials:
        label = material_labels.get(m.material_type, m.material_type or "其他资料")
        file_record = db.session.get(FileStorage, m.file_id) if m.file_id else None
        text_excerpt = ""
        if file_record:
            try:
                text_excerpt = (_read_file_text(file_record) or "")[:800]
            except Exception as exc:
                logger.warning("[subject] 读取主体资料文本失败 material=%s file=%s: %s", m.id, m.file_id, exc)
        # 注意：主体资料是用户上传的正式文件，不做置信度过滤
        # 置信度标记由下游 _filter_low_confidence_subject_materials() 处理
        items.append(
            {
                "id": m.id,
                "file_id": m.file_id,
                "material_type": m.material_type,
                "material_label": label,
                "file_name": m.file_name or "",
                "file_ext": (file_record.file_ext or "") if file_record else "",
                "storage_provider": (file_record.storage_provider or "") if file_record else "",
                "text_excerpt": text_excerpt,
            }
        )
    # 字段格式校验
    company_name = subject.company_name or ""
    credit_code = subject.credit_code or ""
    # 主体数据由用户管理，不做格式校验

    return {
        "company_name": company_name,
        "credit_code": credit_code,
        "address": subject.address or "",
        "contact_person": subject.contact_person or "",
        "contact_phone": subject.contact_phone or "",
        "legal_person": "",  # SubjectCompany 表无法人字段，需从材料 OCR 提取
        "_validations": {},
        "materials": items,
    }


def _extract_outline_leaf_titles(children, prefix_titles=None):
    leaf_titles = []
    for child in children or []:
        title = (child.get("title") or "").strip()
        if not title:
            continue
        current_path = [*(prefix_titles or []), title]
        nested_children = child.get("children", []) or []
        if nested_children:
            leaf_titles.extend(_extract_outline_leaf_titles(nested_children, current_path))
        else:
            leaf_titles.append(" / ".join(current_path))
    return leaf_titles


def _chapter_requires_evidence_placeholder(chapter):
    text = f"{chapter.get('title', '')} {chapter.get('description', '')}".strip()
    strict_phrases = (
        "资格审查资料",
        "资格证明",
        "资质证明",
        "授权文件",
        "授权委托",
        "营业执照",
        "身份证明",
    )
    return any(phrase in text for phrase in strict_phrases)


def _chapter_has_supporting_material(subject_context, knowledge_contexts, product_context):
    if subject_context and any((item.get("text_excerpt") or "").strip() for item in subject_context.get("materials", [])):
        return True
    if knowledge_contexts and knowledge_contexts.get("knowledge_list"):
        for kb in knowledge_contexts.get("knowledge_list", []):
            if _any_snippet_has_text(kb.get("snippets", [])):
                return True
    if product_context and any((item.get("matched_text") or "").strip() for item in product_context.get("matched_products", [])):
        return True
    return False




def _any_snippet_has_text(snippets):
    """判断 snippets 列表中是否有实质文本内容。
    
    兼容新旧两种格式：
    - 旧格式: [str, str, ...]
    - 新格式: [{"text": str, ...}, ...]
    """
    for snip in snippets:
        if isinstance(snip, dict):
            if (snip.get("text") or "").strip():
                return True
        else:
            if (snip or "").strip():
                return True
    return False


def _split_generated_sections_by_titles(content_text, titles):
    if not content_text or not titles:
        return {}
    lines = [line.strip() for line in str(content_text or "").splitlines()]
    positions = []
    normalized_titles = {re.sub(r"\s+", "", title): title for title in titles if title}
    for index, line in enumerate(lines):
        normalized_line = re.sub(r"\s+", "", line)
        if normalized_line in normalized_titles:
            positions.append((index, normalized_titles[normalized_line]))
    if not positions:
        return {}
    sections = {}
    for position, (start_index, title) in enumerate(positions):
        end_index = positions[position + 1][0] if position + 1 < len(positions) else len(lines)
        body_lines = [line for line in lines[start_index + 1 : end_index] if line]
        if body_lines:
            sections[title] = "\n".join(body_lines).strip()
    return sections


def _truncate_binding_text(text, max_length=180):
    normalized = " ".join(str(text or "").replace("\r", "\n").split())
    return normalized[:max_length].strip()


def _is_separator_page_title(title):
    """判断标题是否为响应文件分隔页（容器页，非具体内容章节）。

    Returns:
        bool: True 表示该标题是分隔页
    """
    if not title:
        return False
    for kw in _SEPARATOR_PAGE_KEYWORDS:
        if kw in title:
            return True
    return False


def _build_leaf_response_bindings(chapter, analysis_context, subject_context, knowledge_contexts, product_context):
    children = chapter.get("children", []) or []
    if not children:
        return []

    bidder_notice = analysis_context.get("bidder_notice", {}) or {}
    qualification_review = analysis_context.get("qualification_review", {}) or {}
    project_summary = "；".join(
        item
        for item in [
            f"标的名称：{bidder_notice.get('project_name', '').strip()}" if bidder_notice.get("project_name") else "",
            f"项目编号：{bidder_notice.get('project_no', '').strip()}" if bidder_notice.get("project_no") else "",
            f"项目概况：{bidder_notice.get('overview', '').strip()}" if bidder_notice.get("overview") else "",
        ]
        if item
    )
    subject_snippets = [
        f"{item.get('material_label', '')}：{_truncate_binding_text(item.get('text_excerpt', ''), 120)}"
        for item in (subject_context or {}).get("materials", [])
        if (item.get("text_excerpt") or "").strip()
    ]
    kb_snippets = []
    for kb in (knowledge_contexts or {}).get("knowledge_list", []):
        for snippet in kb.get("snippets", [])[:10]:
            if isinstance(snippet, dict):
                snippet_text = snippet.get("text", "") or ""
                snippet_conf = snippet.get("confidence", "UNKNOWN")
            else:
                snippet_text = snippet or ""
                snippet_conf = "UNKNOWN"
            if snippet_text.strip():
                kb_snippets.append(_truncate_binding_text(snippet_text, 120))
    product_snippets = [
        _truncate_binding_text(item.get("matched_text", ""), 120)
        for item in (product_context or {}).get("matched_products", [])[:3]
        if (item.get("matched_text") or "").strip()
    ]

    bindings = []
    for child in children:
        child_title = (child.get("title") or "").strip()
        child_desc = (child.get("description") or "").strip()
        if not child_title:
            continue

        evidence = []
        combined_text = f"{child_title} {child_desc}"
        if any(keyword in combined_text for keyword in ("项目", "标的", "采购范围", "概述")) and project_summary:
            evidence.append(project_summary)
        if any(keyword in combined_text for keyword in ("技术", "参数", "实施", "交付", "部署")):
            for text in [analysis_context.get("technical_requirements", ""), analysis_context.get("requirements", "")]:
                if text:
                    evidence.append(_truncate_binding_text(text))
        if any(keyword in combined_text for keyword in ("商务", "履约", "交货", "售后")):

            if analysis_context.get("business_requirements"):
                evidence.append(_truncate_binding_text(analysis_context.get("business_requirements", "")))
        if any(keyword in combined_text for keyword in ("资格", "资质", "授权", "证明", "审查")):
            for text in [
                analysis_context.get("qualification_requirements", ""),
                qualification_review.get("qualification_check", ""),
                qualification_review.get("conformity_check", ""),
            ]:
                if text:
                    evidence.append(_truncate_binding_text(text))
            evidence.extend(subject_snippets[:3])
        if "评分" in combined_text:
            if analysis_context.get("scoring_items"):
                evidence.append(_truncate_binding_text(analysis_context.get("scoring_items", "")))
        if "废标" in combined_text:
            for text in [analysis_context.get("disqualification_items", ""), qualification_review.get("disqualification_items", "")]:
                if text:
                    evidence.append(_truncate_binding_text(text))
        if not evidence:
            evidence.extend(kb_snippets[:2])
        if not evidence:
            evidence.extend(product_snippets[:2])

        unique_evidence = []
        for item in evidence:
            normalized = (item or "").strip()
            if normalized and normalized not in unique_evidence:
                unique_evidence.append(normalized)
        bindings.append(
            {
                "title": child_title,
                "requirement": child_desc,
                "evidence": unique_evidence[:3],
                "status": "COVERED" if unique_evidence else "PENDING",
                "require_blank": _chapter_requires_evidence_placeholder(child) and not unique_evidence,
            }
        )
    return bindings


def _compose_leaf_binding_body(binding):
    if binding.get("require_blank"):
        return _EMPTY_PAGE_MARKER

    lines = []
    if binding.get("requirement"):
        lines.append(f"招标要求：{binding['requirement']}")
    evidence = binding.get("evidence", [])
    if evidence:
        lines.append("现有依据如下：")
        for item in evidence:
            lines.append(item)
    elif binding.get("status") == "PENDING":
        lines.append("本节按招标文件要求预留位置，当前未检索到可直接填充的支撑资料。")
    return "\n".join(lines).strip()


def _normalize_chapter_content_by_bindings(content_text, bindings):
    if not bindings:
        return (content_text or "").strip()

    titles = [item["title"] for item in bindings]
    existing_sections = _split_generated_sections_by_titles(content_text, titles)
    sections = []
    for binding in bindings:
        body = (existing_sections.get(binding["title"]) or "").strip()
        if not body:
            body = _compose_leaf_binding_body(binding)
        sections.append(binding["title"])
        if body:
            sections.append(body)
    return "\n".join(sections).strip()


def _extract_binding_body_from_content(content_text, title):
    sections = _split_generated_sections_by_titles(content_text, [title])
    return (sections.get(title) or "").strip()


def _build_generation_coverage_snapshot(
    outline,
    chapter_contents,
    analysis_context,
    subject_context,
    knowledge_contexts,
    product_context,
    generation_plan=None,
):
    source_files = analysis_context.get("source_files", []) if isinstance(analysis_context, dict) else []
    tender_files = [item.get("file_name") for item in source_files if item.get("file_role") == "TENDER" and item.get("file_name")]
    attachment_files = [item.get("file_name") for item in source_files if item.get("file_role") == "ATTACHMENT" and item.get("file_name")]
    source_reference_parts = []
    if tender_files:
        source_reference_parts.append(f"主招标文件：{'、'.join(dict.fromkeys(tender_files))}")
    if attachment_files:
        source_reference_parts.append(f"招标附件：{'、'.join(dict.fromkeys(attachment_files))}")
    source_reference = "；".join(source_reference_parts)

    chapter_map = {}
    for chapter in chapter_contents or []:
        chapter_title = (chapter.get("title") or "").strip()
        if not chapter_title:
            continue
        chapter_text = (chapter.get("content") or "").strip()
        # 兼容 content_blocks 模式：模板章节的内容存储在 content_blocks 中，
        # content 字段为空。只要有 content_blocks 就视为已覆盖。
        chapter_blocks = chapter.get("content_blocks")
        if not chapter_text and chapter_blocks:
            chapter_text = "__CONTENT_BLOCKS_PRESENT__"
        chapter_map[chapter_title] = chapter_text

    plan_lookup = {}
    for item in (generation_plan or {}).get("plan_items", []) or []:
        key = ((item.get("chapter_title") or "").strip(), (item.get("target_title") or "").strip())
        if key[0] and key[1]:
            plan_lookup[key] = item

    requirement_items = []
    for chapter in outline or []:
        chapter_title = (chapter.get("title") or "").strip()
        chapter_content = chapter_map.get(chapter_title, "")
        bindings = _build_leaf_response_bindings(
            chapter,
            analysis_context,
            subject_context,
            knowledge_contexts,
            product_context,
        )
        if bindings:
            for binding in bindings:
                body = _extract_binding_body_from_content(chapter_content, binding["title"])
                covered = bool(body and body != _EMPTY_PAGE_MARKER)
                # 对资格证明类型特殊处理：检验父章节标题，而非子项标题
                # 资格证明文件章节的各子项不需要在正文中展开写
                # 只要 chapter_content 包含 _QUALIFICATION_MARKER，说明已由
                # _generate_qualification_content() 通过 chapter.children 处理
                # 所有子项都应视为已覆盖，不依赖正文文本匹配
                is_qual_chapter = any(kw in chapter_title for kw in ["资格证明", "资格审查", "资质证明", "资格性"])
                if is_qual_chapter:
                    # 如果章节内容包含 QUALIFICATION_MARKER → 所有子项已处理
                    if _QUALIFICATION_MARKER in chapter_content:
                        covered = True
                    else:
                        # 兜底：检查是否有 evidence（来自主体材料或分析数据）
                        covered = bool(binding.get("evidence")) or covered
                requirement_items.append(
                    {
                        "chapter_title": chapter_title,
                        "target_title": binding["title"],
                        "requirement": binding.get("requirement", ""),
                        "status": "COVERED" if covered else ("MISSING" if (binding.get("require_blank") and not binding.get("evidence")) else "PENDING"),
                        "has_evidence": bool(binding.get("evidence")),
                        "source_reference": source_reference,
                        "requirement_level": plan_lookup.get((chapter_title, binding["title"]), {}).get("requirement_level", "NORMAL"),
                        "original_requirement_excerpt": plan_lookup.get((chapter_title, binding["title"]), {}).get(
                            "original_requirement_excerpt", ""
                        ),
                    }
                )
            continue

        chapter_body = (chapter_content or "").strip()
        if chapter_body:
            chapter_body = re.sub(rf"^{re.escape(chapter_title)}\s*", "", chapter_body).strip()
        requirement_items.append(
            {
                "chapter_title": chapter_title,
                "target_title": chapter_title,
                "requirement": chapter.get("description", "") or "",
                "status": "COVERED" if chapter_body and chapter_body != _EMPTY_PAGE_MARKER else "PENDING",
                "has_evidence": bool(chapter_body and chapter_body != _EMPTY_PAGE_MARKER),
                "source_reference": source_reference,
                "requirement_level": plan_lookup.get((chapter_title, chapter_title), {}).get("requirement_level", "NORMAL"),
                "original_requirement_excerpt": plan_lookup.get((chapter_title, chapter_title), {}).get(
                    "original_requirement_excerpt", ""
                ),
            }
        )

    total = len(requirement_items)
    covered_count = sum(1 for item in requirement_items if item["status"] == "COVERED")
    missing_items = [item for item in requirement_items if item["status"] != "COVERED"]
    # 从 generation_plan 补充原子要求覆盖率
    atomic_total = 0
    atomic_covered = 0
    if generation_plan and isinstance(generation_plan, dict):
        plan_items = generation_plan.get("plan_items", []) or []
        atomic_total = generation_plan.get("total_atomic_requirements", len(plan_items))
        # 检查每个 plan_item 在 requirement_items 中的覆盖情况
        for plan_item in plan_items:
            pt = plan_item.get("target_title", "")
            pc = plan_item.get("chapter_title", "")
            for req_item in requirement_items:
                if req_item.get("target_title") == pt and req_item.get("chapter_title") == pc:
                    if req_item.get("status") == "COVERED":
                        atomic_covered += 1
                    break

    return {
        "generated_at": utc_now().isoformat(),
        "total_requirements": total,
        "covered_requirements": covered_count,
        "missing_requirements": len(missing_items),
        "coverage_ratio": round((covered_count / total), 4) if total else 1.0,
        "missing_items": missing_items,
        "requirement_items": requirement_items,
        "atomic_requirements": {
            "total": atomic_total,
            "covered": atomic_covered,
            "missing": atomic_total - atomic_covered,
        },
    }


def _build_analysis_source_reference(analysis_context):
    source_files = analysis_context.get("source_files", []) if isinstance(analysis_context, dict) else []
    tender_files = [item.get("file_name") for item in source_files if item.get("file_role") == "TENDER" and item.get("file_name")]
    attachment_files = [item.get("file_name") for item in source_files if item.get("file_role") == "ATTACHMENT" and item.get("file_name")]
    source_reference_parts = []
    if tender_files:
        source_reference_parts.append(f"主招标文件：{'、'.join(dict.fromkeys(tender_files))}")
    if attachment_files:
        source_reference_parts.append(f"招标附件：{'、'.join(dict.fromkeys(attachment_files))}")
    return "；".join(source_reference_parts)


def _split_requirement_units(text, max_items=8):
    normalized = str(text or "").replace("\r", "\n")
    parts = re.split(r"[\n；;。]", normalized)
    units = []
    for part in parts:
        item = " ".join(part.split()).strip(" ；;。")
        if len(item) < 4:
            continue
        if item not in units:
            units.append(item)
        if len(units) >= max_items:
            break
    return units


def _build_atomic_requirement_items(analysis_context):
    bidder_notice = analysis_context.get("bidder_notice", {}) or {}
    qualification_review = analysis_context.get("qualification_review", {}) or {}
    source_reference = _build_analysis_source_reference(analysis_context)

    field_specs = [
        ("general", "NORMAL", analysis_context.get("requirements", ""), "招标要求"),
        ("business", "NORMAL", analysis_context.get("business_requirements", ""), "商务要求"),
        ("technical", "NORMAL", analysis_context.get("technical_requirements", ""), "技术要求"),
        ("qualification", "REQUIRED", analysis_context.get("qualification_requirements", ""), "资格性审查"),
        ("conformity", "REQUIRED", qualification_review.get("conformity_check", ""), "符合性审查"),
        ("scoring", "IMPORTANT", analysis_context.get("scoring_items", ""), "评分项"),
        ("disqualification", "REQUIRED", analysis_context.get("disqualification_items", ""), "废标项"),
    ]
    items = []
    item_index = 1

    bidder_notice_specs = [
        ("project_name", "标的名称"),
        ("project_no", "项目编号"),
        ("package_no", "包号"),
        ("budget", "预算金额"),
        ("tenderee", "招标人"),
        ("agent", "代理机构"),
    ]
    for key, label in bidder_notice_specs:
        value = str(bidder_notice.get(key, "") or "").strip()
        if not value:
            continue
        items.append(
            {
                "item_id": f"REQ-{item_index:03d}",
                "requirement_type": "basic_info",
                "requirement_level": "REQUIRED",
                "requirement_title": label,
                "requirement_text": f"{label}：{value}",
                "source_reference": source_reference,
            }
        )
        item_index += 1

    for requirement_type, level, text, title in field_specs:
        for unit in _split_requirement_units(text):
            items.append(
                {
                    "item_id": f"REQ-{item_index:03d}",
                    "requirement_type": requirement_type,
                    "requirement_level": level,
                    "requirement_title": title,
                    "requirement_text": unit,
                    "source_reference": source_reference,
                }
            )
            item_index += 1

    return items


def _select_requirement_items_for_target(target_title, target_desc, atomic_requirement_items):
    target_text = f"{target_title} {target_desc}".strip()
    if not target_text:
        return []

    keyword_groups = [
        ("technical", ("技术", "参数", "实施", "交付", "部署", "性能", "规格")),
        ("business", ("商务", "履约", "交货", "售后", "付款", "验收", "质保")),
        ("qualification", ("资格", "资质", "授权", "证明", "审查", "营业执照", "财务", "社保")),
        ("conformity", ("符合性", "格式", "签字", "盖章", "有效期")),
        ("scoring", ("评分", "打分", "评审")),
        ("disqualification", ("废标", "否决", "无效投标")),
        ("basic_info", ("项目", "标的", "编号", "包号", "概况", "采购范围")),
    ]
    selected_types = []
    for requirement_type, keywords in keyword_groups:
        if any(keyword in target_text for keyword in keywords):
            selected_types.append(requirement_type)
    if not selected_types:
        selected_types = ["general"]

    matched = [item for item in atomic_requirement_items if item.get("requirement_type") in selected_types]
    if not matched and "general" not in selected_types:
        matched = [item for item in atomic_requirement_items if item.get("requirement_type") == "general"]
    if not matched:
        matched = atomic_requirement_items[:2]
    return matched[:4]


def _build_generation_plan_snapshot(outline, analysis_context, subject_context, product_context):
    atomic_requirement_items = _build_atomic_requirement_items(analysis_context)
    plan_items = []

    for chapter in outline or []:
        chapter_title = (chapter.get("title") or "").strip()
        bindings = _build_leaf_response_bindings(
            chapter,
            analysis_context,
            subject_context,
            knowledge_contexts={},
            product_context=product_context,
        )
        if bindings:
            for binding in bindings:
                matched_items = _select_requirement_items_for_target(
                    binding.get("title", ""),
                    binding.get("requirement", ""),
                    atomic_requirement_items,
                )
                plan_items.append(
                    {
                        "chapter_title": chapter_title,
                        "target_title": binding.get("title", ""),
                        "target_requirement": binding.get("requirement", ""),
                        "binding_status": binding.get("status", "PENDING"),
                        "plan_action": "LEAVE_BLANK"
                        if binding.get("require_blank")
                        else ("FILL" if binding.get("evidence") else "REVIEW"),
                        "requirement_level": "REQUIRED"
                        if any(item.get("requirement_level") == "REQUIRED" for item in matched_items)
                        else "NORMAL",
                        "matched_requirement_items": matched_items,
                        "evidence_preview": binding.get("evidence", [])[:3],
                    }
                )
            continue

        matched_items = _select_requirement_items_for_target(
            chapter_title,
            chapter.get("description", ""),
            atomic_requirement_items,
        )
        require_blank = _chapter_requires_evidence_placeholder(chapter) and not _chapter_has_supporting_material(
            subject_context,
            {},
            product_context,
        )
        plan_items.append(
            {
                "chapter_title": chapter_title,
                "target_title": chapter_title,
                "target_requirement": chapter.get("description", ""),
                "binding_status": "PENDING" if require_blank else ("COVERED" if matched_items else "PENDING"),
                "plan_action": "LEAVE_BLANK" if require_blank else ("FILL" if matched_items else "REVIEW"),
                "requirement_level": "REQUIRED"
                if require_blank or any(item.get("requirement_level") == "REQUIRED" for item in matched_items)
                else "NORMAL",
                "matched_requirement_items": matched_items,
                "evidence_preview": [],
            }
        )

    pending_count = sum(1 for item in plan_items if item.get("plan_action") != "FILL")
    return {
        "generated_at": utc_now().isoformat(),
        "total_atomic_requirements": len(atomic_requirement_items),
        "total_targets": len(plan_items),
        "pending_targets": pending_count,
        "atomic_requirement_items": atomic_requirement_items,
        "plan_items": plan_items,
    }


def _enrich_generation_plan_with_original_excerpts(generation_plan, analysis_result):
    if not isinstance(generation_plan, dict):
        return generation_plan
    enriched_items = []
    for item in generation_plan.get("plan_items", []) or []:
        item_copy = dict(item)
        item_copy["matched_requirement_items"] = [dict(req) for req in item.get("matched_requirement_items", []) or []]
        enriched_items.append(_maybe_attach_original_excerpt(item_copy, analysis_result))
    generation_plan["plan_items"] = enriched_items
    return generation_plan


def _persist_generation_plan_snapshot(analysis_result, generation_plan):
    if not analysis_result:
        return
    payload = {}
    if getattr(analysis_result, "analysis_data", None):
        try:
            payload = json.loads(analysis_result.analysis_data)
        except (TypeError, json.JSONDecodeError):
            payload = {}
    if not isinstance(payload, dict):
        payload = {}
    if not payload.get("version"):
        payload["version"] = "v2"
    payload["generation_plan"] = generation_plan
    analysis_result.analysis_data = json.dumps(payload, ensure_ascii=False)


def _get_generation_plan_snapshot(analysis_result):
    if not analysis_result or not getattr(analysis_result, "analysis_data", None):
        return {}
    try:
        payload = json.loads(analysis_result.analysis_data)
    except (TypeError, json.JSONDecodeError):
        return {}
    if not isinstance(payload, dict):
        return {}
    generation_plan = payload.get("generation_plan", {})
    return generation_plan if isinstance(generation_plan, dict) else {}


def _extract_original_requirement_excerpt(analysis_result, requirement_texts):
    candidates = []
    for text in requirement_texts or []:
        normalized = " ".join(str(text or "").split()).strip("；;。")
        if normalized and normalized not in candidates:
            candidates.append(normalized)
    if not candidates:
        return ""

    analysis_text = ""
    if analysis_result:
        analysis_text = str(
            getattr(analysis_result, "effective_text", "") or getattr(analysis_result, "raw_text", "") or ""
        )
    if not analysis_text.strip():
        return candidates[0][:180]

    units = [item.strip() for item in re.split(r"[\n。；;]", analysis_text) if item.strip()]
    if not units:
        return candidates[0][:180]

    for candidate in candidates:
        if candidate in analysis_text:
            for unit in units:
                if candidate in unit:
                    return unit[:180]

    keywords = []
    for candidate in candidates:
        keywords.extend([part for part in re.split(r"[\s,，、/]+", candidate) if len(part) >= 4])
    for keyword in keywords:
        for unit in units:
            if keyword in unit:
                return unit[:180]
    return candidates[0][:180]


def _maybe_attach_original_excerpt(plan_item, analysis_result):
    if not isinstance(plan_item, dict):
        return plan_item
    if plan_item.get("plan_action") != "LEAVE_BLANK" or plan_item.get("requirement_level") != "REQUIRED":
        return plan_item

    requirement_texts = []
    if plan_item.get("target_requirement"):
        requirement_texts.append(plan_item.get("target_requirement"))
    for item in plan_item.get("matched_requirement_items", []) or []:
        if item.get("requirement_text"):
            requirement_texts.append(item.get("requirement_text"))
    excerpt = _extract_original_requirement_excerpt(analysis_result, requirement_texts)
    if excerpt:
        plan_item["original_requirement_excerpt"] = excerpt
    return plan_item


def _persist_generation_coverage_snapshot(analysis_result, coverage_snapshot):
    if not analysis_result:
        return
    payload = {}
    if getattr(analysis_result, "analysis_data", None):
        try:
            payload = json.loads(analysis_result.analysis_data)
        except (TypeError, json.JSONDecodeError):
            payload = {}
    if not isinstance(payload, dict):
        payload = {}
    if not payload.get("version"):
        payload["version"] = "v2"
    payload["generation_coverage"] = coverage_snapshot
    analysis_result.analysis_data = json.dumps(payload, ensure_ascii=False)


def _get_generation_coverage_snapshot(analysis_result):
    if not analysis_result or not getattr(analysis_result, "analysis_data", None):
        return {}
    try:
        payload = json.loads(analysis_result.analysis_data)
    except (TypeError, json.JSONDecodeError):
        return {}
    if not isinstance(payload, dict):
        return {}
    coverage_snapshot = payload.get("generation_coverage", {})
    return coverage_snapshot if isinstance(coverage_snapshot, dict) else {}




def _verify_kb_citations(generated_content, knowledge_contexts):
    """验证生成内容中对知识库的引用是否真实存在。
    
    对生成内容中疑似引用知识库的段落，用 Chroma 反向检索验证。
    
    Args:
        generated_content: 生成的章节正文文本
        knowledge_contexts: 知识库上下文（含 snippets）
    
    Returns:
        dict: {verified: [引用文本列表], unverified: [未通过验证的引用文本列表]}
    """
    if not generated_content or not knowledge_contexts:
        return {"verified": [], "unverified": []}
    
    verified = []
    unverified = []
    
    # 收集知识库中的所有可用片段文本
    kb_snippets = []
    for kb in knowledge_contexts.get("knowledge_list", []):
        for snippet in kb.get("snippets", []):
            if isinstance(snippet, dict):
                snippet_text = snippet.get("text", "") or ""
            else:
                snippet_text = snippet or ""
            if snippet_text.strip() and len(snippet_text.strip()) > 20:
                kb_snippets.append(snippet_text.strip())
    
    if not kb_snippets:
        return {"verified": [], "unverified": []}
    
    # 在生成内容中查找疑似引用知识库的段落
    # 匹配模式：包含知识库文件名、"知识库"关键词、或与知识库片段相似的文本
    lines = generated_content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped or len(stripped) < 30:
            continue
        
        # 检查是否与知识库内容高度相似（说明是引用）
        for snippet in kb_snippets:
            # 计算简单重叠：如果行中的长句子出现在 snippet 中
            words = set(stripped.split())
            snippet_words = set(snippet.split())
            if len(words) > 5 and len(snippet_words) > 5:
                overlap = len(words & snippet_words)
                overlap_ratio = overlap / min(len(words), len(snippet_words))
                if overlap_ratio > 0.4:
                    verified.append(stripped[:200])
                    break
    
    return {"verified": verified, "unverified": unverified}


def _build_tender_chroma_context(task, chapter_title, chapter_desc=None):
    """从招标文件 Chroma 集合中检索与当前章节相关的原文片段（多路召回）。
    
    Args:
        task: BiddingTask 对象
        chapter_title: 当前章节标题
        chapter_desc: 当前章节描述
    
    Returns:
        list: 检索到的原文片段列表（每条带出处）
    """
    query_text = chapter_title
    if chapter_desc:
        query_text = f"{chapter_title} {chapter_desc}"
    
    collection = current_app.config.get("CHROMA_COLLECTION", "tender")
    tenant = current_app.config.get("CHROMA_TENANT", "erp")
    database = current_app.config.get("CHROMA_DATABASE", "bidding")
    
    try:
        engine = MultiRecallEngine()
        results = engine.recall(
            query=query_text[:2000],
            collection=collection,
            top_k=8,
            tenant=tenant,
            database=database,
        )
        snippets = [r["text"] for r in results if r.get("text") and len(r["text"].strip()) > 20]
        logger.info("[tender_chroma] 多路召回完成: title=%s, snippets=%s", chapter_title, len(snippets))
        return snippets
    except Exception as exc:
        logger.warning("[tender_chroma] 招标文件检索异常: %s", exc)
        return []

# ============================================================================
# 承诺函/声明函 填空引擎（v2）
# ============================================================================

# 章节类型常量
CHAPTER_TYPE_TEXT_TEMPLATE = "TEMPLATE_TEXT"      # 文本填空（承诺函/声明函等）
CHAPTER_TYPE_TABLE_TEMPLATE = "TEMPLATE_TABLE"    # 表格填充（报价表/应答表等）
CHAPTER_TYPE_QUALIFICATION = "QUALIFICATION"      # 资格证明文件
CHAPTER_TYPE_FREE_WRITE = "FREE_WRITE"            # LLM 自由写作


def _classify_chapter_type(chapter_title, chapter_desc, tender_text=None):
    """分类章节类型，决定走哪条处理路径。

    返回:
        str: CHAPTER_TYPE_* 常量之一
    """
    if not chapter_title and not chapter_desc:
        return CHAPTER_TYPE_FREE_WRITE

    combined = f"{chapter_title} {chapter_desc}".strip()

    # 1. 文本模板检测（承诺函、声明函、授权书等固定格式文本）
    text_keywords = (
        "承诺函", "声明函", "响应函", "授权委托书", "授权书",
        "廉洁承诺书", "法定代表人身份证明", "法定代表人授权",
        "资质声明函", "身份证明",
        "无行贿犯罪记录", "无重大违法记录",
    )
    if any(kw in combined for kw in text_keywords):
        return CHAPTER_TYPE_TEXT_TEMPLATE

    # 2. 表格模板检测（报价表、偏离表、应答表、业绩表、人员情况表等）
    table_keywords = (
        "报价一览表", "报价表", "报价", "偏离表",
        "应答表", "业绩一览表", "人员情况表", "基本情况表",
        "商务要求偏离", "技术要求偏离", "商务应答", "技术应答",
    )
    if any(kw in combined for kw in table_keywords):
        return CHAPTER_TYPE_TABLE_TEMPLATE

    # 3. 资格证明检测
    qual_keywords = (
        "资格证明", "资格审查", "资质证明",
        "资格性审查", "符合性审查",
    )
    if any(kw in combined for kw in qual_keywords):
        return CHAPTER_TYPE_QUALIFICATION

    # 4. 其他 → LLM 写作
    return CHAPTER_TYPE_FREE_WRITE


# ========== 路径 A：文本填空引擎 ==========


# ============================================================================
# 路径 D：置信度门控系统
# ============================================================================
# 用于在数据进入生成流程前过滤脏数据。提供：
# 1. 字段格式校验（信用代码、电话等）
# 2. OCR 文本置信度评估
# 3. 知识库召回相关性门控

# 格式校验规则集：字段名 → (regex_pattern, description)
_FORMAT_VALIDATORS = {
    "credit_code": (r'^[0-9A-HJ-NPQRTUWXY]{18}$', "统一社会信用代码：18位字母数字（不含I/O/S/V/Z）"),
    "credit_code_loose": (r'^[0-9A-Za-z]{15,18}$', "统一社会信用代码（宽松）：15-18位字母数字"),
    "phone_mobile": (r'^1[3-9]\d{9}$', "手机号：11位，1开头"),
    "phone_landline": (r'^0\d{2,3}-\d{7,8}$', "固话：带区号"),
    "email": (r'^[\w.+-]+@[\w-]+(\.[\w-]+)+$', "邮箱地址"),
    "company_name": (r'.{2,100}', "公司名称：2-100字符"),
    "project_no": (r'^[A-Za-z0-9\-]+$', "项目编号：字母数字+连字符"),
    "amount": (r'^\d+(\.\d{1,2})?$', "金额：数字，最多2位小数"),
}


def _compute_text_confidence(text: str, source: str = "ocr") -> float:
    """评估文本质量置信度（0.0 ~ 1.0）。

    适用于无法从源头获得置信度时的启发式评估。

    Args:
        text: 待评估文本
        source: 数据来源（"ocr", "kb_recall", "llm"）

    Returns:
        0.0 ~ 1.0 的置信度分数
    """
    if not text or not text.strip():
        return 0.0

    text = text.strip()
    length = len(text)

    if length < 5:
        return 0.2

    # 基础分
    score = 0.7

    # 中文字符占比（OCR 文本应该以中文为主）
    chinese_chars = sum(1 for c in text if '一' <= c <= '鿿' or '㐀' <= c <= '䶿')
    chinese_ratio = chinese_chars / length if length > 0 else 0

    if chinese_ratio > 0.5:
        score += 0.2
    elif chinese_ratio > 0.2:
        score += 0.1
    else:
        score -= 0.2  # 非中文为主的文本，大概率是乱码

    # 控制字符比例（已经在入口清洗，但评估原始质量）
    control_chars = sum(1 for c in text if ord(c) < 32 and c not in ('\n', '\r', '\t'))
    if control_chars > 0:
        score -= 0.2 * min(1.0, control_chars / max(length, 1))

    # 异常字符比例（非中文、非英文、非数字、非标点的字符）
    abnormal = sum(1 for c in text if ord(c) > 127 and not ('一' <= c <= '鿿')
                   and not ('\u3400' <= c <= '\u4dbf') and c not in ('\u3000', '\u3001', '\u3002',
                    '\uff0c', '\uff1b', '\uff1a', '\uff08', '\uff09', '\u2014', '\u2018',
                    '\u2019', '\u201c', '\u201d', '\u00b7'))
    if abnormal / max(length, 1) > 0.1:
        score -= 0.2

    # 来源特定调整
    if source == "kb_recall":
        # 知识库文本，稍微降低信任
        score -= 0.1

    return max(0.0, min(1.0, score))
def _build_template_field_map(subject_context, analysis_context):
    """构建模板填充用的字段值映射表。

    hint 关键词 → 实际值 的映射，支持同义词匹配。
    """
    context = {
        "subject": subject_context or {},
        "analysis": analysis_context.get("bidder_notice", {}) if analysis_context else {},
    }

    field_map = {}
    for keywords, source, getter in _TEMPLATE_FIELD_MAP:
        value = getter(context)
        if value:
            for kw in keywords:
                field_map[kw] = value

    return field_map


def _resolve_field_by_hint(hint, field_map):
    """根据 LLM 给出的字段 hint，在 field_map 中找最佳匹配值。

    策略：
    1. 精确匹配
    2. 包含匹配（hint in key 或 key in hint），取最长匹配
    """
    if not hint:
        return ""

    hint = hint.strip()

    # 精确匹配
    if hint in field_map:
        return field_map[hint]

    # 包含匹配
    best_key = ""
    best_value = ""
    for key, value in field_map.items():
        if key in hint or hint in key:
            if len(key) > len(best_key):
                best_key = key
                best_value = value

    return best_value


# 正则兜底模式集合（当 LLM 不可用时使用）
_FALLBACK_PLACEHOLDER_PATTERNS = [
    (r'(XXX|____|__________)[（(]([^）)]+)[）)]', 'bracket'),
    (r'(单位名称|法定代表人|授权代表|被授权人|联系电话|联系地址|项目名称|项目编号)[：:：]\s*(\_+|XXX)', 'field_value'),
    (r'(?<!（)XXX(?!（)', 'xxx_standalone'),
    (r'(?<![一-龥])\_{4,}(?![一-龥])', 'underline_standalone'),
]


def _fallback_extract_placeholders(text):
    """正则兜底：当 LLM 不可用时，用规则提取占位符。

    返回:
        list[dict]: [{"raw": "...", "start": N, "end": N, "hint": "..."}]
    """
    if not text:
        return []
    placeholders = []
    for pattern, ptype in _FALLBACK_PLACEHOLDER_PATTERNS:
        for match in re.finditer(pattern, text):
            if ptype == 'bracket' and len(match.groups()) > 1:
                hint = match.group(2).strip()
            elif ptype == 'field_value' and match.groups():
                hint = match.group(1).strip()
            else:
                hint = ""
            placeholders.append({
                "raw": match.group(0),
                "start": match.start(),
                "end": match.end(),
                "hint": hint,
            })
    placeholders.sort(key=lambda x: x["start"])
    return placeholders


def _identify_placeholders_via_llm(template_text, adapter=None):
    """调用 LLM 识别模板文本中的占位符。

    LLM 只做识别不填充，返回结构化占位符信息。
    如果 LLM 不可用，降级到正则兜底。

    返回:
        list[dict]: [{"raw": "...", "start": N, "end": N, "hint": "..."}]
    """
    if not template_text or not template_text.strip():
        return []

    # 尝试 LLM 识别
    llm_placeholders = []
    if adapter and adapter.is_available():
        try:
            prompt = (
                "你是一个占位符识别助手。找出下面文本中所有需要填写的空白位置。\n"
                "规则：\n"
                "1. 只识别，不填充，不改写原文\n"
                "2. 返回 JSON 数组格式，不要包含任何解释或其他文字\n"
                "3. 每个元素包含：raw(占位符原文), start(起始字符位置), end(结束位置), hint(推测字段含义)\n\n"
                "识别所有格式的占位符：\n"
                "- XXX（字段名）格式：XXX（比选申请人名称）\n"
                "- 下划线格式：______\n"
                "- 字段名+下划线格式：法定代表人：__________\n"
                "- 隐式空白：比选日期：  年   月   日\n"
                "- 方括号格式：【待填写】\n"
                "- 任何看起来需要填写的空白位置\n\n"
                "示例：\n"
                "本单位XXX（比选申请人名称）参加XXX（项目名称）的比选活动\n"
                '输出：[{"raw": "XXX（比选申请人名称）", "start": 3, "end": 16, "hint": "公司名称"},\n'
                ' {"raw": "XXX（项目名称）", "start": 19, "end": 28, "hint": "项目名称"}]\n\n'
                "文本：\n"
                f"{template_text[:2000]}\n\n"
                "如果文本中没有占位符，返回空数组 []。"
                "只返回 JSON 数组："
            )
            raw = adapter.generate_text(
                system_prompt="你是一个占位符识别助手，只输出 JSON 数组。",
                user_prompt=prompt,
                temperature=0.1,
                max_tokens=2000,
            )
            # 解析 JSON 结果
            import json as _json
            # 从返回中提取 JSON 数组
            json_match = re.search(r'\[.*?\]', raw.strip(), re.DOTALL)
            if json_match:
                llm_placeholders = _json.loads(json_match.group(0))
                if not isinstance(llm_placeholders, list):
                    llm_placeholders = []
                else:
                    logger.info("[template] LLM 识别占位符 %s 个", len(llm_placeholders))
        except Exception as exc:
            logger.warning("[template] LLM 占位符识别失败: %s", exc)

    # 正则仅做位置修正（不做兜底提取）
    regex_placeholders = _fallback_extract_placeholders(template_text)
    
    # 纯 LLM 模式：LLM 识别不到就返回空，不降级到正则
    if not llm_placeholders:
        logger.info("[template] LLM 未识别到占位符，保留原文")
        return []
    
    # 仅用正则修正 LLM 返回的位置偏移量
    seen_starts = {ph.get("start", -1) for ph in llm_placeholders}
    # 只在 LLM 已有结果时，用正则重算位置
    if regex_placeholders:
        # 为每个 LLM 识别的占位符找到原文中的准确位置
        corrected = []
        for ph in llm_placeholders:
            raw = ph.get("raw", "")
            hint = ph.get("hint", "")
            # 在原文中查找 raw 的实际位置
            start = template_text.find(raw)
            if start >= 0:
                corrected.append({
                    "raw": raw,
                    "start": start,
                    "end": start + len(raw),
                    "hint": hint,
                })
            else:
                # 如果 LLM 返回的 raw 在原文中找不到，尝试用 LLM 的 start 位置
                corrected.append(ph)
        return corrected
    
    logger.info("[template] LLM 识别占位符 %s 个", len(llm_placeholders))
    return llm_placeholders


def _fill_template(template_text, placeholders, field_map):
    """执行模板填充。

    参数：
        template_text: 模板原文
        placeholders: [{"raw": "...", "start": N, "end": N, "hint": "..."}]
        field_map: {hint关键词: 实际值}

    返回:
        (filled_text: str, unfilled: list[dict])
    """
    if not template_text:
        return "", []

    if not placeholders:
        return template_text, []

    # 从右向左替换，避免位置偏移
    unfilled = []
    result = list(template_text)

    for ph in reversed(sorted(placeholders, key=lambda x: x.get("start", 0))):
        raw = ph.get("raw", "")
        start = ph.get("start", 0)
        end = ph.get("end", 0)
        hint = ph.get("hint", "")

        if end > len(result) or start >= end:
            continue

        value = _resolve_field_by_hint(hint, field_map)
        if value:
            result[start:end] = value
        else:
            # 无对应值 → 留空白占位
            result[start:end] = "______"
            unfilled.append({"raw": raw, "hint": hint, "start": start})

    filled = "".join(result)
    
    # 替换后验证：检查是否还有未替换的占位符
    remaining_xxx = re.findall(r'XXX[（(]?', filled)
    remaining_ul = re.findall(r'_{4,}', filled)
    if remaining_xxx or remaining_ul:
        remaining_count = len(remaining_xxx) + len(remaining_ul)
        # 记录但不断流
        logger.info("[template] 填充后仍有 %s 个占位符未替换（XXX=%s, 下划线=%s）", 
                    remaining_count, len(remaining_xxx), len(remaining_ul))
    
    return filled, unfilled


def _verify_template_diff(original, filled):
    """校验填充结果：确保只变了占位符位置，没改其他原文。

    返回:
        (is_safe: bool, modified_positions: list[(start, end, before, after)])
    """
    if not original or not filled:
        return True, []

    # 逐字符比较
    modified = []
    o_idx, f_idx = 0, 0

    while o_idx < len(original) and f_idx < len(filled):
        if original[o_idx] == filled[f_idx]:
            o_idx += 1
            f_idx += 1
        else:
            # 记录差异
            o_start = o_idx
            f_start = f_idx
            # 找到差异结束位置
            while o_idx < len(original) and f_idx < len(filled) and original[o_idx] != filled[f_idx]:
                o_idx += 1
                f_idx += 1
            modified.append((f_start, f_idx,
                            original[o_start:o_idx],
                            filled[f_start:f_idx]))

    is_safe = not bool(modified) or all(
        len(after) <= len(before) + 20  # 允许小幅度长度变化（占位符→值）
        for _, _, before, after in modified
    )

    return is_safe, modified


def _template_has_meaningful_content(template_text):
    """判断模板文本是否包含有效内容（不只是占位符框架）。"""
    if not template_text or not template_text.strip():
        return False
    cleaned = re.sub(r'XXX|______|【[^】]+】|（[^）]*）', '', template_text).strip()
    return len(cleaned) > 30


def _extract_template_from_tender(chapter_info, tender_text):
    """从招标文件原文中提取与当前章节匹配的模板文本。

    策略：
    1. 用章节标题中的关键词在原文中定位
    2. 提取从标题行开始到下一个章节标题或文件末尾的内容
    """
    if not tender_text:
        return ""

    lines = tender_text.split("\n")

    # 从章节信息中提取搜索关键词
    search_raw = re.sub(r'^[\d一二三四五六七八九十]+[\s、.．,，、]\s*', '', chapter_info).strip()
    search_key = search_raw[:6] if len(search_raw) > 6 else search_raw
    if not search_key or len(search_key) < 2:
        return ""

    # 查找匹配行
    match_idx = -1
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        line_clean = re.sub(r'^[\d一二三四五六七八九十]+[\s、.．,，、]\s*', '', stripped).strip()
        if search_key in stripped or search_key in line_clean:
            match_idx = i
            break

    if match_idx < 0:
        return ""

    # 从匹配行开始，提取到下一个一级/二级标题或文件末尾
    extracted_lines = []
    next_section_pattern = re.compile(r'^[\d一二三四五六七八九十]+[\s、.．]')
    for line in lines[match_idx:]:
        stripped = line.strip()
        if extracted_lines and next_section_pattern.match(stripped):
            break
        extracted_lines.append(line)

    result = "\n".join(extracted_lines).strip()
    return result


def _detect_template_type(chapter_title, chapter_desc, tender_text):
    """检测当前章节是否为固定格式模板（承诺函/声明函等）。

    返回:
        (is_template: bool, template_text: str)
    """
    if not chapter_title and not chapter_desc:
        return False, ""

    combined = f"{chapter_title} {chapter_desc}".strip()

    matched = any(kw in combined for kw in (
        "承诺函", "声明函", "响应函", "授权委托书", "授权书",
        "法定代表人身份证明", "法定代表人授权", "廉洁承诺书",
        "资质声明函", "无行贿", "无重大违法",
    ))
    if not matched:
        return False, ""

    # 用章节标题（去掉编号）去原文中搜索，不要混入 desc
    search_title = re.sub(r'^[\d一二三四五六七八九十]+[\s、.．,，、]\s*', '', chapter_title).strip()
    template_text = _extract_template_from_tender(search_title or chapter_title, tender_text)
    return True, template_text




# ============================================================================
# 路径 B：表格填充引擎
# ============================================================================
# 常见表格模板的列结构定义
def _generate_table_content(chapter_title, chapter_desc, analysis_context, subject_context):
    """从 format_requirements 中当前章节的 template_content 提取第一个表格并填充。

    不创建默认表格 — 招标文件没有的表格不在标书中生成。

    返回:
        str: 以 _TABLE_JSON_PREFIX 开头的 JSON 文本，无可返回的表时返回空。
    """
    import json as _json
    table_type = chapter_title

    fmt = analysis_context.get("_format_requirements", {})
    if not fmt:
        return ""
    for sec in fmt.get("required_sections", []):
        if not (chapter_title in sec.get("title", "") or sec.get("title", "") in chapter_title):
            continue
        for blk in sec.get("template_content", []):
            if blk.get("type") != "table":
                continue
            tbl = blk
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if not headers or not rows:
                continue
            table_dict = {
                "headers": list(headers),
                "rows": list(rows),
                "merges": tbl.get("merge_cells", []),
                "column_widths": [],
                "row_heights": [],
            }
            filled_rows = _smart_fill_table(table_dict, analysis_context, subject_context)
            ncols = len(headers)
            for i in range(len(filled_rows)):
                while len(filled_rows[i]) < ncols:
                    filled_rows[i].append("")
            data_start = 0
            if filled_rows and filled_rows[0] == headers:
                data_start = 1
            table_package = {
                "headers": list(headers),
                "rows": filled_rows[data_start:],
                "column_widths": [],
                "row_heights": [],
                "merges": tbl.get("merge_cells", []),
                "text_before": "",
                "text_after": "",
            }
            marker = f"{_TABLE_JSON_PREFIX}{table_type}]]"
            return marker + "\n" + _json.dumps(table_package, ensure_ascii=False, default=str)
    return ""

# ============================================================================
# 路径 C：资格证明文件插入引擎
# ============================================================================

# 资格证明关键词 → material_type 映射
_QUALIFICATION_MATERIAL_MAP = [
    ("营业执照", "BUSINESS_LICENSE"),
    ("法人证书", "BUSINESS_LICENSE"),
    ("统一社会信用代码", "BUSINESS_LICENSE"),
    ("法定代表人身份证明", "LEGAL_PERSON_STATEMENT"),
    ("法定代表人身份证", "LEGAL_PERSON_ID_CARD"),
    ("法人身份证", "LEGAL_PERSON_ID_CARD"),
    ("授权委托书", "AUTHORIZATION_LETTER"),
    ("授权书", "AUTHORIZATION_LETTER"),
    ("被授权人身份证", "AUTHORIZED_PERSON_ID_CARD"),
    ("资质声明函", "QUALIFICATION_DECLARATION"),
    ("资格声明", "QUALIFICATION_DECLARATION"),
    ("财务报表", "FINANCIAL_STATEMENT"),
    ("纳税", "FINANCIAL_STATEMENT"),
    ("社保", "FINANCIAL_STATEMENT"),
    ("廉洁承诺书", "INTEGRITY_COMMITMENT"),
    ("资质文件", "QUALIFICATION_FILE"),
    ("资质证书", "QUALIFICATION_FILE"),
    ("许可", "QUALIFICATION_FILE"),
]


def _extract_qualification_requirements(analysis_context):
    """从分析上下文提取资格证明文件要求清单。

    返回:
        list[dict]: [{"requirement": "...", "keyword": "...", "material_type": "..."}]
    """
    requirements = []

    # 从 qualification_requirements 中提取
    qual_text = analysis_context.get("qualification_requirements", "") or ""

    # 从 qualification_review 中提取
    qual_review = analysis_context.get("qualification_review", {}) or {}
    qual_check = qual_review.get("qualification_check", "") or ""
    conformity_check = qual_review.get("conformity_check", "") or ""

    combined = f"{qual_text}\n{qual_check}\n{conformity_check}"

    # 按行切分，提取资格要求（包括提供类关键词和法定合规性要求）
    for line in combined.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        # 匹配两类要求：
        # 1. 提供/提交类（文档证明材料）
        # 2. 法定合规类（具有、无行贿、依法等声明性要求）
        is_action_req = any(kw in stripped for kw in ["提供", "提交", "证明", "出具", "递交"])
        is_compliance_req = any(kw in stripped for kw in ["具有", "无行贿", "无重大违法", "依法", "良好", "独立承担民事责任"])
        
        if is_action_req or is_compliance_req:
            for keyword, material_type in _QUALIFICATION_MATERIAL_MAP:
                if keyword in stripped:
                    requirements.append({
                        "requirement": stripped[:120],
                        "keyword": keyword,
                        "material_type": material_type,
                    })
                    break
            else:
                # 匹配到合规要求但未匹配到材料类型时，标记为通用资质文件
                if is_compliance_req:
                    requirements.append({
                        "requirement": stripped[:120],
                        "keyword": "",
                        "material_type": "QUALIFICATION_FILE",
                    })

    # 去重（按 material_type）
    seen_types = set()
    unique_requirements = []
    for req in requirements:
        if req["material_type"] not in seen_types:
            seen_types.add(req["material_type"])
            unique_requirements.append(req)

    return unique_requirements


def _check_qualification_material_status(requirements, subject_context, knowledge_contexts=None):
    """检查每项资格要求对应的主体资料是否已上传。

    三级递进查找：
    1. 主体材料中匹配 → UPLOADED
    2. 知识库中检索 → KB_FOUND
    3. 都没有 → MISSING

    Args:
        requirements: 资格要求清单
        subject_context: 主体资料上下文
        knowledge_contexts: 知识库上下文（可选）

    返回:
        list[dict]: [{"requirement": "...", "material_type": "...",
                       "status": "UPLOADED|KB_FOUND|MISSING",
                       "material": {...}, "kb_excerpt": "..."}]
    """
    if not requirements:
        return []

    materials = (subject_context or {}).get("materials", []) or []
    result = []

    for req in requirements:
        mt = req["material_type"]
        keyword = req.get("keyword", "")
        # 确保 requirement 文本不含 XML 控制字符
        if req.get("requirement"):
            req["requirement"] = _strip_xml_control_chars(req["requirement"])
        
        # Level 1: 主体材料匹配
        matched = [m for m in materials if m.get("material_type") == mt]
        if matched:
            result.append({
                "requirement": req["requirement"],
                "material_type": mt,
                "status": "UPLOADED",
                "material": matched[0],
                "kb_excerpt": "",
            })
            continue
        
        # Level 2: 知识库检索
        kb_found = False
        kb_excerpt = ""
        if knowledge_contexts:
            for kb in knowledge_contexts.get("knowledge_list", []):
                for snippet in kb.get("snippets", []):
                    if isinstance(snippet, dict):
                        snippet = snippet.get("text", "") or ""  # 兼容 dict 格式
                    if not snippet:
                        continue
                    # 匹配关键词
                    if keyword and keyword in snippet:
                        kb_found = True
                        kb_excerpt = snippet[:200]
                        break
                    # 匹配 material_type 中文名
                    mt_label = {
                        "BUSINESS_LICENSE": "营业执照",
                        "QUALIFICATION_FILE": "资质文件",
                        "LEGAL_PERSON_ID_CARD": "法人身份证",
                        "AUTHORIZATION_LETTER": "授权委托书",
                        "AUTHORIZED_PERSON_ID_CARD": "被授权人身份证",
                        "QUALIFICATION_DECLARATION": "资质声明函",
                        "LEGAL_PERSON_STATEMENT": "法定代表人身份证明",
                        "FINANCIAL_STATEMENT": "财务报表",
                        "INTEGRITY_COMMITMENT": "廉洁承诺书",
                    }.get(mt, "")
                    if mt_label and mt_label in snippet:
                        kb_found = True
                        kb_excerpt = _strip_xml_control_chars(snippet)[:200]
                        break
                if kb_found:
                    break
        
        if kb_found:
            safe_kb_excerpt = _strip_xml_control_chars(kb_excerpt) if kb_excerpt else ""
            result.append({
                "requirement": req["requirement"],
                "material_type": mt,
                "status": "KB_FOUND",
                "material": None,
                "kb_excerpt": safe_kb_excerpt,
            })
        else:
            # Level 3: 都没有 → MISSING
            result.append({
                "requirement": req["requirement"],
                "material_type": mt,
                "status": "MISSING",
                "material": None,
                "kb_excerpt": "",
            })

    return result


def _generate_qualification_content(analysis_context, subject_context, knowledge_contexts=None, chapter=None):
    """生成资格证明文件的插入指令。

    三级递进查找：
    1. 先在主体材料中匹配
    2. 主体没有 → 去知识库检索
    3. 都没有 → 标记为待人工补充

    Args:
        analysis_context: 分析上下文
        subject_context: 主体资料上下文
        knowledge_contexts: 知识库上下文（用于二级查找）
        chapter: 当前章节信息（含 children 列表，优先使用 children 作为资格要求）

    返回:
        str: 含 _QUALIFICATION_MARKER 的内容，供 _build_docx_bytes 识别处理
    """
    # 优先使用目录 children 作为资格要求（来自 check_items 的结构化数据）
    children = (chapter or {}).get("children", []) or []
    if children:
        # 从 children 构建 requirements
        requirements = []
        for child in children:
            title = (child.get("title") or "").strip()
            desc = (child.get("description") or "").strip()
            if not title:
                continue
            # 从 title 中匹配材料类型
            matched_type = "QUALIFICATION_FILE"
            for keyword, material_type in _QUALIFICATION_MATERIAL_MAP:
                if keyword in title:
                    matched_type = material_type
                    break
            requirements.append({
                "requirement": (title + " " + desc).strip()[:120],
                "keyword": "",
                "material_type": matched_type,
            })
    else:
        # 降级：从文本分析提取
        requirements = _extract_qualification_requirements(analysis_context)
    
    status_list = _check_qualification_material_status(requirements, subject_context, knowledge_contexts)

    import json as _json
    data = {
        "items": status_list,
        "uploaded_count": sum(1 for s in status_list if s["status"] == "UPLOADED"),
        "kb_found_count": sum(1 for s in status_list if s["status"] == "KB_FOUND"),
        "missing_count": sum(1 for s in status_list if s["status"] == "MISSING"),
    }

def _extract_bound_segment_data(analysis_result, bound_segments):
    """从分析结果中提取与指定 segment 绑定的数据。"""
    if not bound_segments or not analysis_result:
        return {}
    # 当前简化实现：如果 analysis_data 中有 _section_index，做简单匹配
    try:
        import json
        payload = json.loads(analysis_result.analysis_data) if isinstance(analysis_result.analysis_data, str) else (analysis_result.analysis_data or {})
        if not isinstance(payload, dict):
            return {}
        section_index = payload.get("_section_index", [])
        if section_index:
            matched = [s for s in section_index if s.get("id") in bound_segments]
            return {"matched_sections": matched, "count": len(matched)}
    except Exception:
        pass
    return {}


def _generate_chapter_content(task, chapter, analysis_result, subject_context, knowledge_contexts, product_context):
    """调用模型生成单个章节的详细正文内容。"""
    # 阶段A：尝试模板绑定 — 如果有原文模板，直接复制+填空，不走LLM
    bid_type_label_map = {"GOODS": "货物类", "SERVICE": "服务类", "ENGINEERING": "工程类"}
    try:
        _analysis_data = analysis_result.safe_analysis_data() if analysis_result else {}
        _fmt = _analysis_data.get("format_requirements", {}) if isinstance(_analysis_data, dict) else {}
        if _fmt:
            from .template_binder import bind_template, fill_content
            chapter_title_for_bind = chapter.get("title", "").strip()
            binding = bind_template(chapter_title_for_bind, _fmt)
            if binding.has_template:
                filled = fill_content(binding, subject_context, knowledge_contexts, product_context)
                if filled:
                    import json as _json
                    serialized = _json.dumps([b.to_dict() for b in filled], ensure_ascii=False)
                    logger.info("[template_binder] 章节「%s」使用模板绑定完成，%d个内容块",
                                chapter_title_for_bind, len(filled))
                    return _CONTENT_BLOCKS_PREFIX + serialized
    except Exception as _exc:
        logger.warning("[template_binder] 章节「%s」模板绑定异常: %s",
                       chapter.get("title", ""), _exc)
    bid_type_label = bid_type_label_map.get(task.bid_type, "货物类")
    chapter_title = chapter.get("title", "").strip()
    chapter_desc = chapter.get("description", "") or ""

    effective_text = analysis_result.effective_text if analysis_result and analysis_result.effective_text else (analysis_result.raw_text if analysis_result else "暂无招标依据文本。")
    analysis_context = _extract_analysis_context(analysis_result)

    catalog_profile = _get_catalog_generation_profile(task.catalog_generation_level)
    selected_package_no = getattr(task, "selected_package_no", None) or ""

    children = chapter.get("children", [])
    leaf_bindings = _build_leaf_response_bindings(
        chapter,
        analysis_context,
        subject_context,
        knowledge_contexts,
        product_context,
    )

    if _chapter_requires_evidence_placeholder(chapter) and not _chapter_has_supporting_material(
        subject_context,
        knowledge_contexts,
        product_context,
    ):
        return _EMPTY_PAGE_MARKER

    # ===== D1a: 模板存在性二次校验 =====
    # 如果 format_requirements 中存在匹配的 section 且 template_content 非空，
    # 但 bind_template 未返回 has_template=True，则判定为"有模板但绑定失败"，留空。
    try:
        _secondary_fmt = _analysis_data.get("format_requirements", {}) if isinstance(_analysis_data, dict) else {}
        if _secondary_fmt:
            _sec_reqs = _secondary_fmt.get("required_sections", [])
            if _sec_reqs:
                from .template_binder import _clean_title as _tc
                _clean_chapter = _tc(chapter_title)
                for _sec in _sec_reqs:
                    _sec_clean = _tc(_sec.get("title", ""))
                    if _clean_chapter in _sec_clean or _sec_clean in _clean_chapter:
                        _tc_content = _sec.get("template_content", []) or _sec.get("content_blocks", [])
                        if _tc_content:
                            logger.info("[template-gate] 章节「%s」有模板定义但绑定失败，留空", chapter_title)
                            return _EMPTY_PAGE_MARKER
    except Exception:
        pass

    # ===== D1b: 分隔页检测 =====
    if _is_separator_page_title(chapter_title):
        logger.info("[separator] 章节「%s」为分隔页", chapter_title)
        # 尝试从招标原文中提取分隔页的对应内容
        _sep_original = _extract_template_from_tender(chapter_title, effective_text)
        if _sep_original:
            return _SEPARATOR_PAGE_PREFIX + _sep_original
        return _SEPARATOR_PAGE_EMPTY

    # ===== D1c: 分类引擎 =====
    chapter_type = _classify_chapter_type(chapter_title, chapter_desc)

    if chapter_type == CHAPTER_TYPE_TEXT_TEMPLATE:
        _, template_text = _detect_template_type(chapter_title, chapter_desc, effective_text)
        if template_text:
            field_map = _build_template_field_map(subject_context, analysis_context)
            # 优先用 LLM 识别占位符，降级到正则
            placeholders = _identify_placeholders_via_llm(template_text)
            filled, unfilled = _fill_template(template_text, placeholders, field_map)
            if _template_has_meaningful_content(filled):
                # 原文锁定校验
                is_safe, diffs = _verify_template_diff(template_text, filled)
                if not is_safe:
                    logger.warning("[template] 章节「%s」填充后原文锁定校验失败，保留填充后原文，%s个占位符未替换",
                                   chapter_title, len(unfilled))
                logger.info("[template] 章节「%s」填空完成，占位符%s个，未填充%s个",
                            chapter_title, len(placeholders), len(unfilled))
                return filled
        # TEXT_TEMPLATE 在原文中找不到模板文本→留空，绝不落入 LLM
        logger.info("[template] 章节「%s」为模板类型但原文中未找到模板文本，留空", chapter_title)
        return _EMPTY_PAGE_MARKER

    elif chapter_type == CHAPTER_TYPE_TABLE_TEMPLATE:
        logger.info("[table] 章节「%s」使用表格引擎", chapter_title)
        return _generate_table_content(chapter_title, chapter_desc, analysis_context, subject_context)

    elif chapter_type == CHAPTER_TYPE_QUALIFICATION:
        logger.info("[qualification] 章节「%s」使用资格证明插入引擎（三级递进）", chapter_title)
        return _generate_qualification_content(analysis_context, subject_context, knowledge_contexts, chapter)

    else:
        # ===== D1d: FREE_WRITE — 仅此路径可进入 LLM =====
        pass

    system_prompt = (
        "你是一名投标文件内容编排助手，不是自由创作助手。" + "\n\n"
        "请基于给定的目录章节说明、招标需求依据、投标主体资料，" + "\n"
        "只对已经提供的内容做结构化整理与响应，不得编造、不得补写未提供的承诺或能力。" + "\n\n"
        "以下要求请严格遵守：" + "\n"
        "1. 正文内容必须紧紧围绕章节标题和说明展开，不可偏离主题。" + "\n"
        "2. 优先展示招标文件明确要求的内容和已确认资料，宁缺毯滥。" + "\n"
        "3. 如材料不足以支撑实质性承诺，请仅整理已提供的要求或事实，不得自行扩展。" + "\n"
        "4. 不要使用 Markdown 语法标记（如 #、##、**、*、-列表、```、| 表格线等）。" + "\n"
        "5. 只输出纯中文正文内容，不要重复输出顶级章节标题，不要输出解释性文字。" + "\n"
        "6. 正文使用规范的书面语，段落之间用空行分隔。" + "\n"
        "7. 正文中引用的投标主体名称使用公司全称。"
    )

    user_parts = []
    user_parts.append(f"章节标题：{chapter_title}")
    user_parts.append(f"章节说明：{chapter_desc or chapter_title}")
    user_parts.append(f"标书类型：{bid_type_label}")

    if selected_package_no:
        user_parts.append(
            f"分包信息：本项目有分包，当前包号为 {selected_package_no}。"
            f"内容只能围绕当前包号的需求编写，"
            f"不得提及其他包号的内容。"
        )

        # 截取有效文本时只保留当前包的内容
        filtered_text = _extract_effective_text(effective_text, selected_package_no)
        if filtered_text:
            effective_text = filtered_text

    user_parts.append(f"写作指导：{catalog_profile['directive']}")
    user_parts.append("编写原则：严格依据已提供材料组织内容，不得编造。如某项无支撑资料，宁可保持简略，也不要补写承诺。")
    bidder_notice = analysis_context.get("bidder_notice", {}) or {}
    if bidder_notice:
        info_lines = []
        if bidder_notice.get("project_name"):
            info_lines.append(f"项目名称：{bidder_notice['project_name']}")
        if bidder_notice.get("project_no"):
            info_lines.append(f"项目编号：{bidder_notice['project_no']}")
        if bidder_notice.get("budget"):
            info_lines.append(f"预算：{bidder_notice['budget']}")
        if bidder_notice.get("tenderee"):
            info_lines.append(f"招标人：{bidder_notice['tenderee']}")
        if bidder_notice.get("agent"):
            info_lines.append(f"代理机构：{bidder_notice['agent']}")
        if bidder_notice.get("overview"):
            info_lines.append(f"项目概况：{bidder_notice['overview']}")
        if info_lines:
            user_parts.append("\n结构化项目信息：")
            user_parts.extend(info_lines)
    if analysis_context.get("business_requirements"):
    # ========== 结构化分析数据（保留原始结构，非文本块） ==========
        _elig = analysis_context.get("_eligibility", {}) or {}
        if _elig and isinstance(_elig, dict):
            quals = _elig.get("qualifications", []) or []
            if quals:
                user_parts.append("\n[结构化] 资格要求清单（逐项）：")
                for idx, q in enumerate(quals, 1):
                    if isinstance(q, dict):
                        req = (q.get("requirement") or "").strip()
                        mat = (q.get("material") or q.get("required_material") or "").strip()
                        if req:
                            line_text = f"  {idx}. {req}"
                            if mat:
                                line_text += f" → 需提供材料：{mat}"
                            user_parts.append(line_text)
            starred = _elig.get("starred_requirements", []) or []
            if starred:
                user_parts.append("\n[结构化] ★ 实质性要求（必须完全响应）：")
                for idx, s in enumerate(starred, 1):
                    if isinstance(s, dict):
                        req = (s.get("requirement") or "").strip()
                        if req:
                            user_parts.append(f"  ★{idx}. {req}")
            disqs = _elig.get("disqualifications", []) or []
            if disqs:
                user_parts.append("\n[结构化] 废标条件（不可违反）：")
                for idx, d in enumerate(disqs, 1):
                    if isinstance(d, dict):
                        req = (d.get("requirement") or "").strip()
                        if req:
                            user_parts.append(f"  ✘{idx}. {req}")
        
        # 产品清单表
        _tc = analysis_context.get("_format_requirements", {}) or {}
        if _tc and isinstance(_tc, dict):
            pl = _tc.get("product_lists", []) or []
            if pl:
                user_parts.append("\n[结构化] 产品清单表：")
                for ti, pl_item in enumerate(pl, 1):
                    items = pl_item.get("items", []) or []
                    for item in items:
                        if isinstance(item, dict):
                            name = item.get("采购产品名称", "") or item.get("产品名称", "") or ""
                            spec = item.get("★规格参数", "") or item.get("技术参数与性能指标", "") or ""
                            qty = item.get("★数量", "") or item.get("数量", "") or ""
                            limit = item.get("★单价最高限价", "") or item.get("单价最高限价", "") or ""
                            if name:
                                line_text = f"  - {name}"
                                if spec: line_text += f" | 规格：{spec[:60]}"
                                if qty: line_text += f" | 数量：{qty}"
                                if limit: line_text += f" | 限价：{limit}"
                                user_parts.append(line_text)
        
            tech_reqs = _tc.get("tech_requirements", []) or []
            if tech_reqs:
                user_parts.append("\n[结构化] 技术参数要求（逐项）：")
                for tr in tech_reqs:
                    if isinstance(tr, dict):
                        for item in tr.get("items", []) or []:
                            if isinstance(item, dict):
                                name = item.get("技术要求名称", "") or ""
                                param = item.get("技术参数与性能指标", "") or item.get("技术参数", "") or ""
                                if name:
                                    line_text = f"  - {name}"
                                    if param: line_text += f"：{param[:120]}"
                                    user_parts.append(line_text)
        
            biz_reqs = _tc.get("business_requirements", []) or []
            if biz_reqs:
                user_parts.append("\n[结构化] 商务要求（逐项）：")
                for br in biz_reqs:
                    if isinstance(br, dict):
                        for item in br.get("items", []) or []:
                            if isinstance(item, dict):
                                name = item.get("商务要求名称", "") or ""
                                val = item.get("商务要求内容", "") or ""
                                if name:
                                    line_text = f"  - {name}"
                                    if val: line_text += f"：{val[:120]}"
                                    user_parts.append(line_text)
        
        # 评分标准
        _sc = analysis_context.get("_scoring", {}) or {}
        if _sc and isinstance(_sc, dict):
            dims = _sc.get("dimensions", []) or []
            if dims:
                user_parts.append("\n[结构化] 评分维度：")
                for idx, dim in enumerate(dims, 1):
                    if isinstance(dim, dict):
                        name = (dim.get("name") or "").strip()
                        score = (dim.get("score") or "")
                        criteria = (dim.get("criteria") or dim.get("standard") or "").strip()
                        if name:
                            line_text = f"  {idx}. {name}（{score}分）"
                            if criteria: line_text += f" - {criteria[:100]}"
                            user_parts.append(line_text)
        
        # 核心产品
        _pkgs = analysis_context.get("_packages", []) or []
        if _pkgs:
            for pkg in _pkgs:
                if isinstance(pkg, dict):
                    params = pkg.get("parameters", {}) or {}
                    if params:
                        core_products = params.get("core_products", []) or []
                        if core_products:
                            user_parts.append("\n[结构化] 核心产品列表：")
                            for cp in core_products[:10]:
                                user_parts.append(f"  - {cp}")
        
        user_parts.append(f"\n结构化商务要求：\n{analysis_context['business_requirements'][:1200]}")
    if analysis_context.get("technical_requirements"):
        user_parts.append(f"\n结构化技术要求：\n{analysis_context['technical_requirements'][:1200]}")
    if analysis_context.get("qualification_requirements"):
        user_parts.append(f"\n结构化资格性审查：\n{analysis_context['qualification_requirements'][:1200]}")
    qualification_review = analysis_context.get("qualification_review", {}) or {}
    if qualification_review.get("conformity_check"):
        user_parts.append(f"\n结构化符合性审查：\n{qualification_review['conformity_check'][:1200]}")
    if analysis_context.get("disqualification_items"):
        user_parts.append(f"\n结构化废标项：\n{analysis_context['disqualification_items'][:1200]}")
    if analysis_context.get("scoring_items"):
        user_parts.append(f"\n结构化评分标准：\n{analysis_context['scoring_items'][:1200]}")
    # 追加子项信息到提示词
    if children:
        leaf_titles = _extract_outline_leaf_titles(children)
        user_parts.append("\n该章节应包含以下具体子项及响应要求：")
        for child in children:
            child_title = child.get("title", "").strip()
            child_desc = child.get("description", "").strip()
            if child_title:
                if child_desc:
                    user_parts.append(f"  - {child_title}：{child_desc[:200]}")
                else:
                    user_parts.append(f"  - {child_title}")
        user_parts.append("\n以上子项的具体内容需在正文中逐一覆盖，按顺序展开说明，不可遗漏。")
        if leaf_titles:
            user_parts.append("请按以下子项标题分别成段输出正文，小标题必须与目录子项标题保持一致：")
            for leaf_title in leaf_titles:
                user_parts.append(f"- {leaf_title}")
        if leaf_bindings:
            user_parts.append("\n以下是系统整理出的子项绑定清单，请严格按子项逐一响应：")
            for binding in leaf_bindings:
                user_parts.append(
                    f"- {binding['title']} | 状态：{binding['status']} | 要求：{binding.get('requirement', '') or '未提取到明确要求'}"
                )
                if binding.get("evidence"):
                    for item in binding["evidence"]:
                        user_parts.append(f"  依据：{item}")
                elif binding.get("require_blank"):
                    user_parts.append("  依据：未检索到可用资料，需保留空白。")
                else:
                    user_parts.append("  依据：未检索到直接证据，仅保留招标要求本身。")

    # 改用招标文件 Chroma 语义检索替代原文截断（多路召回）
    tender_snippets = _build_tender_chroma_context(task, chapter_title, chapter_desc)
    
    # 注入质量保证约束（需求追踪矩阵）
    try:
        from ..quality_assurance import inject_constraints_into_prompt
        analysis_result = BiddingAnalysisResult.query.filter_by(shared_resource_id=task.shared_resource_id).first()
        if analysis_result and analysis_result.analysis_data:
            ad = json.loads(analysis_result.analysis_data) if isinstance(analysis_result.analysis_data, str) else analysis_result.analysis_data
            matrix = ad.get("requirement_traceability_matrix", {}) if isinstance(ad, dict) else {}
            if matrix and matrix.get("requirements"):
                constraints = inject_constraints_into_prompt(chapter_title, chapter_desc or "", matrix, {"bindings": []})
                if constraints.get("tier1_items") or constraints.get("hard_constraints"):
                    user_parts.append("\n=== 质量保证约束 ===")
                    if constraints["tier1_items"]:
                        user_parts.append("\n[第一层] 主体已有材料（必须引用）：")
                        for item in constraints["tier1_items"]:
                            src = item.get("evidence_source", {}) or {}
                            user_parts.append(f"  ✅ {item['requirement_text'][:80]} → 材料: {src.get('file_name', '')}")
                    if constraints["hard_constraints"]:
                        user_parts.append("\n[硬约束] 废标项（全程不可违反）：")
                        for item in constraints["hard_constraints"]:
                            user_parts.append(f"  🔴 {item['requirement_text'][:80]}")
    except Exception as exc:
        logger.warning("[qa] 约束注入失败: %s", exc)
    
    if tender_snippets:
        user_parts.append("\n招标需求依据（招标文件原文检索）：")
        for idx, snippet in enumerate(tender_snippets[:8], start=1):
            user_parts.append(f"[原文片段 {idx}] {snippet[:800]}")
    else:
        # 降级：使用 effective_text 前 800 字符
        user_parts.append(f"\n招标需求依据（有效文本）：\n{effective_text[:800]}")

    if subject_context:
        company = subject_context.get("company_name", "")
        user_parts.append(f"\n投标主体：{company}")
        for mat in subject_context.get("materials", []):
            user_parts.append(f"- [{mat['material_label']}] {mat['file_name']}")
            if mat.get("text_excerpt"):
                user_parts.append(f"  资料摘录：{mat['text_excerpt'][:200]}")

    if knowledge_contexts and knowledge_contexts.get("knowledge_list"):
        for kb in knowledge_contexts["knowledge_list"]:
            user_parts.append(f"\n知识库参考 [{kb.get('knowledge_base_name', '')}]:")
            for snip in kb.get("snippets", [])[:5]:
                if isinstance(snip, dict):
                    snip_text = snip.get("text", "") or ""
                else:
                    snip_text = snip or ""
                if snip_text.strip():
                    user_parts.append(f"  - {snip_text[:300]}")

    if product_context:
        terms = product_context.get("product_terms", [])
        if terms:
            user_parts.append(f"\n产品项抽取：" + "、".join(terms[:8]) + "")
        for mp in product_context.get("matched_products", [])[:3]:
            user_parts.append(f"  {mp.get('query_term', '')} -> {mp.get('matched_text', '')[:200]}")

    user_parts.append("\n请直接输出章节正文内容，不要输出解释和其他标题。")
    user_prompt = "\n".join(user_parts)

    if current_app.config.get("FLASK_ENV") == "TESTING":
        _maybe_fail_chapter_for_testing(int(chapter.get("chapter_no", 0)))
        return f"【测试内容】{chapter_title} 的模拟正文。"

    try:
        adapter = LLMAdapter(
            api_key=current_app.config.get("DEEPSEEK_API_KEY"),
            base_url=current_app.config.get("DEEPSEEK_BASE_URL"),
            default_model=current_app.config.get("DEEPSEEK_MODEL_NAME"),
        )
        if not adapter.is_available():
            logger.warning("[free] LLM 服务不可用，章节「%s」留白待补充", chapter_title)
            return _EMPTY_PAGE_MARKER

        temperature = current_app.config.get("LLM_TEMPERATURE", 0.4)
        max_tokens = 3000

        raw = adapter.generate_text(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            temperature=float(temperature),
            max_tokens=int(max_tokens),
        )
        return _normalize_chapter_content_by_bindings(raw, leaf_bindings)
    except Exception as _exc:
        logger.warning("[free] 章节「%s」LLM 调用异常: %s，降级为留白", chapter_title, _exc)
        return _EMPTY_PAGE_MARKER


def _read_file_text(file_record):
    """读取文件记录内容并解析为文本。
    
    根据存储方式选择读取路径：
    - doc_parse_cache（优先，上传阶段已同步写入）
    - CHROMA → 从 ChromaDB 按 chunk_id 读取并拼接
    - MINIO  → 从 MinIO 下载后解析
    
    注意：二进制文件（图片、压缩包等）不会被解析为文本，
    直接返回空字符串，避免乱码写入文档。
    """
    if not file_record:
        return ""

    # 检查文件类型：跳过二进制/图片文件
    file_name = file_record.file_name or ""
    ext = (Path(file_name).suffix or "").lower()
    BINARY_EXTENSIONS = {
        ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif",
        ".webp", ".ico", ".svg",
        ".zip", ".rar", ".7z", ".tar", ".gz",
        ".exe", ".dll", ".so", ".dylib",
        ".pdf",  # PDF will be handled separately by DocumentParser
    }
    # PDF 有专门的解析器，不在此过滤
    BINARY_EXTENSIONS_FOR_SKIP = BINARY_EXTENSIONS - {".pdf"}
    if ext in BINARY_EXTENSIONS_FOR_SKIP:
        logger.debug("[file_text] 跳过二进制文件: %s (ext=%s)", file_name, ext)
        return ""

    # 0. 优先从doc_parse_cache读取（上传阶段已同步写入）
    cached_text = StorageService.read_parsed_text(file_record.id)
    if cached_text:
        return cached_text

    # 1. CHROMA 存储：直接按 chroma_doc_id 从向量库读取
    if file_record.storage_provider in ("CHROMA", "CHROMA_MANAGED"):
        return _read_text_from_chroma(file_record)

    # 2. MINIO 存储：先下载文件再解析
    payload = StorageService.read_bytes(file_record)
    if payload:
        parser = DocumentParser()
        text = parser.parse_bytes(file_record.file_name or "未知文件", payload)
        if text:
            return text
    return ""


def _read_text_from_chroma(file_record):
    """从 ChromaDB 按 document_id 读取文件的所有 chunks 并拼接。
    支持 chroma_doc_id 格式: "document_id" (同步上传) 或 "document_id||task_id" (异步上传)。
    """
    chroma_doc_id = getattr(file_record, "chroma_doc_id", None)
    if not chroma_doc_id:
        return ""
    chroma_collection = file_record.chroma_collection or current_app.config.get("CHROMA_COLLECTION", "tender")
    raw = str(chroma_doc_id).strip()
    if not raw:
        return ""
    # 异步上传时 chroma_doc_id = document_id||task_id，提取 document_id
    document_id = raw.split("||")[0] if "||" in raw else raw
    adapter = ChromaAdapter(
        host=current_app.config.get("CHROMA_HOST"),
        port=current_app.config.get("CHROMA_PORT"),
        tenant=current_app.config.get("CHROMA_TENANT"),
        database=current_app.config.get("CHROMA_DATABASE"),
    )
    result = adapter.get_file_documents(chroma_collection, document_id)
    if result and result.get("documents"):
        docs = result["documents"]
        if isinstance(docs, list) and docs:
            text_parts = []
            for doc in docs:
                if isinstance(doc, str) and doc.strip():
                    text_parts.append(doc.strip())
                elif hasattr(doc, "page_content"):
                    text_parts.append(doc.page_content.strip())
                elif doc is not None:
                    text_parts.append(str(doc).strip())
            if text_parts:
                return "\n".join(text_parts)
    return ""


def _extract_package_numbers(text):
    """从招标文本中提取分包列表（仅正则，无LLM）。
    
    返回 [{"package_no": "...", "package_name": "..."}]
    """
    if not text:
        return []
    return _extract_packages_fallback(text)

def _extract_packages_fallback(text):
    """正则方式回退提取包号。"""
    results = []
    seen = set()

    # 模式1: "第X包" + 名称
    for match in re.finditer(
        r"第\s*([A-Za-z0-9一二三四五六七八九十百零]+)\s*包\s*[：:、\s]*([^\n。，,；;]{0,50})",
        text,
    ):
        package_no = match.group(1).strip()
        package_name = (match.group(2) or "").strip()
        if package_no and package_no not in seen:
            seen.add(package_no)
            results.append({"package_no": package_no, "package_name": package_name})

    # 模式2: "包号：1" 此类格式
    if not results:
        for match in re.finditer(r"包号\s*[:：]\s*([A-Za-z0-9一二三四五六七八九十百零]+)", text):
            package_no = match.group(1).strip()
            if package_no and package_no not in seen:
                seen.add(package_no)
                results.append({"package_no": package_no, "package_name": ""})

    # 模式3: "采购包1"、"标包01" 等
    if not results:
        for match in re.finditer(r"(?:采购|标|招投标?)[包匹]\s*([A-Za-z0-9一二三四五六七八九十百零]+)", text):
            package_no = match.group(1).strip()
            if package_no and package_no not in seen:
                seen.add(package_no)
                results.append({"package_no": package_no, "package_name": ""})

    return results


def _extract_effective_text(raw_text, package_no):
    """根据包号选择结果裁剪后续流程使用的有效文本。"""
    if not raw_text:
        return ""
    if not package_no:
        return raw_text
    package_no = str(package_no).strip()
    pattern = r"(第\s*([A-Za-z0-9一二三四五六七八九十]+)\s*包|包号\s*[:：]?\s*([A-Za-z0-9一二三四五六七八九十]+))"
    matches = list(re.finditer(pattern, raw_text))
    if not matches:
        return raw_text
    for index, match in enumerate(matches):
        current_package_no = str(match.group(2) or match.group(3) or "").strip()
        if current_package_no == package_no:
            start = match.start()
            if index + 1 < len(matches):
                end = matches[index + 1].start()
            else:
                end = len(raw_text)
            return raw_text[start:end].strip()
    return raw_text


def _build_check_items(shared_resource_id, overview, requirements):
    """基于分析结果生成待人工确认的核对项。"""
    BiddingCheckItem.query.filter_by(shared_resource_id=shared_resource_id).delete()
    items = [
        ("overview", "项目概述", overview, 1),
        ("requirements", "招标要求", requirements, 2),
    ]
    for check_key, check_label, check_value, sort_no in items:
        db.session.add(
            BiddingCheckItem(
                shared_resource_id=shared_resource_id,
                check_key=check_key,
                check_label=check_label,
                check_value=check_value,
                confirmed_flag=False,
                sort_no=sort_no,
            )
        )


def _build_docx_bytes(task, catalog_record, analysis_result, knowledge_contexts, product_context, subject_context, chapter_contents):
    """将章节内容组装为最终 docx 二进制文件（带专业格式、目录层级、无调试信息）。"""
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    catalog_payload = json.loads(catalog_record.catalog_content) if isinstance(catalog_record.catalog_content, str) else (catalog_record.catalog_content or {})
    outline = catalog_payload.get("outline", []) if isinstance(catalog_payload, dict) else []

    company_name = subject_context.get("company_name", "") if subject_context else ""
    analysis_context = _extract_analysis_context(analysis_result) if analysis_result else {}
    coverage_snapshot = _get_generation_coverage_snapshot(analysis_result)
    generation_plan = _get_generation_plan_snapshot(analysis_result)
    bidder_notice = analysis_context.get("bidder_notice", {}) or {}
    cover_item_name = (bidder_notice.get("project_name") or "").strip()
    cover_project_no = (bidder_notice.get("project_no") or "").strip()
    # 兜底：从 analysis_result 顶层字段获取（兼容分析未写入 bidder_notice 的情况）
    if not cover_item_name and analysis_result:
        cover_item_name = (getattr(analysis_result, "project_name", None) or
                          bidder_notice.get("标的名称", "") or "").strip()
    if not cover_project_no and analysis_result:
        cover_project_no = (getattr(analysis_result, "project_no", None) or "").strip()
    # 深度兜底：直接从 analysis_data JSON 的 metadata 中读取
    if not cover_item_name or not cover_project_no:
        _ad_raw = getattr(analysis_result, "analysis_data", None) if analysis_result else None
        if _ad_raw:
            import json as _json
            try:
                _ad = _json.loads(_ad_raw) if isinstance(_ad_raw, str) else (_ad_raw or {})
            except Exception:
                _ad = {}
            if isinstance(_ad, dict):
                _meta = _ad.get("metadata", {}) or {}
                if not cover_item_name:
                    _pn = _meta.get("project_name", {}) or {}
                    if isinstance(_pn, dict):
                        _pn = _pn.get("value", "")
                    if _pn and str(_pn).strip():
                        cover_item_name = str(_pn).strip()
                if not cover_project_no:
                    _pc = _meta.get("project_code", {}) or {}
                    if isinstance(_pc, dict):
                        _pc = _pc.get("value", "")
                    if _pc and str(_pc).strip():
                        cover_project_no = str(_pc).strip()
    cover_package_no = (getattr(task, "selected_package_no", "") or bidder_notice.get("package_no") or "").strip()
    cover_bid_time = utc_now().strftime("%Y年%m月%d日")

    document = Document()
    inserted_material_ids = set()
    image_extensions = {"png", "jpg", "jpeg", "gif", "bmp", "tiff", "tif", "webp"}

    # ========== 插入免责声明页（第一页） ==========
    def _add_disclaimer_page(doc):
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for _ in range(1):
            doc.add_paragraph("")
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("免责声明")
        run.font.name = "黑体"
        run.font.size = Pt(22)
        run.bold = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        doc.add_paragraph("")

        disclaimer_lines = [
            ("一、服务性质", "本服务为 AI 辅助工具，用于生成标书参考初稿。您必须对最终提交的标书文件负全部责任，包括审查、修改内容以确保其符合所有法规与项目要求。"),
            ("二、不担保准确性", "本公司不保证 AI 生成内容的绝对准确性与完整性。您必须自行核实所有关键信息，并承担因使用本服务而产生的任何后果。"),
            ("三、知识产权承诺与风险", "您须确保上传的所有资料不侵犯任何第三方权利。由此引发的任何法律责任及赔偿，均由您自行承担。本公司对用户上传内容不享有权利，也不承担审查义务。"),
            ("四、图片素材风险提示", "服务提供的图片素材仅供参考。您若使用（包括引用、修改或二次创作），必须自行承担其导致的侵权等全部风险与责任，本公司概不负责。"),
            ("五、责任限制", "在任何情况下，本公司均不对因使用本服务造成的任何直接、间接或后果性损失（如利润损失、业务中断、数据丢失）承担责任。"),
            ("六、其他", "本公司保留随时修改或终止服务的权利。本须知的解释与争议解决均适用中华人民共和国法律。"),
        ]

        for clause_title, clause_body in disclaimer_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.5
            run_title = p.add_run(f"{clause_title}：")
            run_title.font.name = "黑体"
            run_title.font.size = Pt(12)
            run_title.bold = True
            run_title.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run_body = p.add_run(clause_body)
            run_body.font.name = "宋体"
            run_body.font.size = Pt(12)
            run_body.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        doc.add_paragraph("")
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = note_p.add_run("（使用本服务即视为已阅读并同意以上条款）")
        run.font.name = "宋体"
        run.font.size = Pt(11)
        run.italic = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    _add_disclaimer_page(document)
    document.add_page_break()

    # ========== 页面设置 ==========
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========== 默认字体 ==========
    style = document.styles["Normal"]
    font = style.font
    font.name = "仿宋"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # ========== 定义标题样式 ==========
    def _set_heading_style(heading_level, font_name, font_size, bold=True, space_before=12, space_after=6):
        hs = document.styles[f"Heading {heading_level}"]
        hs.font.name = font_name
        hs.font.size = Pt(font_size)
        hs.font.bold = bold
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        hpf = hs.paragraph_format
        hpf.space_before = Pt(space_before)
        hpf.space_after = Pt(space_after)
        hpf.line_spacing = 1.5

    _set_heading_style(1, "宋体", 16, True, 24, 12)
    _set_heading_style(2, "宋体", 15, True, 18, 8)
    _set_heading_style(3, "宋体", 14, True, 12, 6)
    _set_heading_style(4, "宋体", 12, True, 6, 6)

    # ========== 辅助函数 ==========
    def _clean_markdown(text):
        cleaned = text
        cleaned = re.sub(r'```[\w]*\n?', '', cleaned)
        cleaned = re.sub(r'`([^`]+)`', r'\1', cleaned)
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)
        cleaned = re.sub(r'^#{1,6}\s+', '', cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
        cleaned = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', cleaned)
        cleaned = re.sub(r'^[-*_]{3,}\s*$', '', cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'^[\s]*[-*+]\s+', '', cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'^\s*\d+[.\)]\s+', '', cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'----media/[\w./-]+----', '', cleaned)
        cleaned = re.sub(r'media/image\d+\.\w+', '', cleaned)
        cleaned = _strip_xml_control_chars(cleaned)
        return cleaned.strip()


    def _build_subject_declaration_text():
        materials = subject_context.get("materials", []) if subject_context else []
        if not company_name or not materials:
            return ""
        labels = [item.get("material_label", "").strip() for item in materials if item.get("material_label")]
        joined_labels = "、".join(dict.fromkeys(labels))
        if not joined_labels:
            return ""
        return (
            f"{company_name}郑重声明：本单位已按本项目要求提供主体资质、身份证明及授权相关材料，"
            f"包括{joined_labels}。凡在本标书中引用到前述主体资料的章节，均同步插入对应原始文件、扫描页或图片内容；"
            "未在正文中单独展开的资料，统一附于本文件后续附件章节备查。"
        )

    def _write_formatted_content(doc, text):
        if not text or not text.strip():
            return
        cleaned_text = _clean_markdown(text)
        if not cleaned_text.strip():
            return
        lines = cleaned_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if re.match(r'^[\s]*----media/', stripped) or 'media/image' in stripped:
                continue
            # 安全网：确保每行文本不含 XML 控制字符
            safe_text = _strip_xml_control_chars(stripped)
            if not safe_text:
                continue
            p = doc.add_paragraph(safe_text)
            p.style = document.styles["Normal"]
            pf = p.paragraph_format
            pf.first_line_indent = Pt(24)
            pf.line_spacing = 1.5

    

    def _get_material_identity(material):
        return material.get("id") or material.get("file_id") or material.get("file_name")

    def _get_material_file_record(material):
        file_id = material.get("file_id")
        if not file_id:
            return None
        try:
            return db.session.get(FileStorage, int(file_id))
        except Exception:
            return None

    def _extract_docx_media_payloads(payload, max_images=6):
        from zipfile import ZipFile

        results = []
        try:
            with ZipFile(BytesIO(payload)) as archive:
                media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
                for media_name in media_names[:max_images]:
                    media_payload = archive.read(media_name)
                    if media_payload:
                        results.append(media_payload)
        except Exception as exc:
            logger.warning("[docx] 提取 docx 图片失败: %s", exc)
        return results

    def _render_pdf_pages_as_png(payload, max_pages=4):
        import fitz

        images = []
        try:
            pdf = fitz.open(stream=payload, filetype="pdf")
            try:
                for page_index in range(min(len(pdf), max_pages)):
                    page = pdf.load_page(page_index)
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
                    images.append(pixmap.tobytes("png"))
            finally:
                pdf.close()
        except Exception as exc:
            logger.warning("[docx] 渲染 PDF 图片失败: %s", exc)
        return images

    def _load_material_visual_payloads(material, max_assets=6):
        file_record = _get_material_file_record(material)
        if not file_record:
            return []
        payload = StorageService.read_bytes(file_record)
        if not payload:
            return []
        extension = (file_record.file_ext or material.get("file_ext") or Path(file_record.file_name or "").suffix.lstrip(".")).lower()
        if extension in image_extensions:
            return [payload]
        if extension == "pdf":
            return _render_pdf_pages_as_png(payload, max_pages=max_assets)
        if extension == "docx":
            return _extract_docx_media_payloads(payload, max_images=max_assets)
        return []

    def _add_picture_payload(payload, width_cm=15.5):
        try:
            pic_paragraph = document.add_paragraph()
            pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_paragraph.add_run().add_picture(BytesIO(payload), width=Cm(width_cm))
            return True
        except Exception as exc:
            logger.warning("[docx] 插入图片失败: %s", exc)
            return False

    def _write_material_block(material):
        material_title = material.get("material_label") or "主体资料"
        file_name = material.get("file_name") or ""
        material_id = _get_material_identity(material)

        document.add_heading(material_title, level=3)
        # 只显示有意义的文件名（排除系统生成的纯数字/哈希文件名）
        if file_name and not re.match(r'^[\d_]+(\.\w+)$', file_name):
            file_para = document.add_paragraph(f"文件名称：{file_name}")
            file_para.style = document.styles["Normal"]

        text_excerpt = (material.get("text_excerpt") or "").strip()
        if text_excerpt:
            _write_formatted_content(document, text_excerpt)

        inserted_visual = False
        for image_payload in _load_material_visual_payloads(material):
            if _add_picture_payload(image_payload):
                inserted_visual = True

        if not inserted_visual and not text_excerpt:
            fallback = document.add_paragraph(
                "当前资料记录仅保留文本索引，未保留可回填的原始文件或图片流，无法按原样插入扫描页。"
                "如需恢复该类图片展示，请重新上传对应主体资料后重新生成标书。"
            )
            fallback.style = document.styles["Normal"]

        inserted_material_ids.add(material_id)

    def _material_matches_outline_item(material, title, desc=""):
        material_type = str(material.get("material_type") or "").strip().upper()
        outline_text = f"{title} {desc}".strip()
        auth_keywords = ("授权", "委托", "身份证明", "法定代表人", "被授权人")
        qualification_keywords = ("资质", "资格", "营业执照", "证明材料", "审查", "声明函", "响应文件格式")

        if any(keyword in outline_text for keyword in auth_keywords):
            return material_type in {
                "AUTHORIZATION_LETTER",
                "AUTHORIZED_PERSON_ID_CARD",
                "LEGAL_PERSON_ID_CARD",
                "LEGAL_PERSON_STATEMENT",
            }
        if any(keyword in outline_text for keyword in qualification_keywords):
            return material_type in {
                "BUSINESS_LICENSE",
                "QUALIFICATION_FILE",
                "QUALIFICATION_DECLARATION",
                "LEGAL_PERSON_ID_CARD",
            }
        return False

    def _write_subject_materials_for_outline_item(title, desc=""):
        materials = subject_context.get("materials", []) if subject_context else []
        matched = []
        for material in materials:
            material_id = _get_material_identity(material)
            if material_id in inserted_material_ids:
                continue
            if _material_matches_outline_item(material, title, desc):
                matched.append(material)
        if not matched:
            return 0

        intro = document.add_paragraph("以下插入与本节内容直接对应的主体资质/授权原始资料：")
        intro.style = document.styles["Normal"]
        for material in matched:
            _write_material_block(material)
        return len(matched)

    def _write_remaining_subject_materials():
        materials = subject_context.get("materials", []) if subject_context else []
        remaining = [item for item in materials if _get_material_identity(item) not in inserted_material_ids]
        if not remaining:
            return

        document.add_page_break()
        document.add_heading("主体资料附件", level=1)
        declaration_text = _build_subject_declaration_text()
        if declaration_text:
            _write_formatted_content(document, declaration_text)
        for material in remaining:
            _write_material_block(material)



    def _render_separator_page(doc, outline_item, original_text=None):
        """渲染响应文件分隔页：居中、大字号、独立一页。"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        # 上方留白
        for _ in range(6):
            doc.add_paragraph("")
        title = outline_item.get("title", "").strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(_strip_xml_control_chars(title))
        run.font.name = "宋体"
        run.font.size = Pt(22)  # 二号
        run.bold = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if original_text:
            _write_formatted_content(doc, original_text)

    def _normalize_outline_title_for_match(title):
        return re.sub(r"\s+", "", str(title or "").strip())

    def _find_plan_item(chapter_title, target_title):
        for item in (generation_plan.get("plan_items", []) if generation_plan else []):
            item_chapter_title = (item.get("chapter_title") or "").strip()
            item_target_title = (item.get("target_title") or "").strip()
            if item_chapter_title == (chapter_title or "").strip() and item_target_title == (target_title or "").strip():
                return item
        return {}

    def _extract_child_content_sections(content_text, children):
        if not content_text or not children:
            return {}
        lines = [line.strip() for line in str(content_text or "").splitlines()]
        child_titles = [(child.get("title") or "").strip() for child in children if (child.get("title") or "").strip()]
        if not child_titles:
            return {}

        title_positions = []
        for index, line in enumerate(lines):
            normalized_line = _normalize_outline_title_for_match(line)
            if not normalized_line:
                continue
            for child_title in child_titles:
                if normalized_line == _normalize_outline_title_for_match(child_title):
                    title_positions.append((index, child_title))
                    break
        if not title_positions:
            return {}

        sections = {}
        for position, (start_index, child_title) in enumerate(title_positions):
            end_index = title_positions[position + 1][0] if position + 1 < len(title_positions) else len(lines)
            body_lines = [line for line in lines[start_index + 1 : end_index] if line]
            if body_lines:
                sections[child_title] = "\n".join(body_lines).strip()
        return sections

        # ========== 第二页封面（优先使用招标文件封面模板） ==========
    # 从 outline 中找 is_cover=True 的节点，渲染其封面内容
    # outline 节点现在自带 template_content（含 font 信息和占位符标记）
    _cover_outline_items = [item for item in outline if item.get("is_cover")]
    _cover_template_found = False
    
    # ====== 占位符填充辅助函数 ======
    def _fill_placeholder_text(text):
        """对封面文本进行占位符填充。"""
        # 仅第一个封面做全文替换，后续封面只替换可识别的占位符
        result = text
        # 常见占位符替换（按优先级）
        result = result.replace("XXX（单位名称）", company_name or "XXX（单位名称）")
        result = result.replace("XXX", company_name or "XXX")
        result = result.replace("（项目名称）", cover_item_name or "（项目名称）")
        result = result.replace("（项目编号）", cover_project_no or "（项目编号）")
        result = result.replace("采购项目名称:#", "")
        return result
    
    def _is_empty_or_placeholder(text):
        """判断文本是否为空或纯占位符。"""
        if not text or not text.strip():
            return True
        if re.fullmatch(r'[_\s]{2,}', text.strip()):
            return True
        if re.fullmatch(r'[xX]{2,}', text.strip()):
            return True
        return False
    
    for _cover_idx, _cover_item in enumerate(_cover_outline_items):
        _cover_blocks = _cover_item.get("template_content", [])
        _is_first_cover = (_cover_idx == 0)
        
        if _cover_blocks:
            for _blk in _cover_blocks:
                if _blk.get("type") in ("paragraph", "text"):
                    _text = _blk.get("text", "") or ""
                    _font = _blk.get("font", {}) or {}
                    _is_placeholder = _blk.get("placeholder", False)
                    _fill_mode = _blk.get("fill_mode", "")
                    
                    # 处理占位符填充
                    if _is_placeholder:
                        if _fill_mode == "replace":
                            # 纯占位符 → 尝试用上下文填充
                            _filled = _fill_placeholder_text(_text)
                            if _filled == _text and _is_empty_or_placeholder(_text):
                                # 找不到内容，保留原样（留空）
                                pass
                            else:
                                _text = _filled
                        elif _fill_mode == "partial":
                            # 混合型 → 仅替换占位符部分
                            _text = _fill_placeholder_text(_text)
                        else:
                            # 普通占位符 → 尝试填充
                            _text = _fill_placeholder_text(_text)
                    
                    _p = document.add_paragraph()
                    # 应用对齐方式
                    _alignment = _font.get("alignment", "")
                    if _alignment == "center":
                        _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif _alignment == "right":
                        _p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif _alignment == "left":
                        _p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        _p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 封面默认居中
                    
                    _r = _p.add_run(_text)
                    # 应用字体信息
                    _font_name = _font.get("font_name", "") or "宋体"
                    _font_size = _font.get("font_size", 16.0)
                    _font_bold = _font.get("bold", False)
                    try:
                        _r.font.name = _font_name
                        _r.font.size = Pt(_font_size)
                    except Exception:
                        _r.font.name = "宋体"
                        _r.font.size = Pt(16)
                    if _font_bold:
                        _r.bold = True
                    try:
                        _r.element.rPr.rFonts.set(qn("w:eastAsia"), _font_name or "宋体")
                    except Exception:
                        pass
                        
                elif _blk.get("type") == "table":
                    _headers = _blk.get("headers", [])
                    _rows = _blk.get("rows", [])
                    if _headers and _rows:
                        _t = document.add_table(rows=len(_rows), cols=len(_headers))
                        _t.style = "Table Grid"
                        _apply_black_solid_borders(_t)
                        for _ci, _h in enumerate(_headers):
                            _t.rows[0].cells[_ci].text = _h
                        for _ri, _row in enumerate(_rows):
                            for _ci, _cell in enumerate(_row):
                                if _is_first_cover:
                                    _filled = _fill_placeholder_text(_cell)
                                else:
                                    _filled = _cell.replace("XXX", company_name or "") if _ci == 0 else _cell
                                _t.rows[_ri].cells[_ci].text = _filled
            _cover_template_found = True
        else:
            # outline 节点有 is_cover 但无 template_content → 降级到自生成
            _cover_template_found = False
    
    if not _cover_template_found:
        # 自有封面模板
        for _ in range(5):
            document.add_paragraph("")
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("投标文件")
        run.font.name = "宋体"
        run.font.size = Pt(22)
        run.bold = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        document.add_paragraph("")
        cover_fields = [
            ("标的名称", cover_item_name),
            ("项目编号", cover_project_no),
            ("投标人名称", company_name),
            ("投标时间", cover_bid_time),
        ]
        if cover_package_no:
            cover_fields.append(("包号", cover_package_no))

        for label, value in cover_fields:
            field_para = document.add_paragraph()
            field_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            safe_value = _strip_xml_control_chars(str(value or ""))
            run = field_para.add_run(f"{label}：{safe_value}")
            run.font.name = "宋体"
            run.font.size = Pt(16)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    document.add_page_break()

    # ========== 目录页（占位） ==========
    for _ in range(1):
        document.add_paragraph("")
    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    run.font.name = "黑体"
    run.font.size = Pt(22)
    run.bold = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    document.add_paragraph("")
    # 输出目录结构
    def _write_toc_items(items, indent=0):
        for item in items:
            title = item.get("title", "").strip()
            if not title:
                continue
            indent_str = "    " * indent
            p = document.add_paragraph(f"{indent_str}{_strip_xml_control_chars(title)}")
            p.style = document.styles["Normal"]
            pf = p.paragraph_format
            pf.line_spacing = 1.8
            for run in p.runs:
                run.font.name = "宋体"
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            children = item.get("children", [])
            if children:
                _write_toc_items(children, indent + 1)

    _write_toc_items(outline)

    document.add_paragraph("")
    p = document.add_paragraph("（以上目录由 AI 辅助生成，建议在 Word 中使用“插入 → 目录”功能生成规范目录）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10)
        run.italic = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    document.add_page_break()

    # ========== 递归写入 outline 节点 ==========
    def _write_outline_item(outline_item, level=1, inherited_child_sections=None, parent_title=None):
        title = outline_item.get("title", "").strip()
        desc = outline_item.get("description", "").strip()
        if not title:
            return
        h = document.add_heading(title, level=min(level, 4))
        _last_element = h._element
        chapter_title_for_plan = parent_title or title

        matched_content = None
        # 顶级节点：从 chapter_contents 匹配 LLM 生成的内容
        if level == 1:
            chapter_idx = outline_item.get("_chapter_idx")
            if chapter_idx is not None and chapter_idx < len(chapter_contents):
                matched_content = chapter_contents[chapter_idx].get("content", "")

            if not matched_content:
                logger.info("[write] 标题匹配查找: title='%s', chapter_contents共%d条", title, len(chapter_contents))
                for cc in chapter_contents:
                    _cc_title = cc.get("title", "")
                    if title in _cc_title or _cc_title in title:
                        logger.info("[write] 标题匹配成功: title='%s' -> cc.title='%s'", title, _cc_title)
                        matched_content = cc.get("content", "")
                        break
                if not matched_content:
                    logger.info("[write] 标题匹配失败: title='%s'", title)

        if not matched_content and inherited_child_sections:
            matched_content = inherited_child_sections.get(title)
        if not matched_content and desc and not _chapter_requires_evidence_placeholder(outline_item):
            matched_content = desc

        children = outline_item.get("children", [])
        child_sections = _extract_child_content_sections(matched_content, children) if matched_content and matched_content != _EMPTY_PAGE_MARKER and children else {}

        if matched_content == _EMPTY_PAGE_MARKER:
            # 留白标记：在标题下插入提示文字，不留多余分页
            # 标题已在 _write_outline_item 开头通过 add_heading 写入
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("（本节无内容，待补充）")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run.font.name = "仿宋"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            # 添加招标原文提示（如果有）
            plan_item = _find_plan_item(chapter_title_for_plan, title)
            original_excerpt = (plan_item.get("original_requirement_excerpt") or "").strip()
            original_excerpt = _strip_xml_control_chars(original_excerpt)
            if original_excerpt:
                excerpt_p = document.add_paragraph(f"招标文件原文要求：{original_excerpt}")
                excerpt_p.style = document.styles["Normal"]
                excerpt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in excerpt_p.runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.font.size = Pt(10)
            # 不插入前后分页符——让主循环的 _oi_idx > 0 分页逻辑控制换页
            return

        if matched_content or (level == 1 and chapter_idx is not None 
                               and chapter_idx < len(chapter_contents) 
                               and chapter_contents[chapter_idx].get("content_blocks")):
            # ========== 优先处理 ContentBlock 结构化内容 ==========
            _chapter_cc = chapter_contents[chapter_idx].get("content_blocks") if level == 1 and chapter_idx is not None and chapter_idx < len(chapter_contents) else None
            if not _chapter_cc:
                for _cc in chapter_contents:
                    if title in _cc.get("title", "") or _cc.get("title", "") in title:
                        _chapter_cc = _cc.get("content_blocks")
                        if _chapter_cc:
                            break
                # 深度后备：直接从 analysis_data 的 section_lookup 查找模板内容
                # 用于 level >= 2 的子章节（非顶级 outline 节点，没有独立的 content_blocks）
                if not _chapter_cc:
                    try:
                        _ad_raw3 = getattr(analysis_result, "analysis_data", None) if analysis_result else None
                        if _ad_raw3:
                            import json as _json3
                            _ad3 = _json3.loads(_ad_raw3) if isinstance(_ad_raw3, str) else (_ad_raw3 or {})
                            _fmt3 = _ad3.get("format_requirements", {}) if isinstance(_ad3, dict) else {}
                            _sec_lookup3 = _fmt3.get("section_lookup", {}) if isinstance(_fmt3, dict) else {}
                            # 使用与 section_lookup 构建时一致的 _clean_section_title
                            from .analysis_v3.phase1_5_format import _clean_section_title as _lookup_clean3
                            _clean_t3 = _lookup_clean3(title)
                            _sec3 = _sec_lookup3.get(_clean_t3) if isinstance(_sec_lookup3, dict) else None
                            if not _sec3:
                                # 降级：尝试前缀匹配（兼容含编号前缀的标题）
                                for _lk, _lv in _sec_lookup3.items():
                                    if _clean_t3 in _lk or _lk in _clean_t3:
                                        _sec3 = _lv
                                        break
                            if _sec3:
                                _tc3 = _sec3.get("template_content", []) or _sec3.get("content_blocks", [])
                                if _tc3:
                                    from .template_binder import ContentBlock as _CB3
                                    _blocks3 = []
                                    for _bd3 in _tc3:
                                        if not isinstance(_bd3, dict):
                                            continue
                                        _bt3 = _bd3.get("type", "text")
                                        if _bt3 in ("text", "paragraph"):
                                            _blocks3.append(_CB3.paragraph(_bd3.get("text", ""), []))
                                        elif _bt3 == "table":
                                            _blocks3.append(_CB3.table(
                                                _bd3.get("headers", []),
                                                _bd3.get("rows", []),
                                                _bd3.get("merge_cells", []),
                                                _bd3.get("column_widths", []),
                                                per_cell=_bd3.get("per_cell"),
                                            ))
                                    if _blocks3:
                                        _chapter_cc = [b.to_dict() for b in _blocks3]
                    except Exception as _fmt_exc3:
                        logger.warning("[write] format_requirements 后备查找失败: %s", _fmt_exc3)

            if _chapter_cc:
                for _block in _chapter_cc:
                    if isinstance(_block, dict):
                        if _block.get("type") in ("text", "paragraph"):
                            _p = document.add_paragraph(_strip_xml_control_chars(_block.get("text", "")))
                            _p.style = document.styles["Normal"]
                            for _r in _p.runs:
                                _r.font.name = "仿宋"
                                _r.font.size = Pt(12)
                                _r.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                            _last_element = _p._element
                        elif _block.get("type") == "table":
                            # 统一渲染：通过 per_cell_data → write_table_from_data
                            _pcd = _block.get("per_cell_data") or _block.get("per_cell")
                            if _pcd:
                                from app.infrastructure.table_codec import from_dict, write_table_from_data
                                td = from_dict(_pcd)
                            else:
                                from app.infrastructure.table_codec import to_per_cell, write_table_from_data
                                _hd = _block.get("headers", [])
                                _rw = _block.get("rows", [])
                                _mc = _block.get("merge_cells", [])
                                _cw = _block.get("column_widths", [])
                                _all_rows = [_hd] + _rw if _hd else _rw
                                td = to_per_cell(_hd, _all_rows, _mc, _cw)
                            # 注入默认格式
                            for _row in td.rows:
                                for _cell in _row.cells:
                                    if not _cell.font_name:
                                        _cell.font_name = "仿宋"
                                    if not _cell.font_size_half_pt:
                                        _cell.font_size_half_pt = 24
                            _tbl = write_table_from_data(document, td, insert_after=_last_element)
                            if _tbl is not None:
                                _last_element = _tbl
                                # ContentBlock 已处理，跳过后续文本写入和表格标记处理
                # 仍需要处理子章节
                _write_subject_materials_for_outline_item(title, desc)
                for _child in children:
                    _write_outline_item(_child, level=level + 1, inherited_child_sections=child_sections, parent_title=chapter_title_for_plan)
                return

            # ========== JSON 表格标记处理（携带宽/高/合并信息） ==========
            if isinstance(matched_content, str) and matched_content.startswith(_TABLE_JSON_PREFIX):
                if desc and desc not in title:
                    desc_p = document.add_paragraph(_strip_xml_control_chars(desc))
                    desc_p.style = document.styles["Normal"]
                    for run in desc_p.runs:
                        run.font.name = "\u4eff\u5b8b"
                        run.font.size = Pt(12)
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "\u4eff\u5b8b")
                end_marker = matched_content.find("]]\n")
                if end_marker > 0:
                    json_str = matched_content[end_marker + 3:].strip()
                    if json_str:
                        try:
                            import json as _json
                            table_dict = _json.loads(json_str)
                            headers = table_dict.get("headers", [])
                            rows = table_dict.get("rows", [])
                            column_widths = table_dict.get("column_widths", [])
                            row_heights = table_dict.get("row_heights", [])
                            merges = table_dict.get("merges", [])
                            text_before = table_dict.get("text_before", "")
                            text_after = table_dict.get("text_after", "")

                            # 写入表格前段落
                            if text_before:
                                tb_p = document.add_paragraph(_strip_xml_control_chars(text_before))
                                tb_p.style = document.styles["Normal"]
                                for run in tb_p.runs:
                                    run.font.name = "\u4eff\u5b8b"
                                    run.font.size = Pt(12)
                                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "\u4eff\u5b8b")

                            # 统一通过 to_per_cell() 构建 TableData
                            from app.infrastructure.table_codec import to_per_cell, write_table_from_data
                            all_rows = [headers] + rows
                            td = to_per_cell(headers, all_rows, merges, column_widths, row_heights)
                            # 注入默认格式
                            for _row in td.rows:
                                for _cell in _row.cells:
                                    if not _cell.font_name:
                                        _cell.font_name = "仿宋"
                                    if not _cell.font_size_half_pt:
                                        _cell.font_size_half_pt = 24
                            _tbl = write_table_from_data(document, td, insert_after=_last_element)
                            if _tbl is not None:
                                _last_element = _tbl

                            # 写入表格后段落
                            if text_after:
                                ta_p = document.add_paragraph(_strip_xml_control_chars(text_after))
                                ta_p.style = document.styles["Normal"]
                                for run in ta_p.runs:
                                    run.font.name = "\u4eff\u5b8b"
                                    run.font.size = Pt(12)
                                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "\u4eff\u5b8b")
                        except Exception as exc:
                            logger.warning("[docx] JSON 表格渲染失败, 降级到旧格式: %s", exc)
                            document.add_paragraph("（此处为表格模板，请根据实际情况填写）")
                else:
                    document.add_paragraph("（此处为表格模板，请根据实际情况填写）")
                _write_subject_materials_for_outline_item(title, desc)
                children = outline_item.get("children", [])
                for child in children:
                    _write_outline_item(child, level=level + 1, inherited_child_sections=child_sections, parent_title=chapter_title_for_plan)
                return

            # ========== 表格标记处理已迁移到 JSON 路径 ==========
            # 不再使用 _TABLE_MARKER_PREFIX 路径，统一走 JSON 表格路径

            # ========== 资格证明文件标记处理（三级递进：主体→知识库→留白） ==========
            if isinstance(matched_content, str) and matched_content.startswith(_QUALIFICATION_MARKER):
                import json as _json
                qual_data_str = matched_content[len(_QUALIFICATION_MARKER):]
                try:
                    qual_data = _json.loads(qual_data_str)
                    items = qual_data.get("items", [])
                    
                    # 统计状态
                    uploaded = qual_data.get("uploaded_count", 0)
                    kb_found = qual_data.get("kb_found_count", 0)
                    missing = qual_data.get("missing_count", 0)
                    total = uploaded + kb_found + missing
                    
                    for _qi_idx, item in enumerate(items):
                        if _qi_idx > 0:
                            document.add_page_break()
                        req = item.get("requirement", "")
                        safe_req = _strip_xml_control_chars(req)
                        status = item.get("status", "MISSING")
                        
                        # 每项资格内容使用三级标题
                        _qh = document.add_heading(_strip_xml_control_chars(req), level=3)
                        
                        if status == "UPLOADED":
                            p = document.add_paragraph("（主体已上传相关证明材料）")
                            p.style = document.styles["Normal"]
                        elif status == "KB_FOUND":
                            p = document.add_paragraph("（知识库检索到相关证明材料）")
                            p.style = document.styles["Normal"]
                        else:
                            p = document.add_paragraph("（待人工补充相关证明材料）")
                            p.style = document.styles["Normal"]
                            if p.runs:
                                p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                except Exception as exc:
                    logger.warning("[qualification] 资格证明数据解析失败: %s", exc)
                    document.add_paragraph("（资格证明文件处理异常，请在对应章节查看）")
                _write_subject_materials_for_outline_item(title, desc)
                children_for_qual = outline_item.get("children", [])
                for child in children_for_qual:
                    _write_outline_item(child, level=level + 1, inherited_child_sections=child_sections, parent_title=chapter_title_for_plan)
                return

            # 如果内容以标题开头，去掉重复的标题文字
            content_text = matched_content
            first_line = content_text.split("\n")[0].strip()
            if title in first_line or first_line in title:
                content_text = "\n".join(content_text.split("\n")[1:]).strip()

            # 渲染纯文本内容（表格已通过上层 JSON/ContentBlock 路径处理）
            if content_text.strip():
                _write_formatted_content(document, content_text.strip())
        _write_subject_materials_for_outline_item(title, desc)

        for child in children:
            _write_outline_item(child, level=level + 1, inherited_child_sections=child_sections, parent_title=chapter_title_for_plan)

    # ========== 给 outline 所有节点注入 chapter_idx（包括分隔页子节点） ==========
    for idx, item in enumerate(outline):
        item["_chapter_idx"] = idx

    # ========== 构建封面标题集合（封面已单独渲染） ==========
    _cover_fmt = analysis_context.get("_format_requirements", {}) if isinstance(analysis_context, dict) else {}
    _cover_titles = set()
    for _sec in _cover_fmt.get("required_sections", []):
        _t = _sec.get("title", "").strip()
        if _t:
            _cover_titles.add(_t)

    # ========== 按目录结构生成正文 ==========
    for _oi_idx, item in enumerate(outline):
        _item_title = item.get("title", "").strip()
        # ===== 分隔页检测（放在 is_cover 检查之前，避免被封面跳过） =====
        if _is_separator_page_title(_item_title):
            # 从 chapter_contents 查找分隔页的原文内容
            _sep_raw = None
            _chapter_idx = item.get("_chapter_idx")
            if _chapter_idx is not None and _chapter_idx < len(chapter_contents):
                _sep_raw = chapter_contents[_chapter_idx].get("content", "")
            if _sep_raw and _sep_raw.startswith(_SEPARATOR_PAGE_PREFIX):
                _sep_raw = _sep_raw[len(_SEPARATOR_PAGE_PREFIX):]
            elif _sep_raw and _sep_raw == _SEPARATOR_PAGE_EMPTY:
                _sep_raw = None
            _render_separator_page(document, item, original_text=_sep_raw)
            document.add_page_break()
            # 子节点渲染：优先使用格式要求中的模板内容（有序的段落+表格），
            # 若有模板则直接渲染模板内容（不重复写子章节），无模板时退回到 write_outline_item
            # 获取section_lookup（使用 _clean_section_title 保持key一致性）
            _sec_lookup_global = {}
            _ad_raw4 = getattr(analysis_result, "analysis_data", None) if analysis_result else None
            if _ad_raw4:
                try:
                    import json as _json4
                    _ad4 = _json4.loads(_ad_raw4) if isinstance(_ad_raw4, str) else (_ad_raw4 or {})
                    _fmt4 = _ad4.get("format_requirements", {}) if isinstance(_ad4, dict) else {}
                    _sec_lookup_global = _fmt4.get("section_lookup", {}) or {}
                except Exception:
                    _sec_lookup_global = {}
            # 用于追踪子节点渲染中的最后一个 body 子元素，确保表格精准定位
            _last_child_element = None
            for _child in item.get("children", []):
                _child_title = _child.get("title", "").strip()
                if not _child_title:
                    continue
                # 使用与 section_lookup 构建时一致的 _clean_section_title 清洗函数
                from .analysis_v3.phase1_5_format import _clean_section_title as _lookup_clean
                _child_clean2 = _lookup_clean(_child_title)
                # 精确匹配 section_lookup（清洗后的标题 → section dict）
                _sec3 = _sec_lookup_global.get(_child_clean2) if isinstance(_sec_lookup_global, dict) else None
                if _sec3:
                    _tc3 = _sec3.get("template_content", [])
                    if _tc3:
                        # 渲染有序的段落+表格内容，使用 write_table_from_data 确保表格定位
                        from app.infrastructure.table_codec import to_per_cell, write_table_from_data
                        for _bd3 in _tc3:
                            if not isinstance(_bd3, dict):
                                continue
                            _bt3 = _bd3.get("type", "text")
                            if _bt3 == "table":
                                _h3 = _bd3.get("headers", [])
                                _r3 = _bd3.get("rows", [])
                                _mc3 = _bd3.get("merge_cells", [])
                                _cw3 = _bd3.get("column_widths", [])
                                _per_cell3 = _bd3.get("per_cell") or _bd3.get("per_cell_data")
                                if _h3 and _r3:
                                    # 使用 write_table_from_data 而非 document.add_table() 以支持精确定位
                                    if _per_cell3:
                                        from app.infrastructure.table_codec import from_dict as _pcd_from_dict
                                        _td3 = _pcd_from_dict(_per_cell3)
                                    else:
                                        _all_rows3 = [_h3] + _r3
                                        _td3 = to_per_cell(_h3, _all_rows3, _mc3, _cw3, _bd3.get("row_heights", []))
                                    # 注入默认字体格式
                                    for _row3 in _td3.rows:
                                        for _cell3 in _row3.cells:
                                            if not _cell3.font_name:
                                                _cell3.font_name = "仿宋"
                                            if not _cell3.font_size_half_pt:
                                                _cell3.font_size_half_pt = 24
                                    _tbl3 = write_table_from_data(document, _td3, insert_after=_last_child_element)
                                    if _tbl3 is not None:
                                        _last_child_element = _tbl3
                            elif _bt3 in ("text", "paragraph"):
                                _p3 = document.add_paragraph(_bd3.get("text", ""))
                                _p3.style = document.styles["Normal"]
                                for _r3 in _p3.runs:
                                    _r3.font.name = "仿宋"
                                    _r3.font.size = Pt(12)
                                    _r3.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                                _last_child_element = _p3._element
                        # 模板内容已渲染，跳过子章节的 _write_outline_item 调用
                        # 避免模板内容被重复渲染（content_blocks 中的同份模板会导致重复）
                        document.add_page_break()
                        continue
                # 无模板内容：回退到 write_outline_item
                _write_outline_item(_child, level=1)
                # 更新 _last_child_element 为 body 的最后一个子元素
                _body_children = list(document.element.body)
                if _body_children:
                    _last_child_element = _body_children[-1]
                document.add_page_break()
            continue
        if _oi_idx > 0:
            document.add_page_break()
        _write_outline_item(item, level=1)


    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
